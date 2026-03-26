import streamlit as st
import pandas as pd
from io import BytesIO
import re
from pathlib import Path
import base64
import openpyxl
from openpyxl import Workbook
from openpyxl.utils import get_column_letter
from openpyxl.styles import Alignment, Font, PatternFill, Border, Side
if "var_catalog" not in st.session_state or not isinstance(st.session_state.var_catalog, dict):
    st.session_state.var_catalog = {}


APP_DIR = Path(__file__).resolve().parent






#--------------------------
#Add logo
#--------------------

from pathlib import Path
import base64
import streamlit as st

APP_DIR = Path(__file__).resolve().parent

def render_global_header():
    logo_path = APP_DIR / "NPA.png"
    if not logo_path.exists():
        st.error(f"Logo not found: {logo_path}")
        return

    b64 = base64.b64encode(logo_path.read_bytes()).decode("utf-8")

    st.markdown(
        f"""
        <style>
        header[data-testid="stHeader"],
        div[data-testid="stHeader"] {{
            display: block !important;
            visibility: visible !important;
            height: 76px !important;

            background-image: url("data:image/png;base64,{b64}") !important;
            background-repeat: no-repeat !important;

            /* ⬇ move down ~50% of logo height */
            background-position: 3% 10% !important;


            /* ⬇ make logo ~20% smaller */
            background-size: 160px auto !important;
            background-size: 140px auto !important;
        }}

        header[data-testid="stHeader"] > div,
        div[data-testid="stHeader"] > div {{
            height: 86px !important;
        }}
        </style>
        """,
        unsafe_allow_html=True
    )


import re

def _spss_escape(s: str) -> str:
    s = "" if s is None else str(s)
    return s.replace("'", "''")

def _is_intish(x: str) -> bool:
    try:
        int(str(x).strip())
        return True
    except Exception:
        return False

def generate_spss_from_module2(questions: list, recodes: list) -> str:
    qmap = {str(q.get("qname", "")).strip(): q for q in (questions or [])}
    lines = []

    for r in recodes or []:
        src = str(r.get("source_qname", "")).strip()
        newv = str(r.get("new_qname", "")).strip()
        newlab = str(r.get("new_label", "")).strip()
        groups = r.get("groups", []) or []

        if not src or not newv:
            continue

        src_to_new = {}
        new_value_labels = {}
        used_src_codes = set()

        for g in groups:
            new_code = str(g.get("new_code", "")).strip()
            new_text = str(g.get("new_text", "")).strip()
            if not new_code or not new_text:
                continue
            new_value_labels[new_code] = new_text

            for f in g.get("from", []) or []:
                sc = str(f.get("code", "")).strip()
                if not sc:
                    continue
                used_src_codes.add(sc)
                src_to_new[sc] = new_code

        newcode_to_srccodes = {}
        for sc, nc in src_to_new.items():
            newcode_to_srccodes.setdefault(nc, []).append(sc)

        def _sort_key(k):
            return (0, int(k)) if _is_intish(k) else (1, str(k))

        def _fmt_code_list(code_list):
            out = []
            for c in code_list:
                c = str(c).strip()
                if _is_intish(c):
                    out.append(str(int(c)))
                else:
                    out.append(f"'{_spss_escape(c)}'")
            return ",".join(out)

        recode_parts = []
        for nc in sorted(newcode_to_srccodes.keys(), key=_sort_key):
            srclist = sorted(newcode_to_srccodes[nc], key=_sort_key)
            if _is_intish(nc):
                recode_parts.append(f"({_fmt_code_list(srclist)}={int(nc)})")
            else:
                recode_parts.append(f"({_fmt_code_list(srclist)}='{_spss_escape(nc)}')")

        lines.append(f"*{newv}.")
        if recode_parts:
            lines.append(f"RECODE {src} " + " ".join(recode_parts) + f" (ELSE=COPY) INTO {newv}.")
        else:
            lines.append(f"COMPUTE {newv} = {src}.")
        lines.append("EXECUTE.")

        if newlab:
            lines.append(f"VARIABLE LABELS  {newv} '{_spss_escape(newlab)}'.")

        src_choices = (qmap.get(src, {}) or {}).get("choices", []) or []
        carry_labels = {}
        for ch in src_choices:
            sc = str(ch.get("code", "")).strip()
            sl = str(ch.get("label", "")).strip()
            if not sc or not sl:
                continue
            if sc in used_src_codes:
                continue
            carry_labels[sc] = sl

        if new_value_labels or carry_labels:
            lines.append(f"VALUE LABELS {newv}")
            for nc in sorted(new_value_labels.keys(), key=_sort_key):
                lab = new_value_labels[nc]
                if _is_intish(nc):
                    lines.append(f"  {int(nc)} '{_spss_escape(lab)}'")
                else:
                    lines.append(f"  '{_spss_escape(nc)}' '{_spss_escape(lab)}'")
            for sc in sorted(carry_labels.keys(), key=_sort_key):
                lab = carry_labels[sc]
                if _is_intish(sc):
                    lines.append(f"  {int(sc)} '{_spss_escape(lab)}'")
                else:
                    lines.append(f"  '{_spss_escape(sc)}' '{_spss_escape(lab)}'")
            lines.append(".")

        lines.append("")

    return ("\n".join(lines).strip() + "\n") if lines else ""













# ============================
# Page
# ============================
st.set_page_config(page_title="Question Builder", layout="wide")
st.title("Polling Script ➜ Toplines & Crosstabs")

# ============================
# Global session init
# ============================
if "questions" not in st.session_state:
    st.session_state.questions = []  # module 1 saved variables

if "choices" not in st.session_state:
    st.session_state.choices = [{"label": "", "code": ""}]  # module 1 working choices list

if "form_version" not in st.session_state:
    st.session_state.form_version = 0

if "selected_q_index" not in st.session_state:
    st.session_state.selected_q_index = None  # module 1 selection

if "mode" not in st.session_state:
    st.session_state.mode = "new"  # module 1: "new" or "edit"

if "_defaults" not in st.session_state:
    st.session_state._defaults = {"qname": "", "prompt": "", "label": ""}

# Module 2 state
if "recodes" not in st.session_state:
    # list of recode variables, each:
    # {
    #   "source_qname": "...",
    #   "qname": "...",
    #   "prompt": "...",
    #   "label": "...",
    #   "groups": [
    #       {"new_label": "...", "new_code": "...", "source_pairs": [{"label": "...","code":"..."}...]}
    #   ]
    # }
    st.session_state.recodes = []

if "recode_selected_source" not in st.session_state:
    st.session_state.recode_selected_source = None

if "recode_work_groups" not in st.session_state:
    st.session_state.recode_work_groups = []  # current groups being built

if "recode_form_version" not in st.session_state:
    st.session_state.recode_form_version = 0


# ============================
# Shared helpers
# ============================
def to_excel_bytes_from_export_df(df: pd.DataFrame, sheet_name: str) -> bytes:
    output = BytesIO()
    with pd.ExcelWriter(output, engine="openpyxl") as writer:
        df.to_excel(writer, index=False, header=False, sheet_name=sheet_name)
    return output.getvalue()


# ============================
# Module selector (top)
# ============================
MODULES = {
    "Module 0: Project": "project",
    "Module 1: Scripting": "scripting",
    "Module 2: Recodes": "recodes",
    "Module 2.5: Derived Variables": "derived_vars",
    "Module 3: Import + Match Data": "import_match",
    "Module 4: Weighting": "weighting",
    "Module 5: Topline Shell": "topline_shell",
    "Module 6: Weighted Topline": "weighted_topline",
    "Module 7: Crosstabs": "crosstabs"
}

module_label = st.selectbox("Select module", list(MODULES.keys()))
active_module = MODULES[module_label]

render_global_header()
st.divider()








# =========================================================
# MODULE 1: SCRIPTING  (with optional Excel/Text import)
# Updated to publish variables into st.session_state.var_catalog
# so Module 2.5 can use the shared variable list.
# =========================================================

import re
import pandas as pd
import streamlit as st

# ----------------------------
# Shared var_catalog helpers
# ----------------------------
if "var_catalog" not in st.session_state:
    st.session_state.var_catalog = {}

def _vc_safe_str(x) -> str:
    if x is None:
        return ""
    return str(x).strip()

def _publish_to_var_catalog(var: str, label: str, choices_list: list, origin: str):
    """
    var_catalog[var] = {"label":..., "choices":{code:text}, "origin":..., "type":"single"}
    choices_list expects list of {"code":..., "label":...} or {"code":..., "text":...}
    """
    var = _vc_safe_str(var)
    if not var:
        return

    choices = {}
    for ch in (choices_list or []):
        c = _vc_safe_str(ch.get("code"))
        t = _vc_safe_str(ch.get("label")) or _vc_safe_str(ch.get("text"))
        if c and t:
            choices[c] = t

    st.session_state.var_catalog[var] = {
        "label": _vc_safe_str(label),
        "choices": choices,
        "origin": origin,
        "type": "single",
    }

def _remove_from_var_catalog_if_origin(var: str, origin: str):
    var = _vc_safe_str(var)
    cat = st.session_state.get("var_catalog", {}) or {}
    if var in cat and _vc_safe_str(cat[var].get("origin")) == origin:
        cat.pop(var, None)
        st.session_state.var_catalog = cat

def _sync_all_module1_questions_to_catalog():
    # Rebuild/refresh module1 entries (don’t touch other origins)
    # Remove old m1 entries first
    cat = st.session_state.get("var_catalog", {}) or {}
    to_del = [k for k, v in cat.items() if _vc_safe_str(v.get("origin")) == "m1"]
    for k in to_del:
        cat.pop(k, None)
    st.session_state.var_catalog = cat

    for q in (st.session_state.get("questions") or []):
        _publish_to_var_catalog(
            var=_vc_safe_str(q.get("qname")),
            label=_vc_safe_str(q.get("label")),
            choices_list=(q.get("choices") or []),
            origin="m1",
        )


# ----------------------------
# Regexes (module scope!)
# ----------------------------
_M1_Q_HEADER_RE = re.compile(r"^\s*(?:\d+\.\s*)?(Q[A-Z0-9_]+)\s*:\s*(.+?)\s*$")

# Battery header variants:
# QIMAGE_A-E: ...
# QIMAGE_A - E: ...
_M1_BATTERY_RE = re.compile(
    r"^\s*(Q[A-Z0-9_]+)_([A-Z]|\d+)\s*-\s*([A-Z]|\d+)\s*:\s*(.+?)\s*$"
)

# ----------------------------
# SPSS helpers (shared by Module 1 + 2)
# ----------------------------
def _spss_escape(s: str) -> str:
    s = "" if s is None else str(s)
    return s.replace("'", "''")

def _is_intish(x: str) -> bool:
    try:
        int(str(x).strip())
        return True
    except Exception:
        return False

def generate_spss_from_module1(questions: list) -> str:
    lines = []
    for q in questions or []:
        qname = str(q.get("qname", "")).strip()
        vlabel = str(q.get("label", "")).strip()
        choices = q.get("choices", []) or []

        if not qname or not vlabel or not choices:
            continue

        clean_choices = []
        for ch in choices:
            code = str(ch.get("code", "")).strip()
            lab = str(ch.get("label", "")).strip()
            if not code or not lab:
                continue
            clean_choices.append((code, lab))

        if not clean_choices:
            continue

        lines.append(f"*{qname}.")
        lines.append(f"VARIABLE LABELS  {qname} '{_spss_escape(vlabel)}'.")
        lines.append(f"VALUE LABELS {qname}")
        for code, lab in clean_choices:
            if _is_intish(code):
                lines.append(f"  {int(code)} '{_spss_escape(lab)}'")
            else:
                lines.append(f"  '{_spss_escape(code)}' '{_spss_escape(lab)}'")
        lines.append(".")
        lines.append("")

    return ("\n".join(lines).strip() + "\n") if lines else ""


# ----------------------------
# Helpers
# ----------------------------
def _m1_strip_bracket_tags(s: str) -> str:
    s = "" if s is None else str(s)
    s = re.sub(r"\[[^\]]+\]", "", s)  # remove [FLIP], [ROTATE], etc.
    s = re.sub(r"\s+", " ", s).strip()
    return s

def _m1_is_terminate(val: str) -> bool:
    if val is None:
        return False
    return "TERMINATE" in str(val).upper()

def _m1_is_noise_line(text: str) -> bool:
    if text is None:
        return True
    t = str(text).strip()
    if t == "":
        return True
    if re.fullmatch(r"\[[^\]]+\]", t):
        return True
    if t.upper().startswith("IF ") and "GO TO" in t.upper():
        return True
    if ":" not in t and len(t) <= 40 and not re.search(r"\d", t) and re.fullmatch(r"[A-Za-z][A-Za-z\s/’'-]*", t):
        return True
    return False

def _m1_extract_choice_row(text_cell, value_cell):
    label = "" if text_cell is None else str(text_cell).strip()
    code  = "" if value_cell is None else str(value_cell).strip()

    if label == "" and code == "":
        return None

    if _m1_is_terminate(code) or _m1_is_terminate(label) or re.search(r"\bTERMINATE\b", code, flags=re.I):
        return None

    m = re.match(r"^\s*(\d+)\b", code)
    if not m:
        return None
    code_num = m.group(1)

    label = _m1_strip_bracket_tags(label)
    if label == "":
        return None

    return {"label": label, "code": code_num}

def parse_script_rows_to_questions(rows):
    questions = []
    report = {
        "questions_added": 0,
        "choices_added": 0,
        "rows_skipped_noise": 0,
        "rows_skipped_terminate_or_bad": 0,
        "battery_blocks_used": 0,
    }

    current_q = None

    battery = {
        "active": False,
        "prefix": None,
        "start": None,
        "end": None,
        "prompt": None,
        "choices": [],
    }

    def flush_current():
        nonlocal current_q
        if current_q and current_q.get("choices"):
            questions.append(current_q)
            report["questions_added"] += 1
            report["choices_added"] += len(current_q["choices"])
        current_q = None

    def _token_in_range(token: str, start: str, end: str) -> bool:
        token = "" if token is None else str(token).strip()
        start = "" if start is None else str(start).strip()
        end   = "" if end is None else str(end).strip()
        if not token or not start or not end:
            return False

        # numeric range
        if token.isdigit() and start.isdigit() and end.isdigit():
            try:
                t, s, e = int(token), int(start), int(end)
                lo, hi = (s, e) if s <= e else (e, s)
                return lo <= t <= hi
            except Exception:
                return False

        # letter range
        token_u, start_u, end_u = token.upper(), start.upper(), end.upper()
        if (
            len(token_u) == 1 and len(start_u) == 1 and len(end_u) == 1
            and token_u.isalpha() and start_u.isalpha() and end_u.isalpha()
        ):
            lo, hi = (start_u, end_u) if start_u <= end_u else (end_u, start_u)
            return ord(lo) <= ord(token_u) <= ord(hi)

        return False

    def _parse_battery_groups(m_b):
        g = list(m_b.groups())
        prefix = g[0]
        start_tok = g[1]
        end_tok = g[2]
        prompt = g[3]
        return prefix, start_tok, end_tok, prompt

    for r in rows:
        text = "" if r.get("Text") is None else str(r.get("Text")).strip()
        val = "" if r.get("Value") is None else str(r.get("Value")).strip()

        if _m1_is_noise_line(text) and val == "":
            report["rows_skipped_noise"] += 1
            continue

        m_b = _M1_BATTERY_RE.match(text)
        if m_b:
            flush_current()
            prefix, start_letter, end_letter, prompt = _parse_battery_groups(m_b)
            battery["active"] = True
            battery["prefix"] = prefix
            battery["start"] = start_letter
            battery["end"] = end_letter
            battery["prompt"] = _m1_strip_bracket_tags(prompt)
            battery["choices"] = []
            report["battery_blocks_used"] += 1
            continue

        m_q = _M1_Q_HEADER_RE.match(text)
        if m_q:
            qname, prompt = m_q.groups()
            qname = qname.strip()
            prompt_clean = _m1_strip_bracket_tags(prompt)

            flush_current()

            if battery["active"] and qname.startswith(battery["prefix"] + "_"):
                parts = qname.split("_")
                suffix = parts[-1] if parts else None
                if suffix and _token_in_range(suffix, battery["start"], battery["end"]):
                    current_q = {
                        "qname": qname,
                        "prompt": prompt_clean,
                        "label": "",
                        "choices": [dict(c) for c in battery["choices"]],
                    }
                    continue

            battery["active"] = False
            battery["prefix"] = None
            battery["start"] = None
            battery["end"] = None
            battery["prompt"] = None
            battery["choices"] = []

            current_q = {"qname": qname, "prompt": prompt_clean, "label": "", "choices": []}
            continue

        if battery["active"] and current_q is None:
            ch = _m1_extract_choice_row(text, val)
            if ch:
                if ch["code"] not in {x["code"] for x in battery["choices"]}:
                    battery["choices"].append(ch)
            else:
                if text or val:
                    report["rows_skipped_terminate_or_bad"] += 1
            continue

        if current_q is not None:
            ch = _m1_extract_choice_row(text, val)
            if ch:
                if ch["code"] not in {x["code"] for x in current_q["choices"]}:
                    current_q["choices"].append(ch)
            else:
                if text or val:
                    report["rows_skipped_terminate_or_bad"] += 1
            continue

        report["rows_skipped_noise"] += 1

    flush_current()
    return questions, report

def parse_text_script_to_rows(script_text: str):
    rows = []
    for raw in (script_text or "").splitlines():
        line = raw.rstrip("\n").strip()
        if line == "":
            rows.append({"Text": "", "Value": ""})
            continue

        if "\t" in line:
            parts = [p.strip() for p in line.split("\t") if p.strip() != ""]
            if len(parts) >= 2:
                rows.append({"Text": parts[0], "Value": parts[1]})
                continue

        parts = re.split(r"\s{2,}", line)
        parts = [p.strip() for p in parts if p.strip() != ""]
        if len(parts) >= 2:
            rows.append({"Text": parts[0], "Value": parts[1]})
        else:
            rows.append({"Text": line, "Value": ""})

    return rows

def parse_excel_script_to_rows(uploaded_file):
    df = pd.read_excel(uploaded_file, sheet_name=0, header=None)
    if df.shape[1] >= 2:
        rows = [{"Text": df.iloc[i, 0], "Value": df.iloc[i, 1]} for i in range(len(df))]
    else:
        rows = [{"Text": df.iloc[i, 0], "Value": ""} for i in range(len(df))]
    return rows


# ----------------------------
# Module 1 UI helpers
# ----------------------------
def _m1_bump_form_version():
    st.session_state.form_version = st.session_state.get("form_version", 0) + 1

def module1_add_choice():
    if "choices" not in st.session_state or not isinstance(st.session_state.choices, list):
        st.session_state.choices = [{"label": "", "code": ""}]
    st.session_state.choices.append({"label": "", "code": ""})

def module1_remove_last_choice():
    if "choices" not in st.session_state or not isinstance(st.session_state.choices, list):
        st.session_state.choices = [{"label": "", "code": ""}]
        return
    if len(st.session_state.choices) > 1:
        st.session_state.choices.pop()

def module1_reset_form():
    st.session_state.choices = [{"label": "", "code": ""}]
    _m1_bump_form_version()

def module1_load_question_into_form(q: dict):
    _m1_bump_form_version()
    choices = [
        {"label": str(c.get("label", "")).strip(), "code": str(c.get("code", "")).strip()}
        for c in (q.get("choices", []) or [])
    ]
    st.session_state.choices = choices if choices else [{"label": "", "code": ""}]
    return (
        str(q.get("qname", "")).strip(),
        str(q.get("prompt", "")).strip(),
        str(q.get("label", "")).strip(),
    )

def module1_validate_and_clean(qname: str, prompt: str, label: str):
    cleaned = [
        {"label": str(c.get("label", "")).strip(), "code": str(c.get("code", "")).strip()}
        for c in (st.session_state.get("choices") or [])
        if str(c.get("label", "")).strip() and str(c.get("code", "")).strip()
    ]

    if not str(qname).strip():
        return None, "Question Name is required."
    if not str(prompt).strip():
        return None, "Question prompt is required."
    if not str(label).strip():
        return None, "Question label is required."
    if len(cleaned) == 0:
        return None, "At least one choice with a value code is required."

    codes = [c["code"] for c in cleaned]
    if len(codes) != len(set(codes)):
        return None, "Choice codes must be unique."

    return cleaned, None

def module1_build_export_df():
    rows = []
    for idx, q in enumerate(st.session_state.get("questions", []), start=1):
        qname = str(q.get("qname", "")).strip()
        prompt = str(q.get("prompt", "")).strip()
        header = f"{idx}.    {qname}: {prompt}"
        rows.append({"Text": header, "Value": ""})
        for ch in (q.get("choices", []) or []):
            rows.append({"Text": str(ch.get("label", "")).strip(), "Value": str(ch.get("code", "")).strip()})
        rows.append({"Text": "", "Value": ""})
    return pd.DataFrame(rows, columns=["Text", "Value"])


def render_module_1():
    # ----------------------------
    # Ensure session keys exist
    # ----------------------------
    if "questions" not in st.session_state:
        st.session_state.questions = []
    if "choices" not in st.session_state:
        st.session_state.choices = [{"label": "", "code": ""}]
    if "selected_q_index" not in st.session_state:
        st.session_state.selected_q_index = None
    if "mode" not in st.session_state:
        st.session_state.mode = "new"
    if "_defaults" not in st.session_state:
        st.session_state._defaults = {"qname": "", "prompt": "", "label": ""}
    if "form_version" not in st.session_state:
        st.session_state.form_version = 0
    if "m1_last_import_hash" not in st.session_state:
        st.session_state.m1_last_import_hash = None

    st.subheader("Optional: Start from an uploaded script")

    with st.expander("📥 Upload Excel script (or paste text) to pre-fill Module 1", expanded=False):
        colA, colB = st.columns([2, 1])

        with colA:
            up = st.file_uploader("Upload .xlsx script", type=["xlsx"], key="m1_script_upload")
            pasted = st.text_area("Or paste script text", value="", height=180, key="m1_script_paste")

        with colB:
            replace = st.checkbox("Replace existing questions", value=True, key="m1_import_replace")
            do_import = st.button("Import into Module 1", key="m1_import_run")

        if do_import:
            rows = None
            source_sig = None

            if up is not None:
                rows = parse_excel_script_to_rows(up)
                source_sig = f"xlsx:{up.name}:{up.size}"
            elif pasted.strip():
                rows = parse_text_script_to_rows(pasted)
                source_sig = f"txt:{hash(pasted)}"
            else:
                st.error("Upload an Excel file or paste script text first.")

            if rows is not None:
                if (not replace) and source_sig and st.session_state.m1_last_import_hash == source_sig:
                    st.warning("That looks like the same import as last time. (Skipped to avoid duplicates.)")
                else:
                    imported_questions, report = parse_script_rows_to_questions(rows)

                    if not imported_questions:
                        st.error("Import ran, but no valid questions were found.")
                    else:
                        if replace:
                            st.session_state.questions = imported_questions
                            st.session_state.selected_q_index = 0 if imported_questions else None
                        else:
                            st.session_state.questions.extend(imported_questions)
                            st.session_state.selected_q_index = len(st.session_state.questions) - len(imported_questions)

                        st.session_state.m1_last_import_hash = source_sig

                        st.session_state.mode = "new"
                        st.session_state._defaults = {"qname": "", "prompt": "", "label": ""}
                        module1_reset_form()

                        # --- NEW: publish to var_catalog ---
                        _sync_all_module1_questions_to_catalog()

                        st.success(
                            f"Imported {report.get('questions_added', 0)} questions "
                            f"({report.get('choices_added', 0)} choices). "
                            f"Battery blocks used: {report.get('battery_blocks_used', 0)}."
                        )
                        st.caption(
                            f"Skipped noise rows: {report.get('rows_skipped_noise', 0)} | "
                            f"Skipped terminate/bad rows: {report.get('rows_skipped_terminate_or_bad', 0)}"
                        )
                        st.rerun()

    st.divider()

    # ----------------------------
    # Sidebar: variable list + actions
    # ----------------------------
    st.sidebar.header("Variables coded")

    if len(st.session_state.questions) == 0:
        st.sidebar.info("No variables yet.")
    else:
        q_idx_list = list(range(len(st.session_state.questions)))

        def _fmt_q(i):
            q = st.session_state.questions[i]
            return f"{i+1}. {q.get('qname','')} — {q.get('label','')}"

        sel_index = st.session_state.selected_q_index
        if sel_index is None or sel_index < 0 or sel_index >= len(q_idx_list):
            sel_index = 0

        selected_i = st.sidebar.radio(
            "Select a variable",
            options=q_idx_list,
            index=sel_index,
            format_func=_fmt_q,
            key="m1_radio",
        )
        st.session_state.selected_q_index = selected_i
        q_selected = st.session_state.questions[selected_i]

        st.sidebar.divider()
        cA, cB, cC = st.sidebar.columns(3)

        if cA.button("Duplicate", use_container_width=True, key="m1_dup"):
            dup = {
                "qname": str(q_selected.get("qname", "")).strip() + "_COPY",
                "prompt": str(q_selected.get("prompt", "")).strip(),
                "label": str(q_selected.get("label", "")).strip(),
                "choices": [dict(x) for x in (q_selected.get("choices", []) or [])],
            }
            st.session_state.questions.append(dup)
            st.session_state.selected_q_index = len(st.session_state.questions) - 1

            # --- NEW: publish dup to var_catalog ---
            _publish_to_var_catalog(dup["qname"], dup["label"], dup["choices"], origin="m1")

            st.session_state.mode = "edit"
            qname_default, prompt_default, label_default = module1_load_question_into_form(dup)
            st.session_state._defaults = {"qname": qname_default, "prompt": prompt_default, "label": label_default}
            st.rerun()

        if cB.button("Edit", use_container_width=True, key="m1_edit"):
            st.session_state.mode = "edit"
            qname_default, prompt_default, label_default = module1_load_question_into_form(q_selected)
            st.session_state._defaults = {"qname": qname_default, "prompt": prompt_default, "label": label_default}
            st.rerun()

        if cC.button("Delete", use_container_width=True, key="m1_del"):
            old_name = str(q_selected.get("qname", "")).strip()
            st.session_state.questions.pop(selected_i)

            # --- NEW: remove from var_catalog if it was a module1 var ---
            _remove_from_var_catalog_if_origin(old_name, origin="m1")

            st.session_state.selected_q_index = None if len(st.session_state.questions) == 0 else 0
            st.session_state.mode = "new"
            st.session_state._defaults = {"qname": "", "prompt": "", "label": ""}
            module1_reset_form()
            st.sidebar.success("Deleted.")
            st.rerun()

    st.sidebar.divider()
    if st.sidebar.button("➕ New variable", use_container_width=True, key="m1_new"):
        st.session_state.mode = "new"
        st.session_state.selected_q_index = None
        st.session_state._defaults = {"qname": "", "prompt": "", "label": ""}
        module1_reset_form()
        st.rerun()

    defaults = st.session_state._defaults or {"qname": "", "prompt": "", "label": ""}

    # ----------------------------
    # Main form
    # ----------------------------
    st.subheader("Build / Edit a variable")

    if st.session_state.mode == "edit" and st.session_state.selected_q_index is not None:
        try:
            cur_q = st.session_state.questions[st.session_state.selected_q_index]
            st.caption(f"Editing: **{cur_q.get('qname','')}**")
        except Exception:
            st.caption("Editing: (selection invalid)")
    else:
        st.caption("Creating: **New variable**")

    form_key = f"script_form_{st.session_state.form_version}"
    with st.form(form_key):
        qname = st.text_input("Question Name", value=defaults.get("qname", ""), placeholder="QAGE")
        prompt = st.text_area("Question prompt", value=defaults.get("prompt", ""), placeholder="What is your age?")
        label = st.text_input("Question label", value=defaults.get("label", ""), placeholder="Age")

        st.subheader("Question choices")
        if not st.session_state.choices:
            st.session_state.choices = [{"label": "", "code": ""}]

        for i, choice in enumerate(st.session_state.choices):
            col1, col2 = st.columns([3, 1])
            with col1:
                st.session_state.choices[i]["label"] = st.text_input(
                    f"Choice {i+1} label",
                    value=str(choice.get("label", "")),
                    key=f"{form_key}_label_{i}",
                    placeholder="18–24",
                )
            with col2:
                st.session_state.choices[i]["code"] = st.text_input(
                    "Code",
                    value=str(choice.get("code", "")),
                    key=f"{form_key}_code_{i}",
                    placeholder="2",
                )

        a1, a2, a3, a4 = st.columns([1, 1, 1, 2])
        with a1:
            add_clicked = st.form_submit_button("➕ Add choice")
        with a2:
            remove_clicked = st.form_submit_button("➖ Remove last choice")
        with a3:
            clear_clicked = st.form_submit_button("🧹 Clear")
        with a4:
            save_clicked = st.form_submit_button("💾 Save (new / update)")

    if add_clicked:
        module1_add_choice()
        st.rerun()

    if remove_clicked:
        module1_remove_last_choice()
        st.rerun()

    if clear_clicked:
        st.session_state.mode = "new"
        st.session_state.selected_q_index = None
        st.session_state._defaults = {"qname": "", "prompt": "", "label": ""}
        module1_reset_form()
        st.rerun()

    if save_clicked:
        cleaned_choices, err = module1_validate_and_clean(qname, prompt, label)
        if err:
            st.error(err)
        else:
            payload = {
                "qname": str(qname).strip(),
                "prompt": str(prompt).strip(),
                "label": str(label).strip(),
                "choices": cleaned_choices,
            }

            # If editing AND qname changed, remove old name from catalog if it was m1
            if (
                st.session_state.mode == "edit"
                and st.session_state.selected_q_index is not None
                and 0 <= st.session_state.selected_q_index < len(st.session_state.questions)
            ):
                old = st.session_state.questions[st.session_state.selected_q_index]
                old_name = str(old.get("qname", "")).strip()
                new_name = payload["qname"]

                st.session_state.questions[st.session_state.selected_q_index] = payload
                st.success("Updated.")

                if old_name and old_name != new_name:
                    _remove_from_var_catalog_if_origin(old_name, origin="m1")

                _publish_to_var_catalog(new_name, payload["label"], payload["choices"], origin="m1")
            else:
                st.session_state.questions.append(payload)
                st.session_state.selected_q_index = len(st.session_state.questions) - 1
                st.success("Saved.")
                _publish_to_var_catalog(payload["qname"], payload["label"], payload["choices"], origin="m1")

            st.session_state._defaults = {"qname": "", "prompt": "", "label": ""}
            st.session_state.mode = "new"
            module1_reset_form()
            st.rerun()

    # ----------------------------
    # Export preview (Excel)
    # ----------------------------
    st.divider()
    st.subheader("Saved questions (export preview)")

    if len(st.session_state.questions) == 0:
        st.info("No questions saved yet.")
    else:
        export_df = module1_build_export_df()
        st.dataframe(export_df, use_container_width=True, hide_index=True)

        st.download_button(
            label="⬇️ Download Excel (SPSS-style)",
            data=to_excel_bytes_from_export_df(export_df, sheet_name="Questions"),
            file_name="questions_spss_style.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        )

        if st.button("🗑️ Delete all saved questions", key="m1_delete_all"):
            # remove m1 entries
            cat = st.session_state.get("var_catalog", {}) or {}
            to_del = [k for k, v in cat.items() if _vc_safe_str(v.get("origin")) == "m1"]
            for k in to_del:
                cat.pop(k, None)
            st.session_state.var_catalog = cat

            st.session_state.questions = []
            st.session_state.selected_q_index = None
            st.session_state.mode = "new"
            st.session_state._defaults = {"qname": "", "prompt": "", "label": ""}
            module1_reset_form()
            st.success("Cleared.")
            st.rerun()

    # ----------------------------
    # SPSS syntax output (Module 1)
    # ----------------------------
    st.divider()
    st.subheader("SPSS syntax (Module 1)")

    spss_m1 = generate_spss_from_module1(st.session_state.questions)
    if not spss_m1.strip():
        st.info("No SPSS syntax yet — each question must have a Question label and at least one choice.")
    else:
        st.code(spss_m1, language="spss")
        st.download_button(
            "⬇️ Download SPSS syntax (Module 1)",
            data=spss_m1.encode("utf-8"),
            file_name="module1_questions.sps",
            mime="text/plain",
            key="m1_download_spss",
        )


# ============================
# Module 2: Recodes
# Updated to publish recode vars into st.session_state.var_catalog
# ============================

def _safe_int(x):
    try:
        return int(str(x).strip())
    except Exception:
        return None

# ---- Module 2 session init (safe if already defined) ----
if "recodes" not in st.session_state:
    st.session_state.recodes = []

if "m2_selected_recode_index" not in st.session_state:
    st.session_state.m2_selected_recode_index = None

if "m2_mode" not in st.session_state:
    st.session_state.m2_mode = "new"

if "m2_ui_version" not in st.session_state:
    st.session_state.m2_ui_version = 0

if "m2_pick" not in st.session_state:
    st.session_state.m2_pick = set()

if "m2_defaults" not in st.session_state:
    st.session_state.m2_defaults = {
        "source_qname": "",
        "new_qname": "",
        "new_label": "",
        "group_text": "",
        "group_code": "",
    }

if "m2_work_groups" not in st.session_state:
    st.session_state.m2_work_groups = []

if "m2_last_sidebar_sel" not in st.session_state:
    st.session_state.m2_last_sidebar_sel = None

if "m2_last_source_qname" not in st.session_state:
    st.session_state.m2_last_source_qname = None


if "m2_preserve_groups_on_source_change" not in st.session_state:
    # When True, changing source_qname will NOT wipe existing recode groups (useful for duplicating batteries).
    st.session_state.m2_preserve_groups_on_source_change = False


def _module2_get_from_catalog(varname: str):
    cat = st.session_state.get("var_catalog", {}) or {}
    varname = "" if varname is None else str(varname).strip()
    return cat.get(varname)

def _module2_find_source_question(qname: str):
    """Backward-compatible helper (Module 2 originally only used Module 1 questions).
    Prefer var_catalog when available; fall back to Module 1 questions list.
    """
    qname = "" if qname is None else str(qname).strip()
    hit = _module2_get_from_catalog(qname)
    if hit is not None:
        return hit
    return next((q for q in st.session_state.get("questions", []) if q.get("qname") == qname), None)

def _module2_publish_recode_to_catalog(payload: dict):
    """
    Create choices for the recoded variable:
    - new_code -> new_text for each group
    - plus carry-over source choices not used (because ELSE=COPY semantics)
    """
    src = str(payload.get("source_qname", "")).strip()
    newv = str(payload.get("new_qname", "")).strip()
    newlab = str(payload.get("new_label", "")).strip()
    groups = payload.get("groups", []) or []

    if not newv:
        return

    used_src_codes = set()
    choices = {}

    # group labels
    for g in groups:
        nc = str(g.get("new_code", "")).strip()
        nt = str(g.get("new_text", "")).strip()
        if nc and nt:
            choices[nc] = nt
        for f in (g.get("from") or []):
            sc = str(f.get("code", "")).strip()
            if sc:
                used_src_codes.add(sc)

    # carryover (unused source codes)
    src_q = _module2_find_source_question(src)
    src_choices_dict = (src_q.get("choices", {}) if src_q else {}) or {}

    for sc, sl in src_choices_dict.items():
        sc = str(sc).strip()
        sl = str(sl).strip()
        if sc and sl and sc not in used_src_codes:
            if sc not in choices:
                choices[sc] = sl

    st.session_state.var_catalog[newv] = {
        "label": newlab,
        "choices": choices,
        "origin": "m2",
        "type": "single",
    }

def _module2_remove_recode_from_catalog_if_m2(varname: str):
    _remove_from_var_catalog_if_origin(varname, origin="m2")



def _module2_reconcile_work_groups_to_source(source_qname: str):
    """When swapping the source variable, keep existing recode groups but:
    - Update 'from' labels to match the new source's labels (by code)
    - Drop any 'from' codes that don't exist in the new source
    - Drop any groups that end up empty after filtering
    """
    src_q = _module2_find_source_question(source_qname)
    src_choices_dict = (src_q.get("choices", {}) if src_q else {}) or {}
    # normalize to dict[str,str]
    if isinstance(src_choices_dict, list):
        tmp = {}
        for ch in src_choices_dict:
            if isinstance(ch, dict):
                c = str(ch.get("code", "")).strip()
                l = str(ch.get("label", "")).strip()
                if c:
                    tmp[c] = l
        src_choices_dict = tmp
    else:
        src_choices_dict = {str(k).strip(): str(v).strip() for k, v in (src_choices_dict or {}).items() if str(k).strip()}

    new_groups = []
    for g in (st.session_state.get("m2_work_groups") or []):
        kept_from = []
        for f in (g.get("from") or []):
            c = str(f.get("code", "")).strip()
            if not c:
                continue
            if c not in src_choices_dict:
                continue
            kept_from.append({"code": c, "label": src_choices_dict.get(c, "")})
        if kept_from:
            new_groups.append({
                "new_text": g.get("new_text", ""),
                "new_code": g.get("new_code", ""),
                "from": kept_from,
            })

    st.session_state.m2_work_groups = new_groups


def build_module2_export_df():
    rows = []
    recodes = st.session_state.get("recodes", [])

    for rec in recodes:
        source = str(rec.get("source_qname", "")).strip()
        new_q = str(rec.get("new_qname", "")).strip()
        if not source or not new_q:
            continue

        src_q = _module2_get_from_catalog(source)
        src_choices_obj = (src_q.get("choices") if isinstance(src_q, dict) else None) if src_q else None
        # Normalize choices into a dict of {code: label}
        src_choices_dict = {}
        if isinstance(src_choices_obj, dict):
            src_choices_dict = {str(k).strip(): str(v).strip() for k, v in src_choices_obj.items() if str(k).strip()}
        elif isinstance(src_choices_obj, list):
            for ch in src_choices_obj:
                if isinstance(ch, dict):
                    c = str(ch.get("code", "")).strip()
                    l = str(ch.get("label", "")).strip()
                    if c:
                        src_choices_dict[c] = l
                elif isinstance(ch, (list, tuple)) and len(ch) == 2:
                    c = str(ch[0]).strip()
                    l = str(ch[1]).strip()
                    if c:
                        src_choices_dict[c] = l
        # else: no choices

        used_source_codes = set()
        groups = rec.get("groups", [])
        for g in groups:
            for f in g.get("from", []):
                c = str(f.get("code", "")).strip()
                if c:
                    used_source_codes.add(c)

        out_items = []
        for g in groups:
            txt = str(g.get("new_text", "")).strip()
            code = str(g.get("new_code", "")).strip()
            if txt and code:
                out_items.append({"Text": txt, "Value": code})
        for ch_code, ch_label in (src_choices_dict or {}).items():
            ch_code = str(ch_code).strip()
            ch_label = str(ch_label).strip()
            if not ch_code:
                continue
            if ch_code in used_source_codes:
                continue
            if not ch_label:
                # fall back to code as label if label is missing
                ch_label = ch_code
            out_items.append({"Text": ch_label, "Value": ch_code})

        out_items.sort(
            key=lambda x: (
                _safe_int(x["Value"]) if _safe_int(x["Value"]) is not None else 10**9,
                x["Text"],
            )
        )

        rows.append({"Text": new_q, "Value": ""})
        rows.extend(out_items)
        rows.append({"Text": "", "Value": ""})

    return pd.DataFrame(rows, columns=["Text", "Value"])


def _module2_load_into_ui(rec: dict, as_new: bool):
    st.session_state.m2_pick = set()

    st.session_state.m2_defaults = {
        "source_qname": rec.get("source_qname", ""),
        "new_qname": rec.get("new_qname", ""),
        "new_label": rec.get("new_label", ""),
        "group_text": "",
        "group_code": "",
    }

    st.session_state.m2_work_groups = [
        {
            "new_text": g.get("new_text", ""),
            "new_code": g.get("new_code", ""),
            "from": [dict(x) for x in g.get("from", [])],
        }
        for g in rec.get("groups", [])
    ]

    if as_new:
        st.session_state.m2_mode = "new"
        st.session_state.m2_selected_recode_index = None
    else:
        st.session_state.m2_mode = "edit"
        st.session_state.m2_selected_recode_index = st.session_state.recodes.index(rec)

    st.session_state.m2_last_source_qname = rec.get("source_qname", "")
    st.session_state.m2_ui_version += 1


def render_module_2():
    st.header("Module 2: Recodes")

    cat = st.session_state.get("var_catalog", {}) or {}
    if len(cat) == 0:
        st.info("No variables yet. Add variables in Module 1 (or create derived vars in Module 2.5) first.")
        return

    st.sidebar.header("Recodes")

    if len(st.session_state.recodes) == 0:
        st.sidebar.info("No recodes yet.")
    else:
        rec_opts = [
            f"{i+1}. {r.get('new_qname','')}  (from {r.get('source_qname','')})"
            for i, r in enumerate(st.session_state.recodes)
        ]
        sel_idx = 0 if st.session_state.m2_selected_recode_index is None else st.session_state.m2_selected_recode_index

        sel = st.sidebar.radio("Select a recode", rec_opts, index=sel_idx, key="m2_recode_radio")
        new_index = rec_opts.index(sel)

        if st.session_state.m2_last_sidebar_sel != new_index:
            st.session_state.m2_last_sidebar_sel = new_index
            rec_selected = st.session_state.recodes[new_index]
            _module2_load_into_ui(rec_selected, as_new=False)
            st.rerun()

        rec_selected = st.session_state.recodes[st.session_state.m2_last_sidebar_sel]

        st.sidebar.divider()
        cA, cB, cC = st.sidebar.columns(3)

        if cA.button("Duplicate", use_container_width=True, key="m2_dup"):
            st.session_state.m2_preserve_groups_on_source_change = True
            dup = {

                "source_qname": rec_selected.get("source_qname", ""),
                "new_qname": rec_selected.get("new_qname", ""),
                "new_label": rec_selected.get("new_label", ""),
                "groups": [dict(g) for g in rec_selected.get("groups", [])],
            }
            _module2_load_into_ui(dup, as_new=True)
            st.rerun()

        if cB.button("Edit", use_container_width=True, key="m2_edit"):
            st.session_state.m2_preserve_groups_on_source_change = False
            _module2_load_into_ui(rec_selected, as_new=False)
            st.rerun()

        if cC.button("Delete", use_container_width=True, key="m2_del"):
            del_idx = st.session_state.m2_last_sidebar_sel
            old_newv = str(st.session_state.recodes[del_idx].get("new_qname", "")).strip()

            st.session_state.recodes.pop(del_idx)

            # --- NEW: remove from var_catalog if origin m2 ---
            _module2_remove_recode_from_catalog_if_m2(old_newv)

            st.session_state.m2_last_sidebar_sel = None
            st.session_state.m2_selected_recode_index = None
            st.session_state.m2_mode = "new"
            st.session_state.m2_defaults = {
                "source_qname": "",
                "new_qname": "",
                "new_label": "",
                "group_text": "",
                "group_code": "",
            }
            st.session_state.m2_work_groups = []
            st.session_state.m2_pick = set()
            st.session_state.m2_last_source_qname = None
            st.session_state.m2_ui_version += 1
            st.sidebar.success("Deleted.")
            st.rerun()

    st.sidebar.divider()
    if st.sidebar.button("➕ New recode", use_container_width=True, key="m2_new"):
        st.session_state.m2_preserve_groups_on_source_change = False
        st.session_state.m2_mode = "new"
        st.session_state.m2_selected_recode_index = None
        st.session_state.m2_last_sidebar_sel = None
        st.session_state.m2_pick = set()
        st.session_state.m2_work_groups = []
        st.session_state.m2_defaults = {
            "source_qname": "",
            "new_qname": "",
            "new_label": "",
            "group_text": "",
            "group_code": "",
        }
        st.session_state.m2_last_source_qname = None
        st.session_state.m2_ui_version += 1
        st.rerun()
    cat = st.session_state.get("var_catalog", {}) or {}

    # Only show variables that actually have choices (needed to recode)
    def _m2_has_choices(vname: str) -> bool:
        d = cat.get(vname, {}) or {}
        choices = d.get("choices", {}) or {}
        return isinstance(choices, dict) and len(choices) > 0

    qnames = sorted([v for v in cat.keys() if _m2_has_choices(v)])

    if not qnames:
        st.info("No variables with value choices yet (nothing to recode).")
        return


    default_source = st.session_state.m2_defaults.get("source_qname") or qnames[0]
    if default_source not in qnames:
        default_source = qnames[0]

    source_qname = st.selectbox(
        "Select source variable to recode",
        qnames,
        index=qnames.index(default_source),
        key=f"m2_source_select_{st.session_state.m2_ui_version}",
    )

    src_q = _module2_find_source_question(source_qname)
    src_choices_dict = (src_q.get("choices", {}) if src_q else {}) or {}

    # Convert dict -> list[{"label","code"}] to match existing UI loops
    src_choices = [{"label": lab, "code": code} for code, lab in src_choices_dict.items()]

    prev_source = st.session_state.m2_last_source_qname
    if prev_source is None:
        prev_source = default_source

    if source_qname != prev_source:
        st.session_state.m2_defaults["source_qname"] = source_qname

        old_auto_qname = f"c{prev_source}" if prev_source else ""
        old_auto_label = ""
        prev_q = _module2_find_source_question(prev_source) if prev_source else None
        if prev_q:
            old_auto_label = str(prev_q.get("label", "")).strip()

        cur_new_qname = str(st.session_state.m2_defaults.get("new_qname", "")).strip()
        cur_new_label = str(st.session_state.m2_defaults.get("new_label", "")).strip()

        allow_qname_autofill = (cur_new_qname == "" or cur_new_qname == old_auto_qname)
        allow_label_autofill = (cur_new_label == "" or cur_new_label == old_auto_label)

        if allow_qname_autofill:
            st.session_state.m2_defaults["new_qname"] = f"c{source_qname}"

        if allow_label_autofill:
            st.session_state.m2_defaults["new_label"] = str(src_q.get("label", "")).strip() if src_q else ""

        st.session_state.m2_pick = set()
        if st.session_state.get("m2_preserve_groups_on_source_change", False):
            _module2_reconcile_work_groups_to_source(source_qname)
        else:
            st.session_state.m2_work_groups = []

        st.session_state.m2_last_source_qname = source_qname
        st.session_state.m2_ui_version += 1
        st.rerun()
    else:
        st.session_state.m2_last_source_qname = source_qname

    st.subheader("Create / edit a recode")

    new_qname = st.text_input(
        "New variable name (recoded)",
        value=st.session_state.m2_defaults.get("new_qname", f"c{source_qname}"),
        key=f"m2_new_qname_{st.session_state.m2_ui_version}",
    )
    new_label = st.text_input(
        "New variable label",
        value=st.session_state.m2_defaults.get(
            "new_label",
            str(src_q.get("label", "")).strip() if src_q else "",
        ),
        key=f"m2_new_label_{st.session_state.m2_ui_version}",
    )

    st.session_state.m2_defaults["new_qname"] = new_qname
    st.session_state.m2_defaults["new_label"] = new_label

    st.caption("Select source choices to combine into a new recoded choice. Repeat to create multiple groups.")

    st.markdown("**Source choices**")
    cols = st.columns(3)

    for i, ch in enumerate(src_choices):
        lab = str(ch.get("label", "")).strip()
        code = str(ch.get("code", "")).strip()
        if not lab or not code:
            continue

        cb_key = f"m2_cb_{source_qname}_{st.session_state.m2_ui_version}_{i}_{code}"
        with cols[i % 3]:
            checked = st.checkbox(f"{lab} ({code})", key=cb_key)

        if checked:
            st.session_state.m2_pick.add(code)
        else:
            st.session_state.m2_pick.discard(code)

    st.markdown("---")
    g1, g2, g3 = st.columns([3, 1, 1])

    with g1:
        group_text = st.text_input(
            "New recoded choice text",
            value=st.session_state.m2_defaults.get("group_text", ""),
            placeholder="18-34",
            key=f"m2_group_text_{st.session_state.m2_ui_version}",
        )
    with g2:
        group_code = st.text_input(
            "New recoded choice code",
            value=st.session_state.m2_defaults.get("group_code", ""),
            placeholder="2",
            key=f"m2_group_code_{st.session_state.m2_ui_version}",
        )
    with g3:
        if st.button("🧽 Clear selection", key=f"m2_clear_pick_{st.session_state.m2_ui_version}"):
            st.session_state.m2_pick = set()
            st.session_state.m2_ui_version += 1
            st.rerun()

    if st.button("➕ Add recode group", key=f"m2_add_group_{st.session_state.m2_ui_version}"):
        if not st.session_state.m2_pick:
            st.error("Select at least one source choice to recode.")
        elif not group_text.strip() or not group_code.strip():
            st.error("Enter both a recoded choice text and code.")
        else:
            from_list = []
            for ch in src_choices:
                c = str(ch.get("code", "")).strip()
                if c in st.session_state.m2_pick:
                    from_list.append({"label": str(ch.get("label", "")).strip(), "code": c})

            st.session_state.m2_work_groups.append({
                "new_text": group_text.strip(),
                "new_code": group_code.strip(),
                "from": from_list,
            })

            st.session_state.m2_pick = set()
            st.session_state.m2_defaults["group_text"] = ""
            st.session_state.m2_defaults["group_code"] = ""
            st.session_state.m2_ui_version += 1
            st.success("Added recode group.")
            st.rerun()

    if st.session_state.m2_work_groups:
        st.subheader("Current recode groups")
        for gi, g in enumerate(st.session_state.m2_work_groups, start=1):
            frm = ", ".join([f"{x['label']}({x['code']})" for x in g.get("from", [])])
            st.write(f"{gi}. **{g.get('new_text','')}** → {g.get('new_code','')}  _(from: {frm})_")

        if st.button("🗑️ Delete last group", key=f"m2_del_last_group_{st.session_state.m2_ui_version}"):
            st.session_state.m2_work_groups.pop()
            st.session_state.m2_ui_version += 1
            st.rerun()

    s1, s2 = st.columns([1, 1])

    with s1:
        if st.button("💾 Save as new recode", key=f"m2_save_new_{st.session_state.m2_ui_version}"):
            if not new_qname.strip():
                st.error("New variable name is required.")
            else:
                payload = {
                    "source_qname": source_qname,
                    "new_qname": new_qname.strip(),
                    "new_label": new_label.strip(),
                    "groups": [dict(g) for g in st.session_state.m2_work_groups],
                }
                st.session_state.recodes.append(payload)
                st.session_state.m2_last_sidebar_sel = len(st.session_state.recodes) - 1
                _module2_load_into_ui(payload, as_new=False)

                # --- NEW: publish to var_catalog ---
                _module2_publish_recode_to_catalog(payload)

                st.success("Saved as new recode.")
                st.rerun()

    with s2:
        if st.button("✅ Update selected recode", key=f"m2_update_{st.session_state.m2_ui_version}"):
            if st.session_state.m2_last_sidebar_sel is None:
                st.error("No recode selected to update.")
            else:
                idx = st.session_state.m2_last_sidebar_sel
                if idx < 0 or idx >= len(st.session_state.recodes):
                    st.error("Selected recode index is invalid.")
                elif not new_qname.strip():
                    st.error("New variable name is required.")
                else:
                    # if new_qname changed, remove old catalog var if it was m2
                    old_newv = str(st.session_state.recodes[idx].get("new_qname", "")).strip()

                    payload = {
                        "source_qname": source_qname,
                        "new_qname": new_qname.strip(),
                        "new_label": new_label.strip(),
                        "groups": [dict(g) for g in st.session_state.m2_work_groups],
                    }
                    st.session_state.recodes[idx] = payload

                    if old_newv and old_newv != payload["new_qname"]:
                        _module2_remove_recode_from_catalog_if_m2(old_newv)

                    _module2_publish_recode_to_catalog(payload)

                    st.success("Updated recode.")
                    st.session_state.m2_ui_version += 1
                    st.rerun()

    st.subheader("Export preview (recodes)")
    export_df = build_module2_export_df()

    if export_df.empty:
        st.info("No recodes created yet.")
    else:
        st.dataframe(export_df, use_container_width=True, hide_index=True)

        st.download_button(
            label="⬇️ Download Excel (recodes)",
            data=to_excel_bytes_from_export_df(export_df, sheet_name="Recodes"),
            file_name="recodes_spss_style.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        )

    # ----------------------------
    # SPSS syntax output (Module 2)
    # ----------------------------
    st.divider()
    st.subheader("SPSS syntax (Module 2 recodes)")

    spss_m2 = generate_spss_from_module2(st.session_state.questions, st.session_state.recodes)
    if not spss_m2.strip():
        st.info("No SPSS recode syntax yet — create at least one recode.")
    else:
        st.code(spss_m2, language="spss")
        st.download_button(
            "⬇️ Download SPSS syntax (Module 2)",
            data=spss_m2.encode("utf-8"),
            file_name="module2_recodes.sps",
            mime="text/plain",
            key="m2_download_spss",
        )














# =========================================================
# MODULE 0: PROJECT (Save / Load)
# UPDATED:
# ✅ Saves/loads Module 2.5 (module25)
# ✅ Saves/loads Module 4 weighting SESSION STATE (module4)
#    - config/progress/factors/SPSS lines
#    - DOES NOT store datasets (m4_dfw / m4_uploaded_df)
# =========================================================

import json
from datetime import datetime
import streamlit as st


# --- Project state init ---
if "project" not in st.session_state:
    st.session_state.project = {"name": "Untitled project", "last_saved": None}

if "project_slots" not in st.session_state:
    st.session_state.project_slots = {}


# Ensure Module 2.5 keys exist
if "derived_vars" not in st.session_state:
    st.session_state.derived_vars = []
if "var_catalog" not in st.session_state:
    st.session_state.var_catalog = {}


def _reset_module_25_ui_state():
    # Prevent stale radio widget state from blanking sidebar
    if "m25_radio" in st.session_state:
        del st.session_state["m25_radio"]

    st.session_state.m25_selected_index = None
    st.session_state.m25_last_sidebar_sel = None
    st.session_state.m25_load_index = None
    st.session_state.m25_reset_editor = True
    st.session_state.m25_clear_rule_inputs = False
    st.session_state.m25_rev = st.session_state.get("m25_rev", 0) + 1


def _reset_module_4_ui_state():
    """
    Resets UI-ish keys for Module 4 so it renders cleanly after load/new.
    We do NOT wipe weighting progress here; that's handled separately.
    """
    for k in ["m4_source_choice", "m4_upload", "m4_out_fmt", "m4_order_sel", "m4_order_pick", "m4_freq_pick"]:
        # Only delete if you want widgets to reset after load; generally safe
        if k in st.session_state:
            del st.session_state[k]


def _project_payload() -> dict:
    """
    Everything needed to restore the app state later.
    Keep this conservative: stable data, not raw datasets.
    """
    # ---- Module 4: serialize factors_by_var safely ----
    # Stored as dict[var] -> list of {code,label,factor}
    m4_factors = {}
    fbv = st.session_state.get("m4_factors_by_var", {}) or {}
    for var, df in fbv.items():
        try:
            # df is a pandas df in module 4; but we store plain records
            m4_factors[str(var)] = df.fillna("").to_dict("records")
        except Exception:
            # fallback: skip anything weird
            continue

    return {
        "meta": {
            "app": "Question Builder ➜ Excel (SPSS Style)",
            "version": 4,  # bumped because module4 now persisted
            "saved_at": datetime.now().isoformat(timespec="seconds"),
            "project_name": st.session_state.project.get("name", "Untitled project"),
        },
        "module1": {"questions": st.session_state.get("questions", [])},
        "module2": {"recodes": st.session_state.get("recodes", [])},
        "module25": {
            "derived_vars": st.session_state.get("derived_vars", []),
            "var_catalog": st.session_state.get("var_catalog", {}),
        },
        "module4": {
            # core session state (NO datasets)
            "m4_data_source": st.session_state.get("m4_data_source", "auto"),
            "m4_weight_order": st.session_state.get("m4_weight_order", []),
            "m4_stage_idx": st.session_state.get("m4_stage_idx", 0),
            "m4_stage_offset": st.session_state.get("m4_stage_offset", 0),

            "m4_order_list": st.session_state.get("m4_order_list", []),
            "m4_order_selected": st.session_state.get("m4_order_selected", None),

            "m4_freq_vars": st.session_state.get("m4_freq_vars", []),
            "m4_weighted_n_decimals": st.session_state.get("m4_weighted_n_decimals", 8),

            "m4_spss_lines": st.session_state.get("m4_spss_lines", []),

            # big one
            "m4_factors_by_var": m4_factors,
        },
    }


def _load_project_payload(payload: dict) -> None:
    """Load a saved project payload into session state.

    Notes
    -----
    * Defensive against older payload shapes.
    * Does NOT restore raw datasets (df_in / m4_dfw / m4_uploaded_df).
    * Resets UI/widget state so sidebars/forms render cleanly after load.
    """

    if not isinstance(payload, dict):
        raise ValueError("Project file is not valid JSON (expected an object).")

    meta = payload.get("meta", {}) if isinstance(payload.get("meta", {}), dict) else {}

    # ---- Project name + last_saved ----
    if "project" not in st.session_state:
        st.session_state.project = {"name": "Untitled project", "last_saved": None}

    st.session_state.project["name"] = (
        str(meta.get("project_name") or st.session_state.project.get("name") or "Untitled project").strip()
        or "Untitled project"
    )
    st.session_state.project["last_saved"] = meta.get("saved_at")

    # ---- Module 1 ----
    m1 = payload.get("module1", {}) if isinstance(payload.get("module1", {}), dict) else {}
    questions = m1.get("questions")
    if questions is None:
        # older saves
        questions = payload.get("questions", [])
    st.session_state.questions = questions if isinstance(questions, list) else []

    # Reset Module 1 UI state
    st.session_state.choices = [{"label": "", "code": ""}]
    st.session_state.form_version = st.session_state.get("form_version", 0) + 1
    st.session_state.selected_q_index = 0 if st.session_state.questions else None
    st.session_state.mode = "new"
    st.session_state._defaults = {"qname": "", "prompt": "", "label": ""}

    # ---- Module 2 ----
    m2 = payload.get("module2", {}) if isinstance(payload.get("module2", {}), dict) else {}
    recodes = m2.get("recodes")
    if recodes is None:
        # older saves
        recodes = payload.get("recodes", [])
    st.session_state.recodes = recodes if isinstance(recodes, list) else []

    # Ensure older recode schemas don't break newer code paths
    # (some very old saves used new_name/new_label)
    fixed = []
    for r in st.session_state.recodes:
        if not isinstance(r, dict):
            continue
        rr = dict(r)
        if "new_qname" not in rr and rr.get("new_name"):
            rr["new_qname"] = rr.get("new_name")
        if "new_label" not in rr and rr.get("label"):
            rr["new_label"] = rr.get("label")
        fixed.append(rr)
    st.session_state.recodes = fixed

    # Reset Module 2 UI state (newer keys)
    st.session_state.m2_selected_recode_index = None
    st.session_state.m2_mode = "new"
    st.session_state.m2_ui_version = st.session_state.get("m2_ui_version", 0) + 1
    st.session_state.m2_pick = set()
    st.session_state.m2_defaults = {
        "source_qname": "",
        "new_qname": "",
        "new_label": "",
        "group_text": "",
        "group_code": "",
    }
    st.session_state.m2_work_groups = []
    st.session_state.m2_last_sidebar_sel = None
    st.session_state.m2_last_source_qname = None

    # Back-compat keys (some earlier builds used these names)
    st.session_state.recode_pick = set()
    st.session_state.recode_form_version = st.session_state.get("recode_form_version", 0) + 1
    st.session_state.recode_selected_source = None
    st.session_state.recode_work_groups = []

    # ---- Module 2.5 ----
    m25 = payload.get("module25", {}) if isinstance(payload.get("module25", {}), dict) else {}
    derived_vars = m25.get("derived_vars")
    if derived_vars is None:
        derived_vars = payload.get("derived_vars", [])
    st.session_state.derived_vars = derived_vars if isinstance(derived_vars, list) else []

    # If present, take var_catalog; otherwise rebuild below
    vc = m25.get("var_catalog")
    st.session_state.var_catalog = vc if isinstance(vc, dict) else st.session_state.get("var_catalog", {})

    # Reset Module 2.5 UI state
    _reset_module_25_ui_state()

    # ---- Module 4 (weighting) ----
    m4 = payload.get("module4", {}) if isinstance(payload.get("module4", {}), dict) else {}

    st.session_state.m4_data_source = m4.get("m4_data_source", "auto")
    st.session_state.m4_weight_order = m4.get("m4_weight_order", []) or []
    st.session_state.m4_stage_idx = int(m4.get("m4_stage_idx", 0) or 0)
    st.session_state.m4_stage_offset = int(m4.get("m4_stage_offset", 0) or 0)
    st.session_state.m4_order_list = m4.get("m4_order_list", []) or []
    st.session_state.m4_order_selected = m4.get("m4_order_selected", None)
    st.session_state.m4_freq_vars = m4.get("m4_freq_vars", []) or []
    st.session_state.m4_weighted_n_decimals = int(m4.get("m4_weighted_n_decimals", 8) or 8)
    st.session_state.m4_spss_lines = m4.get("m4_spss_lines", []) or []

    # Restore factor tables (dict[var] -> list[records]) back into DataFrames
    fbv = {}
    raw_fbv = m4.get("m4_factors_by_var", {}) or {}
    if isinstance(raw_fbv, dict):
        for var, records in raw_fbv.items():
            try:
                if isinstance(records, list):
                    fbv[str(var)] = pd.DataFrame(records)
            except Exception:
                continue
    st.session_state.m4_factors_by_var = fbv

    # Never load datasets
    st.session_state.m4_uploaded_df = None
    st.session_state.m4_dfw = None

    _reset_module_4_ui_state()

    # ---- Rebuild var_catalog so downstream modules always have a consistent view ----
    # 1) Module 1 questions
    try:
        _sync_all_module1_questions_to_catalog()
    except Exception:
        # If module1 helpers aren't available for any reason, ignore.
        pass

    # 2) Module 2 recodes
    try:
        for r in st.session_state.get("recodes", []) or []:
            if isinstance(r, dict):
                _module2_publish_recode_to_catalog(r)
    except Exception:
        pass

    # 3) Keep any module25 catalog entries (e.g., derived vars) but ensure dict
    if "var_catalog" not in st.session_state or not isinstance(st.session_state.var_catalog, dict):
        st.session_state.var_catalog = {}



def _build_var_catalog_from_state():
    base = []
    df = st.session_state.get("df_in", None)
    if df is not None:
        try:
            base = list(df.columns)
        except Exception:
            base = []

    recodes = [
        d.get("new_name") for d in st.session_state.get("recodes", [])
        if isinstance(d, dict) and d.get("new_name")
    ]

    derived = [
        d.get("name") for d in st.session_state.get("derived_vars", [])
        if isinstance(d, dict) and d.get("name")
    ]

    # If Module 1 “questions” contain qname, include them too (optional but useful)
    questions = [
        q.get("qname") for q in st.session_state.get("questions", [])
        if isinstance(q, dict) and q.get("qname")
    ]

    seen = set()
    out = []
    for v in base + questions + recodes + derived:
        v = (v or "").strip()
        if v and v not in seen:
            seen.add(v)
            out.append(v)
    return out


def _refresh_var_catalog():
    st.session_state.var_catalog = _build_var_catalog_from_state()
    st.session_state.var_catalog_rev = st.session_state.get("var_catalog_rev", 0) + 1



def _bytes_for_download(payload: dict) -> bytes:
    return json.dumps(payload, indent=2).encode("utf-8")


def render_module_0():
    st.header("Module 0: Project")

    # Project name
    st.session_state.project["name"] = st.text_input(
        "Project name",
        value=st.session_state.project.get("name", "Untitled project"),
        key="m0_project_name",
    )

    st.caption("Save = download a project file (.json). Open = upload that file later to continue editing.")

    c1, c2, c3, c4 = st.columns([1.2, 1.2, 1.6, 1.4])

    # --- New project ---
    with c1:
        if st.button("🆕 New project", use_container_width=True, key="m0_new_project"):
            st.session_state.questions = []
            st.session_state.recodes = []

            st.session_state.derived_vars = []
            st.session_state.var_catalog = {}

            # Module 4: wipe weighting state
            st.session_state.m4_data_source = "auto"
            st.session_state.m4_uploaded_df = None
            st.session_state.m4_dfw = None
            st.session_state.m4_weight_order = []
            st.session_state.m4_stage_idx = 0
            st.session_state.m4_factors_by_var = {}
            st.session_state.m4_spss_lines = []
            st.session_state.m4_freq_vars = ["QAGE", "QSEX", "QRACE", "QPARTYID", "QCOUNTY", "cQEDUCATION"]
            st.session_state.m4_order_list = []
            st.session_state.m4_order_selected = None
            st.session_state.m4_stage_offset = 0
            st.session_state.m4_weighted_n_decimals = 8
            _reset_module_4_ui_state()

            # Module 1/2 UI reset
            st.session_state.choices = [{"label": "", "code": ""}]
            st.session_state.form_version = st.session_state.get("form_version", 0) + 1
            st.session_state.selected_q_index = None
            st.session_state.mode = "new"
            st.session_state._defaults = {"qname": "", "prompt": "", "label": ""}

            st.session_state.recode_pick = set()
            st.session_state.recode_form_version = st.session_state.get("recode_form_version", 0) + 1
            st.session_state.recode_selected_source = None
            st.session_state.recode_work_groups = []

            _reset_module_25_ui_state()

            st.session_state.project["last_saved"] = None
            st.success("Started a new project.")
            st.rerun()

    # --- Save project (download JSON) ---
    with c2:
        payload = _project_payload()
        default_name = st.session_state.project.get("name", "project").strip().replace(" ", "_") or "project"
        filename = f"{default_name}.json"

        st.download_button(
            "💾 Save project",
            data=_bytes_for_download(payload),
            file_name=filename,
            mime="application/json",
            use_container_width=True,
            key="m0_save_download",
        )

    # --- Open project (upload JSON) ---
    with c3:
        uploaded = st.file_uploader(
            "Open project (.json)",
            type=["json"],
            label_visibility="collapsed",
            key="m0_open_uploader",
        )
        if uploaded is not None:
            try:
                payload_in = json.loads(uploaded.getvalue().decode("utf-8"))
                _load_project_payload(payload_in)
                st.success("Project loaded.")
                st.rerun()
            except Exception as e:
                st.error(f"Could not load project: {e}")

    # --- Quick in-session slots ---
    with c4:
        st.markdown("**Quick slots (this session only)**")
        slot_name = st.text_input("Slot name", value="Slot 1", key="m0_slot_name")
        s1, s2 = st.columns(2)

        if s1.button("Save slot", use_container_width=True, key="m0_save_slot"):
            st.session_state.project_slots[slot_name] = _project_payload()
            st.success(f"Saved to slot: {slot_name}")

        if s2.button("Load slot", use_container_width=True, key="m0_load_slot"):
            if slot_name not in st.session_state.project_slots:
                st.warning("No saved data in that slot yet.")
            else:
                try:
                    _load_project_payload(st.session_state.project_slots[slot_name])
                    st.success(f"Loaded slot: {slot_name}")
                    st.rerun()
                except Exception as e:
                    st.error(f"Could not load slot: {e}")

    st.divider()
    st.subheader("Project summary")
    st.write(
        {
            "Project name": st.session_state.project.get("name"),
            "Last saved (from file)": st.session_state.project.get("last_saved"),
            "Module 1 variables": len(st.session_state.get("questions", [])),
            "Module 2 recodes": len(st.session_state.get("recodes", [])),
            "Module 2.5 derived vars": len(st.session_state.get("derived_vars", [])),
            "Module 4 order vars": len(st.session_state.get("m4_weight_order", [])),
            "Module 4 factor tables": len(st.session_state.get("m4_factors_by_var", {}) or {}),
            "Module 4 SPSS lines": len(st.session_state.get("m4_spss_lines", []) or []),
        }
    )





# ============================
# Module 3: Import + Match Data
# ============================

import difflib
import numpy as np

def _normalize_name(x: str) -> str:
    return "".join(ch for ch in str(x).strip().lower() if ch.isalnum() or ch == "_")

def _similarity(a: str, b: str) -> float:
    return difflib.SequenceMatcher(None, _normalize_name(a), _normalize_name(b)).ratio()

def _best_match(col: str, targets: list[str], threshold: float = 0.80):
    best = None
    best_score = 0.0
    for t in targets:
        s = _similarity(col, t)
        if s > best_score:
            best_score = s
            best = t
    if best is not None and best_score >= threshold:
        return best, best_score
    return None, best_score


# -------------------------
# Module 7: column name resolving (fast)
# -------------------------

def _m7_resolve_name(name: str, resolve_map: dict | None) -> str:
    '''Return the actual dataframe column for a displayed name.'''
    if not resolve_map:
        return name
    return resolve_map.get(name, name)


@st.cache_data(show_spinner=False)
def _m7_build_resolve_map_cached(
    base_cols: tuple,
    canonical_vars: tuple,
    threshold: float = 0.80,
) -> tuple[dict, set]:
    '''Build (resolve_map, matched_raw).

    Cached so checkbox clicking does not redo fuzzy matching work.
    base_cols: dataset columns (after dropping __text)
    canonical_vars: var_catalog keys
    '''
    base_cols_list = [str(c) for c in base_cols]
    canonical_list = [str(v) for v in canonical_vars]

    resolve_map: dict[str, str] = {}
    matched_raw: set[str] = set()

    lower_to_col = {str(c).strip().lower(): str(c) for c in base_cols_list}

    for v in sorted(canonical_list):
        key = str(v).strip().lower()
        match = lower_to_col.get(key)
        if match is None:
            match, _score = _best_match(str(v), base_cols_list, threshold=threshold)
        if match is not None:
            resolve_map[str(v)] = match
            matched_raw.add(match)

    return resolve_map, matched_raw

def _get_series_safe(df: pd.DataFrame, colname: str) -> pd.Series:
    """
    Always return a Series even if df has duplicate colnames.
    If duplicates exist, take the FIRST one.
    """
    if colname not in df.columns:
        return pd.Series([np.nan] * len(df), index=df.index)

    # df.loc[:, colname] returns Series if unique, DataFrame if duplicates
    obj = df.loc[:, colname]
    if isinstance(obj, pd.DataFrame):
        return obj.iloc[:, 0]
    return obj

def _build_choice_code_to_text(qdef: dict) -> dict:
    """
    From Module 1 question definition, build mapping: code(str)->label(str)
    """
    m = {}
    for ch in qdef.get("choices", []):
        code = str(ch.get("code", "")).strip()
        lab = str(ch.get("label", "")).strip()
        if code:
            m[code] = lab
    return m

def _apply_label_rollups(df_out: pd.DataFrame, questions: list[dict]):
    """
    For each Module 1 variable QXXX, create a label column like:
      QAGE__text  (or whatever you want)
    So 3 -> "35-44", etc.
    """
    for q in questions:
        qname = q.get("qname", "").strip()
        if not qname:
            continue
        if qname not in df_out.columns:
            continue

        code_to_text = _build_choice_code_to_text(q)
        s = _get_series_safe(df_out, qname).astype(str).str.strip()

        df_out[f"{qname}__text"] = s.map(lambda v: code_to_text.get(v, ""))

def _apply_recode_definitions(df_out: pd.DataFrame, questions: list[dict], recodes: list[dict]):
    """
    Creates recode columns (new_qname) from source_qname.

    Logic:
    - if value code belongs to any group.from codes -> output group.new_code
    - else keep original code (pass-through)
    """
    # quick lookup for source question defs
    q_lookup = {q.get("qname"): q for q in questions}

    for rec in recodes:
        source_q = str(rec.get("source_qname", "")).strip()
        new_q = str(rec.get("new_qname", "")).strip()
        if not source_q or not new_q:
            continue
        if source_q not in df_out.columns:
            continue

        groups = rec.get("groups", [])

        # Build map source_code(str) -> new_code(str)
        src_to_new = {}
        for g in groups:
            new_code = str(g.get("new_code", "")).strip()
            for f in g.get("from", []):
                src_code = str(f.get("code", "")).strip()
                if src_code and new_code:
                    src_to_new[src_code] = new_code

        s = _get_series_safe(df_out, source_q).astype(str).str.strip()

        # pass-through if not recoded
        df_out[new_q] = s.map(lambda v: src_to_new.get(v, v))

        # Also create text for recoded variable if possible:
        # Prefer group new_text for recoded codes; otherwise fall back to original label text
        newcode_to_text = {}
        for g in groups:
            nc = str(g.get("new_code", "")).strip()
            nt = str(g.get("new_text", "")).strip()
            if nc and nt:
                newcode_to_text[nc] = nt

        # fallback from source label mapping (if pass-through codes)
        src_qdef = q_lookup.get(source_q, {})
        src_code_to_text = _build_choice_code_to_text(src_qdef)

        df_out[f"{new_q}__text"] = df_out[new_q].astype(str).str.strip().map(
            lambda v: newcode_to_text.get(v, src_code_to_text.get(v, ""))
        )


def _apply_derived_vars_rules(df: pd.DataFrame, derived_specs: list, var_catalog: dict | None = None) -> pd.DataFrame:
    """Apply Module 2.5 derived variables (RULES specs) onto a dataframe.

    - Creates the derived code column (spec['name'])
    - Also creates a text column '<name>__text' when catalog choices exist.
    """
    if df is None or not isinstance(df, pd.DataFrame):
        return df
    if not derived_specs:
        return df

    cat = var_catalog or (st.session_state.get("var_catalog", {}) or {})

    def _col_lookup(colname: str):
        if not colname:
            return None
        # case/space-insensitive lookup
        wanted = str(colname).strip().lower()
        for c in df.columns:
            if str(c).strip().lower() == wanted:
                return c
        return None

    def _as_str_series(s: pd.Series) -> pd.Series:
        return s.fillna("").astype(str).str.strip()

    for spec in derived_specs:
        if not isinstance(spec, dict):
            continue
        if str(spec.get("type", "")).strip().upper() != "RULES":
            continue

        name = str(spec.get("name", "")).strip()
        if not name:
            continue

        rules = spec.get("rules") or []
        default_code = str(spec.get("default_code", "")).strip()

        out = pd.Series([default_code] * len(df), index=df.index, dtype="object")

        # Apply rules in order (last match wins for overlapping masks)
        for rule in rules:
            if not isinstance(rule, dict):
                continue
            set_to = str(rule.get("set_to", "")).strip()
            when = rule.get("when") or []
            if not set_to or not when:
                continue

            mask = pd.Series(True, index=df.index)
            for cond in when:
                if not isinstance(cond, dict):
                    continue
                v = str(cond.get("var", "")).strip()
                vals = cond.get("values") or []
                if not v or not vals:
                    mask &= False
                    continue

                col = _col_lookup(v)
                if col is None:
                    mask &= False
                    continue

                vals_set = {str(x).strip() for x in vals if str(x).strip() != ""}
                s = _as_str_series(df[col])
                mask &= s.isin(vals_set)

            out.loc[mask] = set_to

        df[name] = out

        # Create text column if we have choices in catalog
        choices = {}
        cat_entry = cat.get(name) if isinstance(cat, dict) else None
        if isinstance(cat_entry, dict):
            choices = cat_entry.get("choices") or {}
        if isinstance(choices, list):
            # list of dicts {code,label}
            tmp = {}
            for r in choices:
                if isinstance(r, dict):
                    c = str(r.get("code", "")).strip()
                    t = str(r.get("text", "") or r.get("label", "")).strip()
                    if c:
                        tmp[c] = t or c
            choices = tmp

        if isinstance(choices, dict) and len(choices) > 0:
            mapper = {str(k).strip(): str(v).strip() for k, v in choices.items()}
            df[f"{name}__text"] = _as_str_series(df[name]).map(lambda x: mapper.get(x, x))

    return df

def render_module_3():
    st.header("Module 3: Import + Match Data")

    if "module3_mapping" not in st.session_state:
        st.session_state.module3_mapping = {}  # canonical_qname -> input_colname OR None

    uploaded = st.file_uploader("Upload a dataset (.csv or Excel)", type=["csv", "xlsx", "xls"])
    if not uploaded:
        st.info("Upload a file to begin.")
        return

    # Read file
    if uploaded.name.lower().endswith(".csv"):
        df_in = pd.read_csv(uploaded)
    else:
        df_in = pd.read_excel(uploaded)

    st.write(f"Loaded **{df_in.shape[0]:,}** rows × **{df_in.shape[1]:,}** columns.")

    # Canonical vars from Module 1
    if len(st.session_state.questions) == 0:
        st.warning("No variables from Module 1. Build variables in Module 1 first.")
        return

    canonical_vars = [q["qname"] for q in st.session_state.questions if q.get("qname")]
    input_cols = list(df_in.columns)

    st.subheader("Auto-match columns (fuzzy 80%)")

    threshold = st.slider("Match threshold", 0.50, 0.95, 0.80, 0.01)

    # Auto-match: for each input col, find best canonical target
    # We will store mapping as canonical -> input col (one input per canonical)
    # If multiple inputs match same canonical, take the best score.
    auto_map = {}
    auto_scores = {}

    for col in input_cols:
        best, score = _best_match(col, canonical_vars, threshold=threshold)
        if best:
            if best not in auto_map or score > auto_scores.get(best, 0):
                auto_map[best] = col
                auto_scores[best] = score

    # Initialize session mapping if empty (first run)
    if not st.session_state.module3_mapping:
        for qn in canonical_vars:
            st.session_state.module3_mapping[qn] = auto_map.get(qn, None)

    # ---- Manual override UI
    st.subheader("Review / edit matches")

    # Add an explicit UNMATCH option
    picker_options = ["(unmatched)"] + input_cols

    # Show as a table-like UI
    for qn in canonical_vars:
        current = st.session_state.module3_mapping.get(qn, None)
        current_label = current if current in input_cols else "(unmatched)"

        cols = st.columns([2, 3, 1])
        with cols[0]:
            st.markdown(f"**{qn}**")
        with cols[1]:
            selected = st.selectbox(
                "Match",
                options=picker_options,
                index=picker_options.index(current_label),
                key=f"m3_match_{qn}",
                label_visibility="collapsed",
            )
            st.session_state.module3_mapping[qn] = None if selected == "(unmatched)" else selected
        with cols[2]:
            # quick unmatch button
            if st.button("Unmatch", key=f"m3_unmatch_{qn}"):
                st.session_state.module3_mapping[qn] = None
                st.rerun()

    # ---- Warning: columns starting with Q that didn't match to anything
    matched_input_cols = {v for v in st.session_state.module3_mapping.values() if v}
    q_like_unmatched = [c for c in input_cols if str(c).strip().upper().startswith("Q") and c not in matched_input_cols]

    if q_like_unmatched:
        st.warning(
            "These input columns start with **Q** but are currently **not matched** to any scripted variable:\n\n"
            + ", ".join(q_like_unmatched)
        )

    # ---- Build output df: keep all original columns + add canonical copies
    st.subheader("Build output (original columns + canonical + recodes)")

    df_out = df_in.copy()

    # Add canonical columns (do NOT delete originals)
    for qn, in_col in st.session_state.module3_mapping.items():
        if in_col is None:
            continue
        if in_col not in df_out.columns:
            continue

        # Create/overwrite canonical column with values from input column
        df_out[qn] = df_out[in_col]

    # Apply label rollups for canonical columns
    _apply_label_rollups(df_out, st.session_state.questions)

    # Apply recodes (creates cQ vars + cQ__text)
    _apply_recode_definitions(df_out, st.session_state.questions, st.session_state.get("recodes", []))

    # Apply Module 2.5 derived variables onto the mapped dataset (adds columns)
    derived_25 = (
        st.session_state.get("derived_vars_25")
        or st.session_state.get("derived_vars")
        or st.session_state.get("derived_variables_25")
        or []
    )
    df_out = _apply_derived_vars_rules(df_out, derived_25, st.session_state.get("var_catalog", {}))

    # Re-apply recodes AFTER derived vars so you can recode Module 2.5 outputs
    # (e.g., recoding a derived var into c<derived> / cc<derived>).
    _apply_recode_definitions(df_out, st.session_state.questions, st.session_state.get("recodes", []))



    # ---- make Module 3 output available to Module 4
    st.session_state.df_out = df_out
    st.session_state.df_out_source = "module3"


    # Preview a small slice
    st.subheader("Preview (first 25 rows)")
    st.dataframe(df_out.head(25), use_container_width=True)

    # Download
    st.subheader("Download output")
    fmt = st.radio("Output format", ["Excel (.xlsx)", "CSV (.csv)"], horizontal=True)

    if fmt.startswith("Excel"):
        # Write Excel
        output = BytesIO()
        with pd.ExcelWriter(output, engine="openpyxl") as writer:
            df_out.to_excel(writer, index=False, sheet_name="Data")
        st.download_button(
            label="⬇️ Download mapped + recoded dataset (Excel)",
            data=output.getvalue(),
            file_name="mapped_recoded_dataset.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        )
    else:
        csv_bytes = df_out.to_csv(index=False).encode("utf-8")
        st.download_button(
            label="⬇️ Download mapped + recoded dataset (CSV)",
            data=csv_bytes,
            file_name="mapped_recoded_dataset.csv",
            mime="text/csv",
        )




# =========================================================
# MODULE 2.5: Derived Variables (NO dataset required)
# Rule builder using variable definitions from Modules 1 & 2
#
# What this module does:
# - Builds/saves derived-variable SPECS only (no dataframe operations here)
# - Uses st.session_state.questions (Module 1) and st.session_state.recodes (Module 2)
#   to know which variables exist + what codes/labels they have
# - Stores derived specs in st.session_state.derived_vars
# - Publishes derived variables into a central catalog:
#     st.session_state.var_catalog[var] = {"label":..., "choices":{code:text}, "origin":...}
# - Prints SPSS syntax for all derived vars at bottom
#
# Includes:
# - ⭐ Add standard derived variables button:
#     csexage, bcsexage, csexeducation, bcsexeducation,
#     csexpartyid, bcsexpartyid, cpartyeducation
# - Stable “New” behavior: clicking New clears EVERYTHING safely (no StreamlitAPIException)
# - Safe “Add rule” input clearing (no StreamlitAPIException)
#
# NOTE:
# - Codes are stored/compared as STRINGS in specs. (Module 3 should compare as strings.)
# - If a variable has no known choices, you can type codes manually (comma-separated).
# =========================================================

import streamlit as st
import pandas as pd
import numpy as np

# ----------------------------
# Session init
# ----------------------------
if "derived_vars" not in st.session_state:
    st.session_state.derived_vars = []  # list of specs

# Central variable catalog (optional but recommended)
# {var: {"label": str, "choices": {code: text}, "origin": "m1|m2|m25", "type": "single"}}
if "var_catalog" not in st.session_state:
    st.session_state.var_catalog = {}

# Revision token for forcing widget-key refresh when needed
if "m25_rev" not in st.session_state:
    st.session_state.m25_rev = 0

if "m25_selected_index" not in st.session_state:
    st.session_state.m25_selected_index = None

if "m25_last_sidebar_sel" not in st.session_state:
    st.session_state.m25_last_sidebar_sel = None

# Editor fields (we do NOT set these after widgets instantiate; we use pre-run flags)
if "m25_name" not in st.session_state:
    st.session_state.m25_name = ""
if "m25_label" not in st.session_state:
    st.session_state.m25_label = ""
if "m25_default_code" not in st.session_state:
    st.session_state.m25_default_code = ""

if "m25_work_conditions" not in st.session_state:
    st.session_state.m25_work_conditions = []  # list of {"var": "", "picked": set(), "manual_codes":""}
if "m25_work_rules" not in st.session_state:
    st.session_state.m25_work_rules = []  # list of {"when":[...], "set_to":"", "set_label":""}
if "m25_work_vlabels" not in st.session_state:
    st.session_state.m25_work_vlabels = []  # list of {"code":"", "text":""}

# Add-rule inputs (must be keys in session_state to avoid StreamlitAPIException)
if "m25_set_to" not in st.session_state:
    st.session_state.m25_set_to = ""
if "m25_set_label" not in st.session_state:
    st.session_state.m25_set_label = ""

# Pre-run flags (safe clearing / loading)
if "m25_clear_rule_inputs" not in st.session_state:
    st.session_state.m25_clear_rule_inputs = False
if "m25_reset_editor" not in st.session_state:
    st.session_state.m25_reset_editor = False
if "m25_load_index" not in st.session_state:
    st.session_state.m25_load_index = None  # int or None


# ----------------------------
# Helpers
# ----------------------------
def _safe_str(x) -> str:
    if x is None:
        return ""
    if isinstance(x, float) and np.isnan(x):
        return ""
    return str(x).strip()

def _safe_int(x):
    try:
        return int(str(x).strip())
    except Exception:
        return None

def _bump_rev():
    st.session_state.m25_rev += 1

def _key_safe(s: str) -> str:
    s = _safe_str(s)
    if not s:
        return "blank"
    return "".join(ch if ch.isalnum() or ch in ("_", "-") else "_" for ch in s)[:120]


# ----------------------------
# Variable registry from Modules 1 & 2
# ----------------------------
def _choices_lookup_from_module1(qname: str) -> dict:
    """
    Returns code->label from st.session_state.questions
    """
    qname = _safe_str(qname)
    if not qname:
        return {}
    q = next((q for q in st.session_state.get("questions", []) if _safe_str(q.get("qname")) == qname), None)
    if not q:
        return {}
    out = {}
    for ch in (q.get("choices") or []):
        c = _safe_str(ch.get("code"))
        t = _safe_str(ch.get("label"))
        if c:
            out[c] = t
    return out

def _label_lookup_from_module1(qname: str) -> str:
    qname = _safe_str(qname)
    q = next((q for q in st.session_state.get("questions", []) if _safe_str(q.get("qname")) == qname), None)
    return _safe_str(q.get("label")) if q else ""

def _find_recode_spec_by_new_qname(new_qname: str):
    for r in (st.session_state.get("recodes", []) or []):
        if _safe_str(r.get("new_qname")) == _safe_str(new_qname):
            return r
    return None

def _choices_lookup_from_module2_recode(new_qname: str) -> dict:
    """
    For Module 2 recode vars: new_code -> new_text
    Also passes through source choices for any codes not explicitly overwritten.
    """
    rec = _find_recode_spec_by_new_qname(new_qname)
    if not rec:
        return {}

    out = {}
    for g in (rec.get("groups") or []):
        nc = _safe_str(g.get("new_code"))
        nt = _safe_str(g.get("new_text"))
        if nc and nt:
            out[nc] = nt

    src_q = _safe_str(rec.get("source_qname"))
    src_map = _choices_lookup_from_module1(src_q) if src_q else {}
    for c, t in src_map.items():
        if c not in out:
            out[c] = t
    return out

def _label_lookup_from_module2_recode(new_qname: str) -> str:
    rec = _find_recode_spec_by_new_qname(new_qname)
    if not rec:
        return ""
    return _safe_str(rec.get("label")) or _safe_str(new_qname)

def _choices_lookup_for_any_var(qname: str) -> dict:
    """
    code->label mapping for:
    - var_catalog (preferred)
    - Module 2 recode vars (new_qname)
    - Module 1 vars
    """
    qname = _safe_str(qname)
    if not qname:
        return {}

    cat = st.session_state.get("var_catalog", {}) or {}
    if qname in cat and isinstance(cat[qname], dict):
        ch = cat[qname].get("choices")
        if isinstance(ch, dict) and ch:
            return {str(k).strip(): str(v) for k, v in ch.items()}

    rec_map = _choices_lookup_from_module2_recode(qname)
    if rec_map:
        return rec_map

    m1_map = _choices_lookup_from_module1(qname)
    if m1_map:
        return m1_map

    return {}

def _label_lookup_for_any_var(qname: str) -> str:
    qname = _safe_str(qname)
    if not qname:
        return ""

    cat = st.session_state.get("var_catalog", {}) or {}
    if qname in cat and isinstance(cat[qname], dict):
        lab = _safe_str(cat[qname].get("label"))
        if lab:
            return lab

    if _find_recode_spec_by_new_qname(qname):
        return _label_lookup_from_module2_recode(qname)

    return _label_lookup_from_module1(qname)

def _list_all_known_vars() -> list:
    """
    Union of:
    - Module 1 questions qname
    - Module 2 recode new_qname
    - var_catalog keys
    - derived_vars names
    """
    out = set()

    for q in (st.session_state.get("questions") or []):
        out.add(_safe_str(q.get("qname")))
    for r in (st.session_state.get("recodes") or []):
        out.add(_safe_str(r.get("new_qname")))
        out.add(_safe_str(r.get("source_qname")))
    for k in (st.session_state.get("var_catalog") or {}).keys():
        out.add(_safe_str(k))
    for d in (st.session_state.get("derived_vars") or []):
        out.add(_safe_str(d.get("name")))

    out = [x for x in out if x]
    out.sort(key=lambda s: s.lower())
    return out

def _publish_var_to_catalog(var: str, label: str, value_labels: list, origin: str):
    """
    Add/overwrite var_catalog entry for a variable.
    value_labels: list of {"code":..,"text":..}
    """
    var = _safe_str(var)
    if not var:
        return

    choices = {}
    for row in (value_labels or []):
        c = _safe_str(row.get("code"))
        t = _safe_str(row.get("text"))
        if c and t:
            choices[c] = t

    st.session_state.var_catalog[var] = {
        "label": _safe_str(label),
        "choices": choices,
        "origin": origin,
        "type": "single",
    }


# ----------------------------
# Standard derived vars
# ----------------------------
def _make_rules_spec(name: str, label: str, default_code: str, rules: list, value_labels: list):
    return {
        "type": "RULES",
        "name": _safe_str(name),
        "label": _safe_str(label),
        "default_code": _safe_str(default_code),
        "conditions": [],  # not used for these prebuilt specs
        "rules": rules,
        "value_labels": value_labels,
    }

def _add_or_replace_derived_spec(spec: dict):
    nm = _safe_str(spec.get("name"))
    if not nm:
        return
    idx = next((i for i, d in enumerate(st.session_state.derived_vars) if _safe_str(d.get("name")) == nm), None)
    if idx is None:
        st.session_state.derived_vars.append(spec)
    else:
        st.session_state.derived_vars[idx] = spec

    _publish_var_to_catalog(
        var=nm,
        label=_safe_str(spec.get("label")),
        value_labels=(spec.get("value_labels") or []),
        origin="m25",
    )

def _add_standard_derived_variables():
    # csexage
    csexage_rules = [
        {"when": [{"var": "qsex", "values": ["1"]}, {"var": "cqage", "values": ["4"]}], "set_to": "1", "set_label": "Men 18-44"},
        {"when": [{"var": "qsex", "values": ["1"]}, {"var": "cqage", "values": ["6"]}], "set_to": "2", "set_label": "Men 45-64"},
        {"when": [{"var": "qsex", "values": ["1"]}, {"var": "cqage", "values": ["7"]}], "set_to": "3", "set_label": "Men 65+"},
        {"when": [{"var": "qsex", "values": ["2"]}, {"var": "cqage", "values": ["4"]}], "set_to": "4", "set_label": "Women 18-44"},
        {"when": [{"var": "qsex", "values": ["2"]}, {"var": "cqage", "values": ["6"]}], "set_to": "5", "set_label": "Women 45-64"},
        {"when": [{"var": "qsex", "values": ["2"]}, {"var": "cqage", "values": ["7"]}], "set_to": "6", "set_label": "Women 65+"},
    ]
    csexage_vl = [
        {"code": "1", "text": "Men 18-44"},
        {"code": "2", "text": "Men 45-64"},
        {"code": "3", "text": "Men 65+"},
        {"code": "4", "text": "Women 18-44"},
        {"code": "5", "text": "Women 45-64"},
        {"code": "6", "text": "Women 65+"},
        {"code": "99", "text": "Other"},
    ]
    _add_or_replace_derived_spec(_make_rules_spec("csexage", "Age and Sex", "99", csexage_rules, csexage_vl))

    # bcsexage default blank (SPSS missing)
    bcsexage_vl = [v for v in csexage_vl if _safe_str(v.get("code")) != "99"]
    _add_or_replace_derived_spec(_make_rules_spec("bcsexage", "Age and Sex", "", [dict(r) for r in csexage_rules], bcsexage_vl))

    # csexeducation
    csexeducation_rules = [
        {"when": [{"var": "qsex", "values": ["1"]}, {"var": "cqeducation", "values": ["3"]}], "set_to": "1", "set_label": "Men non college"},
        {"when": [{"var": "qsex", "values": ["1"]}, {"var": "cqeducation", "values": ["5"]}], "set_to": "2", "set_label": "Men college grad"},
        {"when": [{"var": "qsex", "values": ["2"]}, {"var": "cqeducation", "values": ["3"]}], "set_to": "3", "set_label": "Women non college"},
        {"when": [{"var": "qsex", "values": ["2"]}, {"var": "cqeducation", "values": ["5"]}], "set_to": "4", "set_label": "Women college grad"},
    ]
    csexeducation_vl = [
        {"code": "1", "text": "Men non college"},
        {"code": "2", "text": "Men college grad"},
        {"code": "3", "text": "Women non college"},
        {"code": "4", "text": "Women college grad"},
        {"code": "99", "text": "Other"},
    ]
    _add_or_replace_derived_spec(_make_rules_spec("csexeducation", "Sex and Education", "99", csexeducation_rules, csexeducation_vl))

    # bcsexeducation default blank
    bcsexeducation_vl = [v for v in csexeducation_vl if _safe_str(v.get("code")) != "99"]
    _add_or_replace_derived_spec(_make_rules_spec("bcsexeducation", "Sex and Education", "", [dict(r) for r in csexeducation_rules], bcsexeducation_vl))

    # csexpartyid
    csexpartyid_rules = [
        {"when": [{"var": "qsex", "values": ["1"]}, {"var": "qpartyid", "values": ["1"]}], "set_to": "1", "set_label": "GOP Men"},
        {"when": [{"var": "qsex", "values": ["1"]}, {"var": "qpartyid", "values": ["2"]}], "set_to": "2", "set_label": "Dem Men"},
        {"when": [{"var": "qsex", "values": ["1"]}, {"var": "qpartyid", "values": ["3"]}], "set_to": "3", "set_label": "Indp Men"},
        {"when": [{"var": "qsex", "values": ["2"]}, {"var": "qpartyid", "values": ["1"]}], "set_to": "4", "set_label": "GOP Women"},
        {"when": [{"var": "qsex", "values": ["2"]}, {"var": "qpartyid", "values": ["2"]}], "set_to": "5", "set_label": "Dem Women"},
        {"when": [{"var": "qsex", "values": ["2"]}, {"var": "qpartyid", "values": ["3"]}], "set_to": "6", "set_label": "Indp Women"},
    ]
    csexpartyid_vl = [
        {"code": "1", "text": "GOP Men"},
        {"code": "2", "text": "Dem Men"},
        {"code": "3", "text": "Indp Men"},
        {"code": "4", "text": "GOP Women"},
        {"code": "5", "text": "Dem Women"},
        {"code": "6", "text": "Indp Women"},
        {"code": "99", "text": "Other"},
    ]
    _add_or_replace_derived_spec(_make_rules_spec("csexpartyid", "Party and Sex", "99", csexpartyid_rules, csexpartyid_vl))

    # bcsexpartyid default blank
    bcsexpartyid_vl = [v for v in csexpartyid_vl if _safe_str(v.get("code")) != "99"]
    _add_or_replace_derived_spec(_make_rules_spec("bcsexpartyid", "Party and Sex", "", [dict(r) for r in csexpartyid_rules], bcsexpartyid_vl))

    # cpartyeducation
    cpartyeducation_rules = [
        {"when": [{"var": "qpartyid", "values": ["1"]}, {"var": "cqeducation", "values": ["3"]}], "set_to": "1", "set_label": "GOP non college"},
        {"when": [{"var": "qpartyid", "values": ["1"]}, {"var": "cqeducation", "values": ["5"]}], "set_to": "2", "set_label": "GOP college grad"},
        {"when": [{"var": "qpartyid", "values": ["1"]}, {"var": "cqeducation", "values": ["6"]}], "set_to": "3", "set_label": "GOP Unsure"},
        {"when": [{"var": "qpartyid", "values": ["2"]}, {"var": "cqeducation", "values": ["3"]}], "set_to": "4", "set_label": "DEM non college"},
        {"when": [{"var": "qpartyid", "values": ["2"]}, {"var": "cqeducation", "values": ["5"]}], "set_to": "5", "set_label": "DEM college grad"},
        {"when": [{"var": "qpartyid", "values": ["2"]}, {"var": "cqeducation", "values": ["6"]}], "set_to": "6", "set_label": "DEM Unsure"},
        {"when": [{"var": "qpartyid", "values": ["3"]}, {"var": "cqeducation", "values": ["3"]}], "set_to": "7", "set_label": "IND non college"},
        {"when": [{"var": "qpartyid", "values": ["3"]}, {"var": "cqeducation", "values": ["5"]}], "set_to": "8", "set_label": "IND college grad"},
        {"when": [{"var": "qpartyid", "values": ["3"]}, {"var": "cqeducation", "values": ["6"]}], "set_to": "9", "set_label": "IND Unsure"},
    ]
    cpartyeducation_vl = [
        {"code": "1", "text": "GOP non college"},
        {"code": "2", "text": "GOP college grad"},
        {"code": "3", "text": "GOP Unsure"},
        {"code": "4", "text": "DEM non college"},
        {"code": "5", "text": "DEM college grad"},
        {"code": "6", "text": "DEM Unsure"},
        {"code": "7", "text": "IND non college"},
        {"code": "8", "text": "IND college grad"},
        {"code": "9", "text": "IND Unsure"},
        {"code": "99", "text": "Other"},
    ]
    _add_or_replace_derived_spec(_make_rules_spec("cpartyeducation", "Party and Education", "99", cpartyeducation_rules, cpartyeducation_vl))


# ----------------------------
# Editor lifecycle helpers (SAFE clearing/loading)
# ----------------------------
def _reset_editor_state(default_var: str = ""):
    # Clear the editor completely so "+ New derived variable" starts from a blank slate.
    st.session_state.m25_name = ""
    st.session_state.m25_label = ""
    st.session_state.m25_default_code = ""
    st.session_state.m25_work_conditions = (
        [{"var": default_var, "picked": set(), "manual_codes": ""}]
        if default_var
        else [{"var": "", "picked": set(), "manual_codes": ""}]
    )
    st.session_state.m25_work_rules = []
    st.session_state.m25_work_vlabels = []
    st.session_state.m25_set_to = ""
    st.session_state.m25_set_label = ""


def _load_spec_into_editor(spec: dict, default_var: str = ""):
    st.session_state.m25_name = _safe_str(spec.get("name", ""))
    st.session_state.m25_label = _safe_str(spec.get("label", ""))
    st.session_state.m25_default_code = _safe_str(spec.get("default_code", ""))

    conds = []
    for c in (spec.get("conditions") or []):
        conds.append({"var": _safe_str(c.get("var")), "picked": set(), "manual_codes": ""})
    if not conds:
        conds = [{"var": default_var, "picked": set(), "manual_codes": ""}] if default_var else [{"var": "", "picked": set(), "manual_codes": ""}]
    st.session_state.m25_work_conditions = conds

    st.session_state.m25_work_rules = [dict(r) for r in (spec.get("rules") or [])]
    st.session_state.m25_work_vlabels = [dict(v) for v in (spec.get("value_labels") or [])]

    st.session_state.m25_set_to = ""
    st.session_state.m25_set_label = ""


# ----------------------------
# SPSS Syntax (for derived vars)
# ----------------------------
def _spss_quote(s: str) -> str:
    return '"' + str(s).replace('"', '""') + '"'

def _build_spss_syntax_for_all_derived_25() -> str:
    """
    Generates SPSS syntax for all saved RULES derived variables in st.session_state.derived_vars.
    Mirrors:
    - AND across condition variables in a rule
    - OR within each condition's list of values
    """
    lines = []
    derived = [d for d in (st.session_state.get("derived_vars", []) or []) if d.get("type") == "RULES"]
    if not derived:
        return ""

    for spec in derived:
        name = _safe_str(spec.get("name"))
        if not name:
            continue

        label = _safe_str(spec.get("label"))
        default_code = _safe_str(spec.get("default_code", ""))

        if default_code == "":
            lines.append(f"COMPUTE {name} = $SYSMIS.")
        else:
            lines.append(f"COMPUTE {name} = {default_code}.")

        for rule in (spec.get("rules") or []):
            set_to = _safe_str(rule.get("set_to", ""))
            when = rule.get("when", []) or []
            if not set_to or not when:
                continue

            parts = []
            for cond in when:
                v = _safe_str(cond.get("var"))
                vals = [_safe_str(x) for x in (cond.get("values") or []) if _safe_str(x)]
                if not v or not vals:
                    continue
                or_terms = [f"{v} = {val}" for val in vals]
                parts.append("(" + " OR ".join(or_terms) + ")")

            if not parts:
                continue

            expr = " AND ".join(parts)
            lines.append(f"IF {expr} {name} = {set_to}.")

        lines.append("EXECUTE.")
        lines.append("")

        if label:
            lines.append(f"VARIABLE LABELS {name} {_spss_quote(label)}.")
        else:
            lines.append(f"* VARIABLE LABELS {name} {_spss_quote(name)}.")

        vlabels = [(_safe_str(v.get("code")), _safe_str(v.get("text"))) for v in (spec.get("value_labels") or [])]
        vlabels = [(c, t) for c, t in vlabels if c and t]
        if vlabels:
            lines.append(f"VALUE LABELS {name}")
            for c, t in vlabels:
                lines.append(f"  {c} {_spss_quote(t)}")
            lines.append(".")

        lines.append("")
        lines.append("* ----------------------------")
        lines.append("")

    return "\n".join(lines).strip()


# ----------------------------
# Render Module 2.5
# ----------------------------
def render_module_25():
    st.header("Module 2.5: Derived Variables (no dataset)")

    # -------- pre-run safe actions (must happen BEFORE widgets instantiate) --------
    if st.session_state.get("m25_clear_rule_inputs", False):
        st.session_state.m25_set_to = ""
        st.session_state.m25_set_label = ""
        st.session_state.m25_clear_rule_inputs = False

    if st.session_state.get("m25_reset_editor", False):
        all_known = _list_all_known_vars()
        default_var = all_known[0] if all_known else ""
        _reset_editor_state(default_var=default_var)
        st.session_state.m25_reset_editor = False

    if st.session_state.get("m25_load_index", None) is not None:
        idx = st.session_state.m25_load_index
        all_known = _list_all_known_vars()
        default_var = all_known[0] if all_known else ""
        if 0 <= idx < len(st.session_state.derived_vars):
            _load_spec_into_editor(st.session_state.derived_vars[idx], default_var=default_var)
        st.session_state.m25_load_index = None

    # -------- known vars --------
    all_vars = _list_all_known_vars()
    if not all_vars:
        st.warning("No variables found yet. Create variables in Module 1 (and recodes in Module 2) first.")
        st.info("This module builds derived-variable specs using known variable codes/labels.")
        return

    st.caption(f"Known variables available: **{len(all_vars)}** (from Modules 1, 2, 2.5)")

    # ----------------------------
    # Sidebar controls
    # ----------------------------
    st.sidebar.header("Derived variables (2.5)")

    if st.sidebar.button("⭐ Add standard derived variables", use_container_width=True, key="m25_add_standard"):
        _add_standard_derived_variables()
        st.sidebar.success("Added/updated standard derived variables.")
        _bump_rev()
        st.rerun()

    st.sidebar.divider()

    if len(st.session_state.derived_vars) == 0:
        st.sidebar.info("No derived variables yet.")
    else:
        # Include an explicit "New" option so the sidebar doesn't auto-load an existing spec
        # right after you click New/Save (radio widgets always need a selection).
        existing = [f"{i+1}. {d.get('name','')} — {d.get('label','')}" for i, d in enumerate(st.session_state.derived_vars)]
        opts = ["➕ (New derived variable)"] + existing

        # index 0 = New; otherwise index = selected_index + 1
        sel_idx = 0 if st.session_state.m25_selected_index is None else (int(st.session_state.m25_selected_index) + 1)
        sel_idx = min(max(sel_idx, 0), len(opts) - 1)

        sel = st.sidebar.radio("Select", options=opts, index=sel_idx, key=f"m25_radio_{st.session_state.m25_rev}")

        # If user picks "New", reset editor and do NOT load any existing spec.
        if sel == opts[0]:
            # Only trigger a reset when switching FROM an existing selection to New.
            if st.session_state.m25_selected_index is not None or st.session_state.m25_last_sidebar_sel is not None:
                st.session_state.m25_selected_index = None
                st.session_state.m25_last_sidebar_sel = None
                st.session_state.m25_reset_editor = True
                _bump_rev()
                st.rerun()
        else:
            new_index = opts.index(sel) - 1  # shift down because of the New option

            if st.session_state.m25_last_sidebar_sel != new_index:
                st.session_state.m25_last_sidebar_sel = new_index
                st.session_state.m25_selected_index = new_index
                st.session_state.m25_load_index = new_index
                _bump_rev()
                st.rerun()

            st.sidebar.divider()
            c1, c2 = st.sidebar.columns(2)

            if c1.button("🗑️ Delete", use_container_width=True, key="m25_del"):
                idx = st.session_state.m25_last_sidebar_sel
                if idx is not None and 0 <= idx < len(st.session_state.derived_vars):
                    nm = _safe_str(st.session_state.derived_vars[idx].get("name"))
                    st.session_state.derived_vars.pop(idx)
                    if nm and nm in st.session_state.var_catalog and st.session_state.var_catalog[nm].get("origin") == "m25":
                        st.session_state.var_catalog.pop(nm, None)

                st.session_state.m25_last_sidebar_sel = None
                st.session_state.m25_selected_index = None
                st.session_state.m25_reset_editor = True
                _bump_rev()
                st.rerun()

            if c2.button("📄 Duplicate", use_container_width=True, key="m25_dup"):
                idx = st.session_state.m25_last_sidebar_sel
                if idx is not None and 0 <= idx < len(st.session_state.derived_vars):
                    base = st.session_state.derived_vars[idx]
                    dup = {
                        **base,
                        "name": (_safe_str(base.get("name")) + "_COPY").strip(),
                        "label": (_safe_str(base.get("label")) + " (copy)").strip(),
                        "rules": [dict(x) for x in (base.get("rules") or [])],
                        "value_labels": [dict(x) for x in (base.get("value_labels") or [])],
                        "conditions": [dict(x) for x in (base.get("conditions") or [])],
                    }
                    st.session_state.derived_vars.append(dup)
                    new_i = len(st.session_state.derived_vars) - 1
                    st.session_state.m25_last_sidebar_sel = new_i
                    st.session_state.m25_selected_index = new_i
                    st.session_state.m25_load_index = new_i
                    _bump_rev()
                    st.rerun()
    st.sidebar.divider()
    if st.sidebar.button("➕ New derived variable", use_container_width=True, key="m25_new"):
        st.session_state.m25_selected_index = None
        st.session_state.m25_last_sidebar_sel = None
        st.session_state.m25_reset_editor = True
        _bump_rev()
        st.rerun()

    # ----------------------------
    # Variable list (registry view)
    # ----------------------------
    with st.expander("📋 Variable list (from Modules 1, 2, 2.5)", expanded=False):
        rows = []
        cat = st.session_state.get("var_catalog", {}) or {}
        for v in all_vars:
            if v in cat:
                origin = _safe_str(cat[v].get("origin"))
                label = _safe_str(cat[v].get("label"))
                ch = cat[v].get("choices") or {}
                nchoices = len(ch) if isinstance(ch, dict) else ""
            else:
                origin = ""
                label = _label_lookup_for_any_var(v)
                ch = _choices_lookup_for_any_var(v)
                nchoices = len(ch) if ch else ""

            rows.append({"var": v, "label": label, "origin": origin, "n_choices": nchoices})

        df_list = pd.DataFrame(rows, columns=["var", "label", "origin", "n_choices"])
        st.dataframe(df_list, use_container_width=True, hide_index=True)

    st.divider()

    # ----------------------------
    # Editor
    # ----------------------------
    st.subheader("Build a derived variable (rules with optional conditions)")

    c1, c2, c3 = st.columns([1.2, 2.2, 1.2])
    with c1:
        st.text_input("New variable name", key="m25_name")
    with c2:
        st.text_input("Variable label", key="m25_label")
    with c3:
        st.text_input("Default code (optional)", key="m25_default_code")

    if not st.session_state.m25_work_conditions:
        st.session_state.m25_work_conditions = [{"var": all_vars[0], "picked": set(), "manual_codes": ""}]

    st.caption(
        "Add condition slots (variables). For each rule, pick codes. "
        "If a condition has no codes selected (or manual codes empty), it is ignored for that rule."
    )

    st.markdown("### Conditions (optional per rule)")
    cc1, cc2, cc3 = st.columns([1, 1, 2])
    with cc1:
        if st.button("➕ Add condition", key=f"m25_add_cond_{st.session_state.m25_rev}"):
            st.session_state.m25_work_conditions.append({"var": all_vars[0], "picked": set(), "manual_codes": ""})
            _bump_rev()
            st.rerun()
    with cc2:
        if st.button("➖ Remove last", key=f"m25_rm_cond_{st.session_state.m25_rev}") and len(st.session_state.m25_work_conditions) > 1:
            st.session_state.m25_work_conditions.pop()
            _bump_rev()
            st.rerun()
    with cc3:
        if st.button("🧽 Clear all picks", key=f"m25_clear_picks_{st.session_state.m25_rev}"):
            for c in st.session_state.m25_work_conditions:
                c["picked"] = set()
                c["manual_codes"] = ""
            _bump_rev()
            st.rerun()

    for idx, cond in enumerate(st.session_state.m25_work_conditions):
        st.markdown(f"**Condition {idx+1}** (leave empty to ignore)")

        v = st.selectbox(
            f"Variable for condition {idx+1}",
            options=all_vars,
            index=all_vars.index(cond.get("var")) if cond.get("var") in all_vars else 0,
            key=f"m25_cond_var_{idx}_{st.session_state.m25_rev}",
        )
        cond["var"] = v

        code_to_text = _choices_lookup_for_any_var(v)  # {code:text}
        codes = list(code_to_text.keys()) if code_to_text else []
        codes_sorted = sorted(codes, key=lambda x: (_safe_int(x) if _safe_int(x) is not None else 10**9, x))

        if codes_sorted:
            st.caption(f"Known codes for {v}: {len(codes_sorted)}")
            cols = st.columns(3)
            picked = cond.get("picked") or set()
            new_picked = set(picked)

            for i, code in enumerate(codes_sorted):
                lab = _safe_str(code_to_text.get(code, ""))
                display = f"{lab} ({code})" if lab else f"{code}"
                k = f"m25_cb_{idx}_{_key_safe(v)}_{_key_safe(code)}_{st.session_state.m25_rev}"
                with cols[i % 3]:
                    checked = st.checkbox(display, value=(code in picked), key=k)
                if checked:
                    new_picked.add(code)
                else:
                    new_picked.discard(code)

            cond["picked"] = new_picked
            cond["manual_codes"] = ""
        else:
            st.warning(f"No choice list found for {v}. Enter codes manually for this condition.")
            manual = st.text_input(
                f"Codes for {v} (comma-separated)",
                key=f"m25_manual_codes_{idx}_{st.session_state.m25_rev}",
                value=_safe_str(cond.get("manual_codes", "")),
                placeholder="e.g., 1,2,3",
            )
            cond["manual_codes"] = manual
            cond["picked"] = set()

        st.markdown("---")

    # Add rule UI
    st.markdown("### Add a rule")
    with st.form("m25_add_rule_form", clear_on_submit=False):
        r1, r2 = st.columns([1, 2])
        with r1:
            st.text_input("Set derived code to", key="m25_set_to", placeholder="1")
        with r2:
            st.text_input("Value label (optional but recommended)", key="m25_set_label", placeholder="Some label...")
        add_rule_clicked = st.form_submit_button("➕ Add rule")

    if add_rule_clicked:
        new_name = _safe_str(st.session_state.m25_name)
        set_to = _safe_str(st.session_state.m25_set_to)
        set_label = _safe_str(st.session_state.m25_set_label)

        if not new_name:
            st.error("New variable name is required.")
        elif not set_to:
            st.error("Derived code is required.")
        else:
            active_conds = []
            for c in (st.session_state.m25_work_conditions or []):
                picked = c.get("picked") or set()
                manual = _safe_str(c.get("manual_codes"))
                if picked or manual:
                    active_conds.append(c)

            if len(active_conds) == 0:
                st.error("Pick at least one code in at least one condition (or enter manual codes).")
            else:
                when = []
                for c in active_conds:
                    v = _safe_str(c.get("var"))
                    manual = _safe_str(c.get("manual_codes"))
                    if manual:
                        vals = [x.strip() for x in manual.split(",") if x.strip()]
                    else:
                        vals = list(c.get("picked") or set())

                    vals = [str(x).strip() for x in vals if str(x).strip() != ""]
                    vals = sorted(vals, key=lambda vv: (_safe_int(vv) if _safe_int(vv) is not None else 10**9, vv))

                    if v and vals:
                        when.append({"var": v, "values": vals})

                if not when:
                    st.error("Your rule had no valid conditions after parsing. Check your picks/manual codes.")
                else:
                    st.session_state.m25_work_rules.append({
                        "when": when,
                        "set_to": set_to,
                        "set_label": set_label,
                    })

                    if set_label:
                        vls = st.session_state.m25_work_vlabels
                        vls = [v for v in vls if _safe_str(v.get("code")) != set_to]
                        vls.append({"code": set_to, "text": set_label})
                        st.session_state.m25_work_vlabels = vls

                    st.session_state.m25_clear_rule_inputs = True
                    st.success("Rule added.")
                    _bump_rev()
                    st.rerun()

    # Show rules
    if st.session_state.m25_work_rules:
        st.markdown("### Current rules")
        for i, rule in enumerate(st.session_state.m25_work_rules, start=1):
            parts = []
            for cond in (rule.get("when") or []):
                parts.append(f"{cond.get('var')} in {{{', '.join(cond.get('values', []))}}}")
            cond_text = " AND ".join(parts) if parts else "(no conditions)"
            st.write(f"{i}. IF {cond_text}  →  **{rule.get('set_to','')}**")

        rr1, rr2 = st.columns([1, 3])
        with rr1:
            if st.button("🗑️ Delete last rule", key=f"m25_del_last_rule_{st.session_state.m25_rev}"):
                st.session_state.m25_work_rules.pop()
                _bump_rev()
                st.rerun()
        with rr2:
            st.caption("Rules apply in order. Later rules overwrite earlier ones if both match.")

    # Value labels editor
    st.markdown("### Value labels")
    vdf = pd.DataFrame(st.session_state.m25_work_vlabels or [], columns=["code", "text"])
    st.session_state.m25_work_vlabels = st.data_editor(
        vdf,
        use_container_width=True,
        hide_index=True,
        num_rows="dynamic",
        column_config={
            "code": st.column_config.TextColumn("Code"),
            "text": st.column_config.TextColumn("Label"),
        },
        key=f"m25_vl_editor_{st.session_state.m25_rev}",
    ).fillna("").to_dict("records")

    # Save / Update
    st.divider()
    s1, s2 = st.columns(2)


    def _build_spec_payload():
        # Build value labels from:
        # 1) explicit value-label edits (m25_work_vlabels)
        # 2) any rule outputs (set_to / set_label)
        # 3) default_code (if provided)
        vlabels_map = {}

        # (1) manual value labels first (highest priority)
        for v in (st.session_state.m25_work_vlabels or []):
            if not isinstance(v, dict):
                continue
            c = _safe_str(v.get("code"))
            lab = _safe_str(v.get("text"))
            if c and lab:
                vlabels_map[c] = lab

        # (2) rule outputs
        for r in (st.session_state.m25_work_rules or []):
            if not isinstance(r, dict):
                continue
            c = _safe_str(r.get("set_to"))
            lab = _safe_str(r.get("set_label"))
            if c and c not in vlabels_map:
                vlabels_map[c] = lab or c

        # (3) default code (falls back to code as label)
        dcode = _safe_str(st.session_state.m25_default_code)
        if dcode and dcode not in vlabels_map:
            vlabels_map[dcode] = dcode

        value_labels_out = [{"code": c, "text": vlabels_map[c]} for c in vlabels_map.keys()]

        return {
            "type": "RULES",
            "name": _safe_str(st.session_state.m25_name),
            "label": _safe_str(st.session_state.m25_label),
            "default_code": dcode,
            "conditions": [{"var": _safe_str(c.get("var"))} for c in (st.session_state.m25_work_conditions or [])],
            "rules": [dict(r) for r in (st.session_state.m25_work_rules or [])],
            "value_labels": value_labels_out,
        }

    with s1:

        if st.button("💾 Save as new derived var", use_container_width=True, key=f"m25_save_new_{st.session_state.m25_rev}"):
            new_name = _safe_str(st.session_state.m25_name)
            if not new_name:
                st.error("New variable name is required.")
            elif new_name in [_safe_str(d.get("name")) for d in st.session_state.derived_vars]:
                st.error("A derived variable with that name already exists. Rename it.")
            elif not st.session_state.m25_work_rules:
                st.error("Add at least one rule first.")
            else:
                payload = _build_spec_payload()
                st.session_state.derived_vars.append(payload)

                _publish_var_to_catalog(
                    var=_safe_str(payload.get("name")),
                    label=_safe_str(payload.get("label")),
                    value_labels=(payload.get("value_labels") or []),
                    origin="m25",
                )

                st.success("Saved derived variable (spec only).")

                st.session_state.m25_selected_index = None
                st.session_state.m25_last_sidebar_sel = None
                st.session_state.m25_reset_editor = True
                _bump_rev()
                st.rerun()

    with s2:
        if st.button("✅ Update selected derived var", use_container_width=True, key=f"m25_update_{st.session_state.m25_rev}"):
            idx = st.session_state.m25_selected_index
            new_name = _safe_str(st.session_state.m25_name)

            if idx is None or idx < 0 or idx >= len(st.session_state.derived_vars):
                st.error("No derived variable selected to update.")
            elif not new_name:
                st.error("New variable name is required.")
            elif not st.session_state.m25_work_rules:
                st.error("Add at least one rule first.")
            else:
                payload = _build_spec_payload()
                st.session_state.derived_vars[idx] = payload

                _publish_var_to_catalog(
                    var=_safe_str(payload.get("name")),
                    label=_safe_str(payload.get("label")),
                    value_labels=(payload.get("value_labels") or []),
                    origin="m25",
                )

                st.success("Updated derived variable (spec only).")

                st.session_state.m25_selected_index = None
                st.session_state.m25_last_sidebar_sel = None
                st.session_state.m25_reset_editor = True
                _bump_rev()
                st.rerun()

    # ----------------------------
    # Saved derived specs preview
    # ----------------------------
    st.divider()
    st.subheader("Saved derived variables (specs)")
    if not st.session_state.derived_vars:
        st.info("No derived variables saved yet.")
    else:
        preview_rows = []
        for d in st.session_state.derived_vars:
            preview_rows.append({
                "name": _safe_str(d.get("name")),
                "label": _safe_str(d.get("label")),
                "default": _safe_str(d.get("default_code")),
                "n_rules": len(d.get("rules") or []),
                "n_vlabels": len([v for v in (d.get("value_labels") or []) if _safe_str(v.get("code")) and _safe_str(v.get("text"))]),
            })
        st.dataframe(pd.DataFrame(preview_rows), use_container_width=True, hide_index=True)

    # ----------------------------
    # SPSS syntax (bottom)
    # ----------------------------
    st.divider()
    st.subheader("SPSS syntax for derived variables")

    spss = _build_spss_syntax_for_all_derived_25()
    if not spss:
        st.info("No derived variables saved yet.")
    else:
        st.code(spss, language="spss")


















# ============================
# Module 4: Weighting (manual factors, SPSS-style) — MODULE 3 DATASET ONLY
#
# What this version does:
# ✅ Uses ONLY Module 3 output dataset from st.session_state["df_out"]
# ✅ Optional fallback: user can upload a dataset if df_out is missing
# ✅ Keeps your SPSS-style factor workflow (wvar1, wvar2, WEIGHT_CUM, etc.)
# ✅ Label lookup supports:
#    - Module 1 scripted vars (questions)
#    - Module 2 recodes (recodes) with pass-through fallback to source labels
# ✅ Project Save/Load compatible:
#    - m4_factors_by_var may be restored as DataFrames OR list-of-dict records
#    - This module auto-normalizes either format
#
# Notes:
# - Weighted % displayed as percent points and rounded to xx.xx
# - Weighted N display precision controlled by m4_weighted_n_decimals
# ============================

import numpy as np
import pandas as pd
import streamlit as st
from io import BytesIO

# ----------------------------
# State
# ----------------------------
if "m4_uploaded_df" not in st.session_state:
    st.session_state.m4_uploaded_df = None

if "m4_dfw" not in st.session_state:
    st.session_state.m4_dfw = None  # working df with weights

if "m4_weight_order" not in st.session_state:
    st.session_state.m4_weight_order = []

if "m4_stage_idx" not in st.session_state:
    st.session_state.m4_stage_idx = 0  # index into m4_weight_order

# manual factor tables per variable: var -> df(code,label,factor) OR list-of-dicts (after project load)
if "m4_factors_by_var" not in st.session_state:
    st.session_state.m4_factors_by_var = {}

# SPSS syntax output lines
if "m4_spss_lines" not in st.session_state:
    st.session_state.m4_spss_lines = []

# “freq” list controls SPSS output and bottom toplines
if "m4_freq_vars" not in st.session_state:
    st.session_state.m4_freq_vars = ["QAGE", "QSEX", "QRACE", "QPARTYID", "QCOUNTY", "cQEDUCATION"]

# order UI state
if "m4_order_list" not in st.session_state:
    st.session_state.m4_order_list = []
if "m4_order_selected" not in st.session_state:
    st.session_state.m4_order_selected = None

# passes / stage numbering offset
if "m4_stage_offset" not in st.session_state:
    st.session_state.m4_stage_offset = 0

# display precision
if "m4_weighted_n_decimals" not in st.session_state:
    st.session_state.m4_weighted_n_decimals = 8


# ----------------------------
# Helpers
# ----------------------------
def _safe_int(x):
    try:
        return int(str(x).strip())
    except Exception:
        return None


def _safe_str(x) -> str:
    try:
        if pd.isna(x):
            return ""
    except Exception:
        pass
    return str(x).strip()


def _ensure_factor_df(obj) -> pd.DataFrame:
    """
    Project load may restore m4_factors_by_var[var] as list-of-dicts.
    Normalize to a DataFrame with columns: code,label,factor.
    """
    if isinstance(obj, pd.DataFrame):
        df = obj.copy()
    elif isinstance(obj, list):
        df = pd.DataFrame(obj)
    else:
        df = pd.DataFrame(columns=["code", "label", "factor"])

    # normalize columns
    for col in ["code", "label", "factor"]:
        if col not in df.columns:
            df[col] = np.nan if col == "factor" else ""
    df = df[["code", "label", "factor"]].copy()

    df["code"] = df["code"].astype(str).str.strip()
    df["label"] = df["label"].astype(str)
    df["factor"] = pd.to_numeric(df["factor"], errors="coerce")
    return df


def _get_df_for_module4():
    """
    Module 4 uses ONLY Module 3 output:
      - Preferred: st.session_state["df_out"]
      - Fallback: uploaded df
    """
    def _has_df(obj):
        return isinstance(obj, pd.DataFrame) and not obj.empty

    df3 = st.session_state.get("df_out", None)
    up = st.session_state.get("m4_uploaded_df", None)

    if _has_df(df3):
        return df3, "Module 3 output (df_out)"
    if _has_df(up):
        return up, "Uploaded dataset"
    return None, None


def _normalize_series(df: pd.DataFrame, col: str) -> pd.Series:
    return df[col].astype(str).str.strip()


def _build_code_to_label_from_module1(qname: str) -> dict:
    qname = _safe_str(qname)
    qdef = next((q for q in st.session_state.get("questions", []) if _safe_str(q.get("qname")) == qname), None)
    m = {}
    if qdef:
        for ch in qdef.get("choices", []) or []:
            code = _safe_str(ch.get("code", ""))
            lab = _safe_str(ch.get("label", ""))
            if code:
                m[code] = lab
    return m


def _find_recode_spec_by_new_qname(new_qname: str):
    new_qname = _safe_str(new_qname)
    for r in st.session_state.get("recodes", []) or []:
        if _safe_str(r.get("new_qname")) == new_qname:
            return r
    return None


def _build_code_to_label_from_module2_recode(qname: str) -> dict:
    """
    For Module 2 recode variables:
      - map new_code -> new_text for grouped codes
      - for pass-through codes, fall back to source variable's Module 1 labels
    """
    qname = _safe_str(qname)
    rec = _find_recode_spec_by_new_qname(qname)
    if not rec:
        return {}

    out = {}
    for g in rec.get("groups", []) or []:
        nc = _safe_str(g.get("new_code", ""))
        nt = _safe_str(g.get("new_text", ""))
        if nc and nt:
            out[nc] = nt

    src_q = _safe_str(rec.get("source_qname", ""))
    src_map = _build_code_to_label_from_module1(src_q) if src_q else {}
    for code, lab in src_map.items():
        if code not in out:
            out[code] = lab

    return out


def _build_code_to_label_any(qname: str) -> dict:
    """
    Best-effort label mapping for either:
      - Module 2 recode var (cQ...)  OR
      - Module 1 scripted var (Q...)
    """
    m2 = _build_code_to_label_from_module2_recode(qname)
    if m2:
        return m2
    return _build_code_to_label_from_module1(qname)


def _init_working_df(df: pd.DataFrame):
    """
    Initialize working df only once per session OR after reset.
    IMPORTANT: If you load a new dataset, m4_dfw should be reset externally or by button.
    """
    if st.session_state.m4_dfw is None:
        dfw = df.copy()
        dfw["WEIGHT_CUM"] = 1.0
        st.session_state.m4_dfw = dfw


def _weighted_freq_table(dfw: pd.DataFrame, var: str, n_decimals: int = 3) -> pd.DataFrame:
    if var not in dfw.columns:
        return pd.DataFrame()

    s = _normalize_series(dfw, var)
    w = dfw["WEIGHT_CUM"].astype(float)

    wt = w.groupby(s).sum()
    uw = s.value_counts(dropna=False)
    total_wt = float(wt.sum()) if len(wt) else 0.0

    codes = sorted(
        wt.index.tolist(),
        key=lambda x: (_safe_int(x) if _safe_int(x) is not None else 10**9, str(x)),
    )

    code_to_label = _build_code_to_label_any(var)

    rows = []
    for c in codes:
        cs = str(c).strip()
        wn = float(wt.get(cs, 0.0))
        wp = (wn / total_wt) * 100.0 if total_wt > 0 else 0.0
        rows.append(
            {
                "code": cs,
                "label": code_to_label.get(cs, ""),
                "unweighted_n": int(uw.get(cs, 0)),
                "weighted_n": round(wn, n_decimals),
                "weighted_pct": round(wp, 2),
            }
        )
    return pd.DataFrame(rows)


def _build_factor_table_for_var(dfw: pd.DataFrame, var: str) -> pd.DataFrame:
    """
    Builds/refreshes the per-code factor entry table for var.
    Keeps existing factor entries if present (even after project load).
    """
    s = _normalize_series(dfw, var)
    codes_present = sorted(
        s.dropna().astype(str).str.strip().unique().tolist(),
        key=lambda x: (_safe_int(x) if _safe_int(x) is not None else 10**9, str(x)),
    )

    code_to_label = _build_code_to_label_any(var)

    fresh = pd.DataFrame(
        {
            "code": [str(c).strip() for c in codes_present],
            "label": [code_to_label.get(str(c).strip(), "") for c in codes_present],
            "factor": [np.nan for _ in codes_present],
        }
    )

    if var not in st.session_state.m4_factors_by_var:
        return fresh

    old = _ensure_factor_df(st.session_state.m4_factors_by_var[var])

    # merge old factor values by code
    keep = old[["code", "factor"]].copy()
    out = fresh.merge(keep, on="code", how="left", suffixes=("", "_old"))
    if "factor_old" in out.columns:
        out["factor"] = pd.to_numeric(out["factor"], errors="coerce").combine_first(
            pd.to_numeric(out["factor_old"], errors="coerce")
        )
        out = out.drop(columns=["factor_old"])
    return out


def _apply_stage_manual(dfw: pd.DataFrame, var: str, stage_number: int, factor_df: pd.DataFrame):
    """
    Apply manually entered factors, SPSS-style.
    Keeps wvar numbering increasing across the full session.
    """
    existing_wvars = [c for c in dfw.columns if isinstance(c, str) and c.lower().startswith("wvar")]
    existing_nums = []
    for c in existing_wvars:
        n = _safe_int(c.lower().replace("wvar", ""))
        if n is not None:
            existing_nums.append(n)
    max_wvar = max(existing_nums) if existing_nums else 0

    factor_idx = 1 if max_wvar == 0 else (max_wvar + 1)
    factor_name = f"wvar{factor_idx}"

    prev_cum_idx = None
    if max_wvar > 0:
        prev_cum_idx = max_wvar if (max_wvar % 2 == 1) else (max_wvar - 1)
    prev_cum_name = f"wvar{prev_cum_idx}" if prev_cum_idx else None

    if prev_cum_name is None:
        cum_name = factor_name
    else:
        cum_idx = factor_idx + 1 if (factor_idx % 2 == 0) else factor_idx + 1
        cum_name = f"wvar{cum_idx}"

    factor_df = factor_df.copy()
    factor_df["code"] = factor_df["code"].astype(str).str.strip()
    factor_df["factor"] = pd.to_numeric(factor_df["factor"], errors="coerce")
    m = {r["code"]: float(r["factor"]) for _, r in factor_df.iterrows() if pd.notna(r["factor"])}

    s = _normalize_series(dfw, var)
    dfw[factor_name] = s.map(lambda v: float(m.get(str(v).strip(), 1.0)))

    if prev_cum_name is None:
        dfw[cum_name] = dfw[factor_name]
    else:
        dfw[cum_name] = dfw[prev_cum_name] * dfw[factor_name]

    dfw["WEIGHT_CUM"] = dfw[cum_name]

    # ---- SPSS syntax lines
    lines = []
    if prev_cum_name is None:
        lines.append("filter off.")
        lines.append("weight off.")
        lines.append("")
        lines.append(f"*weight {var}*")
        lines.append("")
        lines.append(f"Compute {factor_name}=0.")
    else:
        lines.append("")
        lines.append(f"*weight {var}*.")
        lines.append("")
        lines.append("weight off.")
        lines.append(f"weight by {prev_cum_name}.")
        lines.append("")
        lines.append(f"compute {factor_name}=0.")

    for _, r in factor_df.iterrows():
        code = str(r["code"]).strip()
        wt = r["factor"]
        if code and pd.notna(wt):
            lines.append(f"if {var}={code} {factor_name}={float(wt):.9f}.")

    lines.append("")
    lines.append(f"weight by {factor_name}.")
    lines.append("")

    if prev_cum_name is not None:
        lines.append(f"compute {cum_name} = {prev_cum_name}*{factor_name}.")
        lines.append("")
        lines.append(f"weight by {cum_name}.")
        lines.append("")

    lines.append("freq ")
    for v in st.session_state.m4_freq_vars:
        lines.append(f"    {v}")
    lines.append(".")

    st.session_state.m4_spss_lines.extend(lines)


def _df_to_excel_bytes(df: pd.DataFrame, sheet_name: str = "Data") -> bytes:
    output = BytesIO()
    with pd.ExcelWriter(output, engine="openpyxl") as writer:
        df.to_excel(writer, index=False, sheet_name=sheet_name)
    return output.getvalue()


def _is_reasonable_weight_var(df: pd.DataFrame, col: str) -> bool:
    """
    Heuristic: allow columns that look categorical-ish.
    - exclude weight helper columns and __text columns
    - must have <= 75 unique non-blank values
    """
    c = str(col)
    if c == "WEIGHT_CUM":
        return False
    if c.lower().startswith("wvar"):
        return False
    if c.endswith("__text") or c.endswith("__Text"):
        return False
    if c not in df.columns:
        return False

    s = df[c].astype(str).str.strip()
    s = s[s != ""]
    nunq = int(s.nunique(dropna=True))
    return nunq > 0 and nunq <= 75


# ----------------------------
# Renderer
# ----------------------------

# =========================================================
# MODULE 4: Weighting history helpers (undo last applied stage)
# =========================================================
import copy as _copy

_M4_HISTORY_KEYS = [
    "m4_dfw",
    "m4_stage_idx",
    "m4_stage_offset",
    "m4_order_list",
    "m4_order_selected",
    "m4_weight_order",
    "m4_spss_lines",
    "m4_factors_by_var",
    "m4_weighted_n_decimals",
]

def _m4_ensure_history():
    if "m4_history" not in st.session_state or not isinstance(st.session_state.m4_history, list):
        st.session_state.m4_history = []

def _m4_snapshot():
    snap = {}
    for k in _M4_HISTORY_KEYS:
        if hasattr(st.session_state, k):
            snap[k] = _copy.deepcopy(getattr(st.session_state, k))
    return snap

def _m4_restore(snap: dict):
    for k in _M4_HISTORY_KEYS:
        if k in snap:
            setattr(st.session_state, k, _copy.deepcopy(snap[k]))
        else:
            if hasattr(st.session_state, k):
                delattr(st.session_state, k)

def _m4_push_history():
    _m4_ensure_history()
    st.session_state.m4_history.append(_m4_snapshot())

def render_module_4():
    st.header("Module 4: Weighting (enter factors manually)")

    # ---- Undo history (stores snapshots after each applied weight stage) ----
    _m4_ensure_history()
    if len(st.session_state.m4_history) == 0:
        _m4_push_history()  # baseline snapshot


    # Always use Module 3 dataset (df_out). Allow upload only if df_out missing.
    df3 = st.session_state.get("df_out", None)
    if not (isinstance(df3, pd.DataFrame) and not df3.empty):
        st.warning("Module 3 output (df_out) not found. Upload a dataset to continue weighting.")
        up = st.file_uploader(
            "Upload Module 3 output (.csv or .xlsx)",
            type=["csv", "xlsx", "xls"],
            key="m4_upload",
        )
        if up is not None:
            if up.name.lower().endswith(".csv"):
                st.session_state.m4_uploaded_df = pd.read_csv(up)
            else:
                st.session_state.m4_uploaded_df = pd.read_excel(up)

    df, df_label = _get_df_for_module4()
    if df is None:
        st.info("No dataset available. Run Module 3 to produce df_out, or upload the exported file here.")
        return

    _init_working_df(df)
    dfw = st.session_state.m4_dfw
    st.caption(f"Using dataset: **{df_label}** — {dfw.shape[0]:,} rows × {dfw.shape[1]:,} columns")

    # Display formatting controls
    st.subheader("Display formatting")
    cD1, cD2 = st.columns([1, 3])
    with cD1:
        st.session_state.m4_weighted_n_decimals = st.number_input(
            "Weighted N decimals",
            min_value=0,
            max_value=8,
            value=int(st.session_state.m4_weighted_n_decimals),
            step=1,
        )
    with cD2:
        st.caption("Weighted % is shown as percent-points with 2 decimals (xx.xx).")

    # Reset
    cR1, cR2 = st.columns([1, 3])
    with cR1:
        if st.button("🧼 Reset weighting session", use_container_width=True):
            st.session_state.m4_dfw = None
            st.session_state.m4_factors_by_var = {}
            st.session_state.m4_spss_lines = []
            st.session_state.m4_stage_idx = 0
            st.session_state.m4_order_list = []
            st.session_state.m4_order_selected = None
            st.session_state.m4_weight_order = []
            st.session_state.m4_stage_offset = 0
            st.success("Reset.")
            st.rerun()
    with cR2:
        st.write("Reset clears weights + factor entries + SPSS syntax for this session.")

    st.divider()

    # ----------------------------
    # Eligible variables
    # ----------------------------
    eligible = [c for c in dfw.columns if _is_reasonable_weight_var(dfw, c)]
    eligible = sorted(eligible)

    if not eligible:
        st.warning("No eligible variables found in the dataset (categorical-ish columns with <= 75 categories).")
        return

    # Freq list
    st.subheader("Topline / freq list (shown after each weight)")
    st.session_state.m4_freq_vars = st.multiselect(
        "Variables to show in the freq block + bottom toplines",
        options=eligible,
        default=[v for v in st.session_state.m4_freq_vars if v in eligible],
        key="m4_freq_pick",
    )

    st.divider()

    # ----------------------------
    # Weight order selection (FORM) + reorder controls
    # ----------------------------
    st.subheader("Weight order")

    with st.form("m4_order_form"):
        picked = st.multiselect(
            "Variables to weight (pick all that apply)",
            options=eligible,
            default=[v for v in st.session_state.m4_order_list if v in eligible],
            help="Pick the set of variables you will weight. Order is set below.",
            key="m4_order_pick",
        )
        save_order = st.form_submit_button("💾 Save selection")

    if save_order:
        old = [v for v in st.session_state.m4_order_list if v in picked]
        new_only = [v for v in picked if v not in old]
        st.session_state.m4_order_list = old + new_only
        if st.session_state.m4_stage_idx >= len(st.session_state.m4_order_list):
            st.session_state.m4_stage_idx = 0
        st.success("Saved. Now set the order below.")

    if not st.session_state.m4_order_list:
        st.info("Pick variables above, click Save selection, then order them below.")
        return

    left, right = st.columns([2, 1])
    with left:
        st.markdown("**Current order**")
        order_opts = [f"{i+1}. {v}" for i, v in enumerate(st.session_state.m4_order_list)]
        # safe index
        if st.session_state.m4_order_selected is None or st.session_state.m4_order_selected >= len(order_opts):
            st.session_state.m4_order_selected = 0
        sel = st.selectbox(
            "Select an item to move",
            options=order_opts,
            index=st.session_state.m4_order_selected,
            key="m4_order_sel",
            label_visibility="collapsed",
        )
        st.session_state.m4_order_selected = order_opts.index(sel)

    with right:
        idx = st.session_state.m4_order_selected
        can_up = idx > 0
        can_down = idx < len(st.session_state.m4_order_list) - 1

        if st.button("⬆️ Move up", use_container_width=True, disabled=not can_up):
            lst = st.session_state.m4_order_list
            lst[idx - 1], lst[idx] = lst[idx], lst[idx - 1]
            st.session_state.m4_order_list = lst
            st.session_state.m4_order_selected = idx - 1
            st.rerun()

        if st.button("⬇️ Move down", use_container_width=True, disabled=not can_down):
            lst = st.session_state.m4_order_list
            lst[idx + 1], lst[idx] = lst[idx], lst[idx + 1]
            st.session_state.m4_order_list = lst
            st.session_state.m4_order_selected = idx + 1
            st.rerun()

        if st.button("🗑️ Remove", use_container_width=True):
            lst = st.session_state.m4_order_list
            removed = lst.pop(idx)
            st.session_state.m4_order_list = lst
            st.session_state.m4_order_selected = 0 if lst else None
            if st.session_state.m4_stage_idx >= len(lst):
                st.session_state.m4_stage_idx = 0
            st.success(f"Removed {removed}")
            st.rerun()

    st.session_state.m4_weight_order = st.session_state.m4_order_list

    st.divider()

    # ----------------------------
    # Stage navigation + apply
    # ----------------------------
    st.subheader("Current stage")

    nav1, nav2, nav3, nav4 = st.columns([1, 1, 2, 2])
    with nav1:
        if st.button("⬅️ Previous", use_container_width=True, disabled=(st.session_state.m4_stage_idx <= 0)):
            st.session_state.m4_stage_idx = max(0, st.session_state.m4_stage_idx - 1)
            st.rerun()
    with nav2:
        if st.button(
            "Next ➡️",
            use_container_width=True,
            disabled=(st.session_state.m4_stage_idx >= len(st.session_state.m4_weight_order) - 1),
        ):
            st.session_state.m4_stage_idx = min(len(st.session_state.m4_weight_order) - 1, st.session_state.m4_stage_idx + 1)
            st.rerun()
    with nav3:
        can_undo = ("m4_history" in st.session_state and isinstance(st.session_state.m4_history, list) and len(st.session_state.m4_history) > 1)
        if st.button("↩️ Delete previous weight", use_container_width=True, disabled=not can_undo):
            # Drop current snapshot and restore the previous one
            st.session_state.m4_history.pop()
            _m4_restore(st.session_state.m4_history[-1])
            st.success("Reverted to the previous weighting stage.")
            st.rerun()
        st.write(f"Stage: **{st.session_state.m4_stage_idx + 1}** of **{len(st.session_state.m4_weight_order)}**")
    with nav4:
        done = (st.session_state.m4_stage_idx == len(st.session_state.m4_weight_order) - 1)
        st.write("Status: " + ("✅ last stage" if done else "in progress"))

    var = st.session_state.m4_weight_order[st.session_state.m4_stage_idx]
    st.markdown(f"### Weight: **{var}**")

    # Build/refresh factor entry table (preserve prior entries)
    factor_tbl = _build_factor_table_for_var(dfw, var)
    st.session_state.m4_factors_by_var[var] = factor_tbl

    # Show current topline for this variable
    if st.session_state.m4_freq_vars:
        st.markdown("**Current topline for this variable (using current WEIGHT_CUM):**")
        st.dataframe(
            _weighted_freq_table(dfw, var, n_decimals=int(st.session_state.m4_weighted_n_decimals)),
            use_container_width=True,
            hide_index=True,
        )

    # Factor entry form
    with st.form(key=f"m4_factor_form_{var}"):
        st.markdown("**Enter factor values (from your spreadsheet) for each code:**")
        edited = st.data_editor(
            st.session_state.m4_factors_by_var[var],
            use_container_width=True,
            hide_index=True,
            column_config={
                "code": st.column_config.TextColumn("Code", disabled=True),
                "label": st.column_config.TextColumn("Label", disabled=True),
                "factor": st.column_config.NumberColumn("Factor (wvar value)", format="%.9f"),
            },
            key=f"m4_factor_editor_{var}",
        )

        a1, a2, a3 = st.columns([1, 1, 2])
        with a1:
            apply_now = st.form_submit_button("✅ Apply this weight")
        with a2:
            clear_factors = st.form_submit_button("🧽 Clear factors")
        with a3:
            st.caption("Blanks default to 1.0 when applying.")

    # Save edits (also keeps project-save compatibility: stored as DataFrame in-session)
    st.session_state.m4_factors_by_var[var] = edited

    if clear_factors:
        tmp = edited.copy()
        tmp["factor"] = np.nan
        st.session_state.m4_factors_by_var[var] = tmp
        st.success("Cleared factors.")
        st.rerun()

    if apply_now:
        tmp = edited.copy()
        tmp["factor"] = pd.to_numeric(tmp["factor"], errors="coerce")
        if tmp["factor"].notna().sum() == 0:
            st.error("Enter at least one factor before applying.")
        else:
            global_stage = st.session_state.m4_stage_offset + (st.session_state.m4_stage_idx + 1)
            _apply_stage_manual(dfw, var, global_stage, edited)

            # Snapshot the post-apply state so we can undo back to the prior stage
            _m4_push_history()

            st.success("Applied weight. Toplines below reflect the new cumulative weight.")
            st.rerun()

    st.divider()

    # ----------------------------
    # Bottom toplines (weighted) — always current
    # ----------------------------
    st.subheader("Toplines (weighted) — current WEIGHT_CUM")

    if not st.session_state.m4_freq_vars:
        st.info("Select variables in the topline list above to show toplines here.")
    else:
        for v in st.session_state.m4_freq_vars:
            if v not in dfw.columns:
                continue
            st.markdown(f"**{v}**")
            t = _weighted_freq_table(dfw, v, n_decimals=int(st.session_state.m4_weighted_n_decimals))
            if t.empty:
                st.info("No categories found.")
            else:
                st.dataframe(
                    t[["code", "label", "unweighted_n", "weighted_n", "weighted_pct"]],
                    use_container_width=True,
                    hide_index=True,
                )

    st.divider()

    # ----------------------------
    # Another pass option
    # ----------------------------
    st.subheader("Run another pass")
    st.caption(
        "If you want to weight the same set of variables again (using the current WEIGHT_CUM as the baseline), "
        "click below. This does NOT reset — it simply lets you start again from stage 1."
    )

    cP1, cP2 = st.columns([1, 3])
    with cP1:
        if st.button("🔁 Start another pass", use_container_width=True):
            st.session_state.m4_stage_offset += len(st.session_state.m4_weight_order)
            st.session_state.m4_stage_idx = 0
            st.success("New pass started. Enter the next set of factors and apply again.")
            st.rerun()
    with cP2:
        st.write(f"Pass stage offset: **{st.session_state.m4_stage_offset}** (internal counter)")

    st.divider()

    # ----------------------------
    # SPSS syntax output
    # ----------------------------
    st.subheader("SPSS syntax (generated)")
    if not st.session_state.m4_spss_lines:
        st.info("No SPSS syntax generated yet. Apply at least one stage.")
    else:
        st.code("\n".join(st.session_state.m4_spss_lines), language="spss")

    st.divider()

    # ----------------------------
    # Download
    # ----------------------------
    st.subheader("Download weighted dataset")
    fmt = st.radio("Output format", ["Excel (.xlsx)", "CSV (.csv)"], horizontal=True, key="m4_out_fmt")

    if fmt.startswith("Excel"):
        st.download_button(
            "⬇️ Download weighted dataset (Excel)",
            data=_df_to_excel_bytes(dfw, sheet_name="Data"),
            file_name="weighted_dataset.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            use_container_width=True,
        )
    else:
        st.download_button(
            "⬇️ Download weighted dataset (CSV)",
            data=dfw.to_csv(index=False).encode("utf-8"),
            file_name="weighted_dataset.csv",
            mime="text/csv",
            use_container_width=True,
        )





# =========================================================
# MODULE 5: Topline shell (3-column order + optional cQ injection)
# Updates:
# ✅ FIX "do everything twice": wrap editor in st.form so edits batch-apply on submit
# ✅ Add a spacer row AFTER any cQ injection block (extra separation)
# ✅ Column 2 blank option shows as "None"
# ✅ No auto-fill fighting: inject defaults only once + only for truly new/blank inject cells
# =========================================================

import re
from io import BytesIO

import pandas as pd
import streamlit as st

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH


# ----------------------------
# Constants
# ----------------------------
EXCLUDE_CQ_INJECT = {"cqage", "cqincome", "cqeducation"}
INJECT_NONE = "None"  # <-- shown in UI for Column 2

IMAGES_PROMPT_TEXT = (
    "The following is a list of names of various people who may have been mentioned in the news recently. "
    "For each one, please indicate if you have heard of the person, and if you have, whether you have a "
    "favorable or unfavorable opinion of them. If you don’t recognize a name, please say so."
)

GREEN = tuple(int("70AD47"[j:j + 2], 16) for j in (0, 2, 4))
ORANGE = tuple(int("ED7D31"[j:j + 2], 16) for j in (0, 2, 4))


# ----------------------------
# Basic helpers
# ----------------------------
def _safe_str(x) -> str:
    if x is None:
        return ""
    try:
        if pd.isna(x):
            return ""
    except Exception:
        pass
    return str(x).strip()


def _safe_int(x):
    try:
        return int(str(x).strip())
    except Exception:
        return None


def _clean_text(text):
    if text is None:
        return ""
    t = str(text)
    t = re.sub(r"^\d+\.\s*", "", t)
    t = re.sub(r"\[FLIP[^\]]*\]", "", t, flags=re.IGNORECASE)
    t = re.sub(r"\(ROTATE[^\)]*\)", "", t, flags=re.IGNORECASE)
    t = re.sub(r"\[RANDOMIZE[^\]]*\]", "", t, flags=re.IGNORECASE)
    return t.strip()


def _is_terminate_code(code: str) -> bool:
    return "terminate" in str(code or "").strip().lower()


def _questions_map(questions: list[dict]) -> dict:
    out = {}
    for q in questions or []:
        qn = _safe_str(q.get("qname"))
        if qn:
            out[qn] = q
    return out


# ----------------------------
# Variable catalog (SlotVar options)
# ----------------------------
def _catalog_vars(questions: list[dict], recodes: list[dict], derived_25: list[dict] | None) -> list[str]:
    s = set()

    for q in questions or []:
        s.add(_safe_str(q.get("qname")))

    for r in recodes or []:
        s.add(_safe_str(r.get("new_qname")))

    for d in (derived_25 or []) or []:
        nm = _safe_str(d.get("name")) or _safe_str(d.get("qname"))
        if nm:
            s.add(nm)

    df = None
    if isinstance(st.session_state.get("m3_df_out"), pd.DataFrame):
        df = st.session_state.m3_df_out
    elif isinstance(st.session_state.get("m35_df_out"), pd.DataFrame):
        df = st.session_state.m35_df_out

    if isinstance(df, pd.DataFrame) and not df.empty:
        for c in df.columns:
            s.add(_safe_str(c))

    s = {x for x in s if x}
    return sorted(s, key=lambda x: x.lower())


def _recode_outputs(recodes: list[dict]) -> list[str]:
    out = []
    for r in recodes or []:
        nm = _safe_str(r.get("new_qname"))
        if nm:
            out.append(nm)
    return sorted(set(out), key=lambda x: x.lower())


def _find_recode_by_newq(recodes: list[dict], new_qname: str) -> dict | None:
    tgt = _safe_str(new_qname)
    if not tgt:
        return None
    return next((r for r in (recodes or []) if _safe_str(r.get("new_qname")) == tgt), None)


def _default_inject_for_slot(recodes: list[dict], slot_var: str) -> str:
    candidate = f"c{_safe_str(slot_var)}"
    if candidate.strip().lower() in EXCLUDE_CQ_INJECT:
        return INJECT_NONE
    rec = _find_recode_by_newq(recodes, candidate)
    return candidate if rec else INJECT_NONE


# ----------------------------
# Build choices for SlotVar (Module 1 OR Module 2 recode)
# ----------------------------
def _choices_from_module1(qdef: dict) -> list[tuple[str, str]]:
    out = []
    for ch in qdef.get("choices", []) or []:
        lab = _safe_str(ch.get("label"))
        code = _safe_str(ch.get("code"))
        if lab and code and not _is_terminate_code(code):
            out.append((lab, code))
    return out


def _choices_from_module2_recode(recode_def: dict, qmap: dict) -> list[tuple[str, str]]:
    out = []

    source = _safe_str(recode_def.get("source_qname"))
    groups = recode_def.get("groups", []) or []

    used_source_codes = set()
    new_value_labels = []

    for g in groups:
        nt = _safe_str(g.get("new_text"))
        nc = _safe_str(g.get("new_code"))
        if nt and nc:
            new_value_labels.append((nt, nc))
        for f in g.get("from", []) or []:
            sc = _safe_str(f.get("code"))
            if sc:
                used_source_codes.add(sc)

    out.extend(new_value_labels)

    if source and source in qmap:
        src_choices = (qmap[source].get("choices", []) or [])
        for ch in src_choices:
            sc = _safe_str(ch.get("code"))
            sl = _safe_str(ch.get("label"))
            if not sc or not sl:
                continue
            if sc in used_source_codes:
                continue
            if _is_terminate_code(sc):
                continue
            out.append((sl, sc))

    def _sort_key(pair):
        lab, code = pair
        ii = _safe_int(code)
        return (0, ii) if ii is not None else (1, code.lower(), lab.lower())

    head_n = len(new_value_labels)
    head = out[:head_n]
    tail = out[head_n:]
    tail = sorted(tail, key=_sort_key)

    return head + tail


def _get_var_label_prompt(slot: str, qmap: dict, recodes: list[dict]) -> tuple[str, str]:
    """Return (label, prompt) for a variable name used in Module 5.

    Priority:
    1) Module 1 question definition (label + prompt)
    2) Module 2 recode definition (new_label)
    3) var_catalog entry (label)
    """
    slot = _safe_str(slot)

    if slot in qmap:
        label = _safe_str(qmap[slot].get("label")) or slot
        prompt = _safe_str(qmap[slot].get("prompt"))
        return label, prompt

    rec = _find_recode_by_newq(recodes, slot)
    if rec:
        label = _safe_str(rec.get("new_label")) or slot
        return label, ""

    cat = st.session_state.get("var_catalog", {}) or {}
    if slot in cat and isinstance(cat.get(slot), dict):
        label = _safe_str(cat[slot].get("label")) or slot
        return label, ""

    return slot, ""


# ----------------------------
# Word formatting helpers
# ----------------------------
def _set_font_cell(cell, font_name="Arial", font_size=11, color=None, bold=False):
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.line_spacing = 1.0
        for run in paragraph.runs:
            run.font.name = font_name
            run.font.size = Pt(font_size)
            run.bold = bold
            if color:
                run.font.color.rgb = RGBColor(*color)
            r = run._element
            r.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def _set_row_height(row, height=259):
    tr = row._tr
    trHeight = OxmlElement("w:trHeight")
    trHeight.set(qn("w:val"), str(height))
    trHeight.set(qn("w:hRule"), "exact")
    tr.append(trHeight)


def _remove_borders(table):
    tbl = table._element
    tblPr = tbl.find(qn("w:tblPr")) or OxmlElement("w:tblPr")
    if tblPr not in tbl:
        tbl.insert(0, tblPr)

    tblBorders = OxmlElement("w:tblBorders")
    for border_name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        border = OxmlElement(f"w:{border_name}")
        border.set(qn("w:val"), "none")
        tblBorders.append(border)

    existing = tblPr.find(qn("w:tblBorders"))
    if existing is not None:
        tblPr.remove(existing)
    tblPr.append(tblBorders)


def _shade_row(row_cells, fill_hex="F2F2F2"):
    for cell in row_cells:
        tc_pr = cell._element.get_or_add_tcPr()
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), fill_hex)
        tc_pr.append(shading)


def _make_2col_table(doc: Document):
    table = doc.add_table(rows=0, cols=2)
    table.autofit = False
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.columns[0].width = Inches(4.5)
    table.columns[1].width = Inches(0.63)
    _remove_borders(table)
    return table


def _m5_add_option_row(table, text, placeholder="%", *, bold=False):
    row_cells = table.add_row().cells
    row_cells[0].text = str(text)
    row_cells[1].text = str(placeholder)
    _set_font_cell(row_cells[0], bold=bold)
    # If this is an injected (bold) row, bold the % cell too.
    _set_font_cell(row_cells[1], bold=bold)
    for para in row_cells[1].paragraphs:
        para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _set_row_height(table.rows[-1])


def _add_question_paragraph(doc: Document, label: str, prompt: str):
    label = _clean_text(label)
    prompt = _clean_text(prompt)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(11)
    p.paragraph_format.space_after = Pt(11)
    p.paragraph_format.line_spacing = 1.0

    r1 = p.add_run(f"{label}: ")
    r1.bold = True
    r1.font.name = "Arial"
    r1.font.size = Pt(11)
    r1._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")

    if prompt:
        r2 = p.add_run(prompt)
        r2.bold = False
        r2.font.name = "Arial"
        r2.font.size = Pt(11)
        r2._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")


# ----------------------------
# QIMAGE table (ONLY place we shade rows)
# ----------------------------
def _insert_qimage_table(doc: Document, image_labels: list[str]):
    headers = ["", "Total fav", "Total unfav", "Very fav", "Smwt fav",
               "Smwt unfav", "Very unfav", "No opin", "NHO", "Net fav"]

    table = doc.add_table(rows=1, cols=len(headers))
    table.autofit = False
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER

    col_widths = [1.55, 0.55, 0.56, 0.55, 0.55, 0.55, 0.53, 0.47, 0.51, 0.55]
    for i, w in enumerate(col_widths):
        table.columns[i].width = Inches(w)

    _remove_borders(table)

    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        col = GREEN if h == "Total fav" else (ORANGE if h == "Total unfav" else None)
        _set_font_cell(hdr_cells[i], bold=True, color=col)
        for para in hdr_cells[i].paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_row_height(table.rows[0])

    for idx, nm in enumerate(image_labels):
        row_cells = table.add_row().cells
        row_cells[0].text = nm
        _set_font_cell(row_cells[0])

        for j in range(1, len(headers)):
            row_cells[j].text = "%"
            bold = True if j in [1, 2] else False
            col = GREEN if j == 1 else (ORANGE if j == 2 else None)
            _set_font_cell(row_cells[j], bold=bold, color=col)
            for para in row_cells[j].paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        if idx % 2 != 0:
            _shade_row(row_cells, "F2F2F2")

        _set_row_height(table.rows[-1])


# ----------------------------
# Order table init (NO rerun autofill fighting)
# ----------------------------
def _ensure_m5_table(scripted_vars: list[str], recodes: list[dict]) -> pd.DataFrame:
    """
    m5_table_df columns: SlotVar, InjectVar, Order

    KEY:
    - Default InjectVar autofill only once (init), not on every rerun.
    - For new rows: only fill InjectVar if it is truly blank/NaN.
    """
    df = None

    if "m5_table_df" in st.session_state and isinstance(st.session_state.m5_table_df, pd.DataFrame):
        df = st.session_state.m5_table_df.copy()

    if df is None or df.empty:
        df = pd.DataFrame(
            [{"SlotVar": v, "InjectVar": INJECT_NONE, "Order": (i + 1) * 10} for i, v in enumerate(scripted_vars)],
            columns=["SlotVar", "InjectVar", "Order"],
        )

    for col in ["SlotVar", "InjectVar", "Order"]:
        if col not in df.columns:
            df[col] = INJECT_NONE if col == "InjectVar" else (10 if col == "Order" else "")

    df = df[["SlotVar", "InjectVar", "Order"]].copy()

    # robust cleaning
    df["SlotVar"] = df["SlotVar"].apply(_safe_str)
    df["InjectVar"] = df["InjectVar"].apply(_safe_str)
    df["InjectVar"] = df["InjectVar"].replace({"": INJECT_NONE}).fillna(INJECT_NONE)
    df["Order"] = pd.to_numeric(df["Order"], errors="coerce").fillna(10_000)
    df = df[df["SlotVar"] != ""].copy()

    # ---- Auto-fill defaults only once ----
    if "m5_inject_autofill_done" not in st.session_state:
        for i in range(len(df)):
            if _safe_str(df.iloc[i]["InjectVar"]) in ("", INJECT_NONE):
                slot = _safe_str(df.iloc[i]["SlotVar"])
                df.iloc[i, df.columns.get_loc("InjectVar")] = _default_inject_for_slot(recodes, slot)
        st.session_state.m5_inject_autofill_done = True

    # ---- New rows: fill only if truly blank ----
    for i in range(len(df)):
        if _safe_str(df.iloc[i]["InjectVar"]) == "":
            slot = _safe_str(df.iloc[i]["SlotVar"])
            df.iloc[i, df.columns.get_loc("InjectVar")] = _default_inject_for_slot(recodes, slot)

    st.session_state.m5_table_df = df
    return df


def _first_qimage_index(ordered_slots: list[str]) -> int | None:
    for i, v in enumerate(ordered_slots):
        vu = str(v).upper()
        if vu.startswith("QIMAGE_") or vu.startswith("CQIMAGE_"):
            return i
    return None


# ----------------------------
# Build preview rows
# ----------------------------
def _build_rows(questions: list[dict], recodes: list[dict], order_df: pd.DataFrame) -> pd.DataFrame:
    qmap = _questions_map(questions)

    tmp = order_df.copy()
    tmp["Order"] = pd.to_numeric(tmp["Order"], errors="coerce").fillna(10_000)
    tmp["SlotVar"] = tmp["SlotVar"].apply(_safe_str)
    tmp["InjectVar"] = tmp["InjectVar"].apply(_safe_str).replace({"": INJECT_NONE}).fillna(INJECT_NONE)
    tmp = tmp[tmp["SlotVar"] != ""].copy()
    tmp = tmp.sort_values(["Order", "SlotVar"], ascending=[True, True])

    ordered_slots = tmp["SlotVar"].tolist()
    qimage_vars = [v for v in ordered_slots if v.upper().startswith("QIMAGE_")]
    qimage_labels = []
    for v in qimage_vars:
        qd = qmap.get(v, {})
        qimage_labels.append((_safe_str(qd.get("label")) or v))
    st.session_state.m5_qimage_labels = qimage_labels

    inject_map = {_safe_str(r["SlotVar"]): _safe_str(r["InjectVar"]) for _, r in tmp.iterrows() if _safe_str(r["SlotVar"])}

    rows = []
    rid = 1

    qimage_pos = _first_qimage_index(ordered_slots)
    inserted_qimage = False

    blocks = []
    if qimage_pos is not None:
        blocks.append("__QIMAGE_TABLE__")
    for slot in ordered_slots:
        if slot.upper().startswith("QIMAGE_"):
            continue
        blocks.append(slot)

    def _is_last_block(name: str) -> bool:
        return bool(blocks) and blocks[-1] == name

    for si, slot in enumerate(ordered_slots):
        slot = _safe_str(slot)
        var_order = (si + 1) * 10

        if (qimage_pos is not None) and (si == qimage_pos) and (not inserted_qimage):
            rows.append({
                "RowID": f"RID{rid:06d}", "VarOrder": var_order, "Var": "__QIMAGE_TABLE__",
                "RowType": "qimage_header", "Text": "Images", "Value": "", "Bold": False
            }); rid += 1
            rows.append({
                "RowID": f"RID{rid:06d}", "VarOrder": var_order, "Var": "__QIMAGE_TABLE__",
                "RowType": "qimage_placeholder", "Text": "[QIMAGE_TABLE_PLACEHOLDER]", "Value": "", "Bold": False
            }); rid += 1
            if not _is_last_block("__QIMAGE_TABLE__"):
                rows.append({
                    "RowID": f"RID{rid:06d}", "VarOrder": var_order, "Var": "__QIMAGE_TABLE__",
                    "RowType": "blank", "Text": "", "Value": "", "Bold": False
                }); rid += 1
            inserted_qimage = True

        if slot.upper().startswith("QIMAGE_"):
            continue

        label, prompt = _get_var_label_prompt(slot, qmap, recodes)

        rows.append({
            "RowID": f"RID{rid:06d}", "VarOrder": var_order, "Var": slot,
            "RowType": "header", "Text": f"{slot}: {label}: {prompt}".strip(), "Value": "", "Bold": False
        }); rid += 1

        inject_var = _safe_str(inject_map.get(slot, INJECT_NONE))
        injected_any = False

        # ---- Inject cQ groups (bold) ----
        if inject_var and inject_var != INJECT_NONE:
            rec = _find_recode_by_newq(recodes, inject_var)
            if rec:
                for g in rec.get("groups", []) or []:
                    t = _safe_str(g.get("new_text"))
                    c = _safe_str(g.get("new_code"))
                    if t and c:
                        rows.append({
                            "RowID": f"RID{rid:06d}", "VarOrder": var_order, "Var": slot,
                            "RowType": "choice", "Text": t, "Value": c, "Bold": True
                        }); rid += 1
                        injected_any = True

        # ✅ Spacer row AFTER injection block (extra separation)
        # (keeps same table; does NOT end question)
        if injected_any:
            rows.append({
                "RowID": f"RID{rid:06d}", "VarOrder": var_order, "Var": slot,
                "RowType": "spacer", "Text": "", "Value": "", "Bold": False
            }); rid += 1

        # ---- Normal labels ----
        choices = []
        if slot in qmap:
            choices = _choices_from_module1(qmap[slot])
        else:
            rec_slot = _find_recode_by_newq(recodes, slot)
            if rec_slot:
                choices = _choices_from_module2_recode(rec_slot, qmap)
            else:
                # Module 2.5 (and other non-M1/M2) variables: use var_catalog choices
                cat = st.session_state.get("var_catalog", {}) or {}
                if slot in cat and isinstance(cat.get(slot), dict):
                    ch = cat[slot].get("choices") or {}
                    if isinstance(ch, dict) and ch:
                        # list[(label, code)] like Module 1/2
                        tmp = []
                        for code, lab in ch.items():
                            code_s = _safe_str(code)
                            lab_s = _safe_str(lab)
                            if code_s and lab_s and not _is_terminate_code(code_s):
                                tmp.append((lab_s, code_s))
                        # sort codes numerically when possible
                        def _sort_key(pair):
                            lab, code = pair
                            ii = _safe_int(code)
                            return (0, ii) if ii is not None else (1, code.lower(), lab.lower())
                        choices = sorted(tmp, key=_sort_key)

        if choices:
            for (opt, code) in choices:
                rows.append({
                    "RowID": f"RID{rid:06d}", "VarOrder": var_order, "Var": slot,
                    "RowType": "choice", "Text": opt, "Value": code, "Bold": False
                }); rid += 1
        else:
            rows.append({
                "RowID": f"RID{rid:06d}", "VarOrder": var_order, "Var": slot,
                "RowType": "choice",
                "Text": "(No labels found for this variable — add labels in Module 1 or recode it in Module 2.)",
                "Value": "",
                "Bold": False
            }); rid += 1

        if not _is_last_block(slot):
            rows.append({
                "RowID": f"RID{rid:06d}", "VarOrder": var_order, "Var": slot,
                "RowType": "blank", "Text": "", "Value": "", "Bold": False
            }); rid += 1

    return pd.DataFrame(rows, columns=["RowID", "VarOrder", "Var", "RowType", "Text", "Value", "Bold"])


# ----------------------------
# Render
# ----------------------------
def render_module_5():
    st.header("Module 5: Topline shell")

    questions = st.session_state.get("questions", []) or []
    recodes = st.session_state.get("recodes", []) or []

    derived_25 = (
        st.session_state.get("derived_vars_25")
        or st.session_state.get("derived_vars")
        or st.session_state.get("derived_variables_25")
        or []
    )

    if not questions and not recodes and not derived_25 and not isinstance(st.session_state.get("m3_df_out"), pd.DataFrame):
        st.info("No variables found yet. Build variables in Modules 1/2/2.5 (or load a dataset in Module 3) first.")
        return

    # ---------------- Header inputs (lightweight; ok outside form) ----------------
    st.subheader("Header (optional)")
    c1, c2, c3 = st.columns([2, 1, 1])
    with c1:
        survey_title = st.text_input("Survey title", value=st.session_state.get("project", {}).get("name", ""))
    with c2:
        sample_n = st.text_input("Sample size (N)", value="")
    with c3:
        field_dates = st.text_input("Field dates", value="")

    st.divider()

    catalog_vars = _catalog_vars(questions, recodes, derived_25)
    if not catalog_vars:
        st.error("No variables available to place in Column 1 yet.")
        return

    qnames_scripted = [_safe_str(q.get("qname")) for q in questions if _safe_str(q.get("qname"))]
    order_df = _ensure_m5_table(qnames_scripted, recodes)

    slot_options = catalog_vars
    inject_options = [INJECT_NONE] + _recode_outputs(recodes)  # <-- Column 2 shows None

    st.subheader("Topline order + optional cQ injection")
    st.caption(
        "Edits are batch-applied only when you click **Save** (no rerun per cell). "
        "Column 2 defaults to c{SlotVar} when that recode exists (except cQAGE/cQINCOME/cQEDUCATION). "
        "Set Column 2 to **None** to cancel injection for that variable."
    )

    # =========================================================
    # ✅ FORM: stops rerun while you edit the table
    # =========================================================
    with st.form("m5_order_form", clear_on_submit=False):
        edited = st.data_editor(
            order_df,
            use_container_width=True,
            hide_index=True,
            num_rows="dynamic",
            column_config={
                "SlotVar": st.column_config.SelectboxColumn("Variable", options=slot_options, required=True),
                "InjectVar": st.column_config.SelectboxColumn("Inject cQ (optional)", options=inject_options, required=False),
                "Order": st.column_config.NumberColumn("Order", step=10),
            },
            key="m5_table_editor_v4",
        )

        cA, cB = st.columns(2)
        with cA:
            save_only = st.form_submit_button("💾 Save order", use_container_width=True)
        with cB:
            save_and_preview = st.form_submit_button("💾 Save + 🔄 Build preview", use_container_width=True)

    if save_only or save_and_preview:
        edited = edited.copy()
        edited["SlotVar"] = edited["SlotVar"].apply(_safe_str)
        edited["InjectVar"] = edited["InjectVar"].apply(_safe_str).replace({"": INJECT_NONE}).fillna(INJECT_NONE)
        edited["Order"] = pd.to_numeric(edited["Order"], errors="coerce").fillna(10_000)
        edited = edited[edited["SlotVar"] != ""].copy()
        st.session_state.m5_table_df = edited
        st.success("Order saved.")

        if save_and_preview:
            st.session_state.m5_rows_df = _build_rows(questions, recodes, st.session_state.m5_table_df)
            st.success("Preview built.")

    st.divider()

    # ---------------- Preview ----------------
    st.subheader("Preview (Module 1 style)")
    st.caption("Injected rows are bold in Word. Spacer row appears after injection for separation.")

    if "m5_rows_df" not in st.session_state:
        st.session_state.m5_rows_df = pd.DataFrame()

    df = st.session_state.m5_rows_df.copy()
    if df.empty:
        st.info("Click **Save + Build preview** to generate the shell preview.")
        return

    edited_rows = st.data_editor(
        df,
        use_container_width=True,
        hide_index=True,
        num_rows="fixed",
        column_order=["Text", "Value"],
        column_config={
            "Text": st.column_config.TextColumn(label="Text"),
            "Value": st.column_config.TextColumn(label="Value"),
        },
        key="m5_export_editor_simple_v4",
    )
    st.session_state.m5_rows_df = edited_rows

    st.divider()

    # ---------------- Word build ----------------
    st.subheader("Download Word topline shell")

    if st.button("📝 Build Word topline shell", use_container_width=True, key="m5_build_word_v4"):
        df_final = st.session_state.m5_rows_df.copy()
        if df_final.empty:
            st.error("No rows to build.")
            return

        df_final["VarOrder"] = pd.to_numeric(df_final.get("VarOrder"), errors="coerce").fillna(10_000)
        df_final = df_final.sort_values(["VarOrder", "RowID"], ascending=[True, True])

        doc = Document()

        if any([survey_title.strip(), sample_n.strip(), field_dates.strip()]):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(8)
            p.paragraph_format.line_spacing = 1.0

            r = p.add_run(survey_title.strip() if survey_title.strip() else "Topline")
            r.bold = True
            r.font.name = "Arial"
            r.font.size = Pt(12)
            r._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")

            meta_parts = []
            if sample_n.strip():
                meta_parts.append(f"N = {sample_n.strip()}")
            if field_dates.strip():
                meta_parts.append(field_dates.strip())

            if meta_parts:
                p2 = doc.add_paragraph()
                p2.paragraph_format.space_before = Pt(0)
                p2.paragraph_format.space_after = Pt(12)
                p2.paragraph_format.line_spacing = 1.0
                r2 = p2.add_run(" • ".join(meta_parts))
                r2.font.name = "Arial"
                r2.font.size = Pt(11)
                r2._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")

        qimage_labels = st.session_state.get("m5_qimage_labels", []) or []

        current_table = None
        qimage_inserted = False

        for _, r in df_final.iterrows():
            row_type = _safe_str(r.get("RowType"))
            text = _safe_str(r.get("Text"))
            bold = bool(r.get("Bold", False))

            if row_type == "qimage_header":
                _add_question_paragraph(doc, "Images", IMAGES_PROMPT_TEXT)
                current_table = None
                continue

            if row_type == "qimage_placeholder":
                if (not qimage_inserted) and qimage_labels:
                    _insert_qimage_table(doc, qimage_labels)
                qimage_inserted = True
                continue

            if row_type == "header":
                current_table = None
                t = _clean_text(text)
                parts = [p.strip() for p in t.split(":", 2)]
                if len(parts) == 3:
                    _, lab, prm = parts
                elif len(parts) == 2:
                    lab, prm = parts
                else:
                    lab, prm = t, ""
                _add_question_paragraph(doc, lab, prm)
                current_table = _make_2col_table(doc)
                continue

            if current_table is None:
                continue

            if row_type == "blank":
                current_table = None
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(6)
                p.paragraph_format.line_spacing = 0
                continue

            # ✅ Spacer row stays inside the table
            if row_type == "spacer":
                _m5_add_option_row(current_table, "", placeholder="", bold=False)
                continue

            _m5_add_option_row(current_table, text, placeholder="%", bold=bold)

        if (not qimage_inserted) and qimage_labels:
            _add_question_paragraph(doc, "Images", IMAGES_PROMPT_TEXT)
            _insert_qimage_table(doc, qimage_labels)

        buf = BytesIO()
        doc.save(buf)
        buf.seek(0)

        filename = (survey_title.strip().replace(" ", "_") or "topline_shell") + ".docx"
        st.download_button(
            label="⬇️ Download topline shell (.docx)",
            data=buf.getvalue(),
            file_name=filename,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
            key="m5_download_docx_v4",
        )
        st.success("Topline shell built.")











# =========================================================
# MODULE 6: Weighted topline (fills in numbers) — UPDATED
#
# Update included:
# ✅ Option to weight by NOTHING (unweighted mode)
#   - Weight picker includes "(none — unweighted)"
#   - Internally uses all-ones weights when none selected
#
# Notes:
# - NO shading anywhere except Images table (kept like Module 5)
# - Images header prompt EXACT text you provided
# - cQ rule kept:
#     If matching recode exists (new_qname == "c" + QNAME) and it is NOT cQAGE/cQINCOME/cQEDUCATION,
#     then each recode group is inserted at TOP of that Q's options and BOLDED in Word.
#     Those injected rows are computed as SUM of weights over the group’s source codes.
# - QIMAGE handled as a single table inserted where the first QIMAGE appears in order,
#   and it fills Total fav/unfav/etc using standard image coding (1-6).
#
# Output:
# - Download a .docx with weighted topline numbers filled in.
# =========================================================

import re
import pandas as pd
import numpy as np
import streamlit as st
from io import BytesIO

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH


# ----------------------------
# Constants / formatting
# ----------------------------
EXCLUDE_CQ_INJECT = {"cqage", "cqincome", "cqeducation"}

IMAGES_PROMPT_TEXT = (
    "The following is a list of names of various people who may have been mentioned in the news recently. "
    "For each one, please indicate if you have heard of the person, and if you have, whether you have a "
    "favorable or unfavorable opinion of them. If you don’t recognize a name, please say so."
)

# Image codes (your standard)
IMG_VF = "1"
IMG_SF = "2"
IMG_SU = "3"
IMG_VU = "4"
IMG_NO = "5"
IMG_NHO = "6"

GREEN = tuple(int("70AD47"[j:j+2], 16) for j in (0, 2, 4))
ORANGE = tuple(int("ED7D31"[j:j+2], 16) for j in (0, 2, 4))


def _fmt_pct(x: float) -> str:
    if x is None or (isinstance(x, float) and np.isnan(x)):
        return ""

    # true zero stays 0%
    if x == 0:
        return "0%"

    r = round(x)

    # if it rounds to 0 but isn't actually 0 → "<1%"
    if r == 0:
        return "<1%"

    return f"{r}%"


def _fmt_wn(x: float) -> str:
    if x is None or (isinstance(x, float) and np.isnan(x)):
        return ""
    # show a few decimals per your earlier preference
    return f"{x:.3f}"


def _clean_text(text):
    if text is None:
        return ""
    t = str(text)
    t = re.sub(r"^\d+\.\s*", "", t)
    t = re.sub(r"\[FLIP[^\]]*\]", "", t, flags=re.IGNORECASE)
    t = re.sub(r"\(ROTATE[^\)]*\)", "", t, flags=re.IGNORECASE)
    t = re.sub(r"\[RANDOMIZE[^\]]*\]", "", t, flags=re.IGNORECASE)
    return t.strip()


def _is_terminate_code(code: str) -> bool:
    return "terminate" in str(code or "").strip().lower()


def _as_code_series(df: pd.DataFrame, col: str) -> pd.Series:
    # Always compare codes as stripped strings
    s = df[col]
    # If pandas sees numbers, keep them but stringify cleanly
    return s.map(lambda v: "" if pd.isna(v) else str(v).strip())


def _as_weight_series(df: pd.DataFrame, wcol: str | None) -> pd.Series:
    """
    If wcol is None/empty => unweighted mode: all-ones.
    Otherwise coerce chosen weight column to numeric.
    """
    if not wcol:
        return pd.Series(1.0, index=df.index)

    w = pd.to_numeric(df[wcol], errors="coerce").fillna(0.0)
    return w


# ----------------------------
# Word helpers
# ----------------------------
def _set_font_cell(cell, font_name="Arial", font_size=11, color=None, bold=False):
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.line_spacing = 1.0
        for run in paragraph.runs:
            run.font.name = font_name
            run.font.size = Pt(font_size)
            run.bold = bold
            if color:
                run.font.color.rgb = RGBColor(*color)
            r = run._element
            r.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def _set_row_height(row, height=259):  # 259 twips ≈ 0.18"
    tr = row._tr
    trHeight = OxmlElement("w:trHeight")
    trHeight.set(qn("w:val"), str(height))
    trHeight.set(qn("w:hRule"), "exact")
    tr.append(trHeight)


def _remove_borders(table):
    tbl = table._element
    tblPr = tbl.find(qn("w:tblPr")) or OxmlElement("w:tblPr")
    if tblPr not in tbl:
        tbl.insert(0, tblPr)

    tblBorders = OxmlElement("w:tblBorders")
    for border_name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        border = OxmlElement(f"w:{border_name}")
        border.set(qn("w:val"), "none")
        tblBorders.append(border)

    existing = tblPr.find(qn("w:tblBorders"))
    if existing is not None:
        tblPr.remove(existing)
    tblPr.append(tblBorders)


def _shade_row(row_cells, fill_hex="F2F2F2"):
    for cell in row_cells:
        tc_pr = cell._element.get_or_add_tcPr()
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), fill_hex)
        tc_pr.append(shading)


def _add_question_paragraph(doc: Document, label: str, prompt: str):
    label = _clean_text(label)
    prompt = _clean_text(prompt)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(11)
    p.paragraph_format.space_after = Pt(11)
    p.paragraph_format.line_spacing = 1.0

    r1 = p.add_run(f"{label}: ")
    r1.bold = True
    r1.font.name = "Arial"
    r1.font.size = Pt(11)
    r1._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")

    r2 = p.add_run(prompt)
    r2.bold = False
    r2.font.name = "Arial"
    r2.font.size = Pt(11)
    r2._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")


def _make_2col_table(doc: Document):
    table = doc.add_table(rows=0, cols=2)
    table.autofit = False
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.columns[0].width = Inches(4.5)
    table.columns[1].width = Inches(0.63)
    _remove_borders(table)
    return table


# Renamed to avoid collisions with Module 5
def _m6_add_option_row(table, text, value_text, *, bold_left=False):
    row_cells = table.add_row().cells
    row_cells[0].text = str(text)
    row_cells[1].text = str(value_text)
    _set_font_cell(row_cells[0], bold=bold_left)
    # If this is an injected (bold) row, bold the % cell too.
    _set_font_cell(row_cells[1], bold=bold_left)
    for para in row_cells[1].paragraphs:
        # keep consistent with your earlier preference (feel free to change to RIGHT)
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_row_height(table.rows[-1])


# ----------------------------
# Images table (ONLY shaded section)
# ----------------------------
def _insert_qimage_table_filled(doc: Document, image_labels: list[str], image_stats_by_var: dict):
    """
    image_stats_by_var: { "QIMAGE_A": {"tot_fav":.., ...}, ... } in 0-100 percent units
    """
    headers = ["", "Total fav", "Total unfav", "Very fav", "Smwt fav",
               "Smwt unfav", "Very unfav", "No opin", "NHO", "Net fav"]

    table = doc.add_table(rows=1, cols=len(headers))
    table.autofit = False
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER

    col_widths = [1.55, 0.55, 0.56, 0.55, 0.55, 0.55, 0.53, 0.47, 0.51, 0.55]
    for i, w in enumerate(col_widths):
        table.columns[i].width = Inches(w)

    _remove_borders(table)

    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        col = GREEN if h == "Total fav" else (ORANGE if h == "Total unfav" else None)
        _set_font_cell(hdr_cells[i], bold=True, color=col)
        for para in hdr_cells[i].paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_row_height(table.rows[0])

    qimage_vars = st.session_state.get("m6_qimage_vars", [])

    for idx, nm in enumerate(image_labels):
        var = qimage_vars[idx] if idx < len(qimage_vars) else None
        stats = image_stats_by_var.get(var, {}) if var else {}

        row_cells = table.add_row().cells
        row_cells[0].text = nm
        _set_font_cell(row_cells[0])

        vals = [
            _fmt_pct(stats.get("tot_fav")),
            _fmt_pct(stats.get("tot_unfav")),
            _fmt_pct(stats.get("vf")),
            _fmt_pct(stats.get("sf")),
            _fmt_pct(stats.get("su")),
            _fmt_pct(stats.get("vu")),
            _fmt_pct(stats.get("no")),
            _fmt_pct(stats.get("nho")),
            _fmt_pct(stats.get("net_fav")),
        ]

        for j, txt in enumerate(vals, start=1):
            row_cells[j].text = txt
            bold = True if j in [1, 2] else False
            col = GREEN if j == 1 else (ORANGE if j == 2 else None)
            _set_font_cell(row_cells[j], bold=bold, color=col)
            for para in row_cells[j].paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        if idx % 2 != 0:
            _shade_row(row_cells, "F2F2F2")

        _set_row_height(table.rows[-1])


# ----------------------------
# Definition helpers
# ----------------------------
def _find_qdef(questions: list[dict], qname: str) -> dict | None:
    return next((q for q in questions if str(q.get("qname", "")).strip() == qname), None)


def _find_matching_recode(recodes: list[dict], qname: str) -> dict | None:
    target = f"c{qname}"
    return next((r for r in recodes if str(r.get("new_qname", "")).strip() == target), None)


def _default_var_order_df(var_list: list[str]) -> pd.DataFrame:
    return pd.DataFrame(
        [{"Var": v, "Order": (i + 1) * 10} for i, v in enumerate(var_list)],
        columns=["Var", "Order"],
    )


def _first_qimage_index(ordered_vars: list[str]) -> int | None:
    for i, v in enumerate(ordered_vars):
        if str(v).upper().startswith("QIMAGE_"):
            return i
    return None


# ----------------------------
# Weighted computations
# ----------------------------
def _weighted_dist(df: pd.DataFrame, var: str, wcol: str | None):
    """
    Returns:
      denom_w (float): sum of weights over valid cases
      by_code_w (dict): code(str) -> sum(weights)
    """
    if var not in df.columns:
        return 0.0, {}

    codes = _as_code_series(df, var)
    w = _as_weight_series(df, wcol)

    valid = codes != ""
    denom = float(w[valid].sum())

    by_code = {}
    tmp = pd.DataFrame({"code": codes, "w": w})
    tmp = tmp[tmp["code"] != ""]
    grp = tmp.groupby("code", dropna=False)["w"].sum()
    for k, v in grp.items():
        by_code[str(k).strip()] = float(v)

    return denom, by_code


def _pct_from_w(sum_w: float, denom_w: float) -> float:
    if denom_w <= 0:
        return np.nan
    return 100.0 * (sum_w / denom_w)


def _compute_image_stats(df: pd.DataFrame, var: str, wcol: str | None):
    denom, by_code = _weighted_dist(df, var, wcol)

    vf = by_code.get(IMG_VF, 0.0)
    sf = by_code.get(IMG_SF, 0.0)
    su = by_code.get(IMG_SU, 0.0)
    vu = by_code.get(IMG_VU, 0.0)
    no = by_code.get(IMG_NO, 0.0)
    nho = by_code.get(IMG_NHO, 0.0)

    tot_fav_w = vf + sf
    tot_unfav_w = su + vu
    net_fav_pct = _pct_from_w(tot_fav_w, denom) - _pct_from_w(tot_unfav_w, denom)

    return {
        "vf": _pct_from_w(vf, denom),
        "sf": _pct_from_w(sf, denom),
        "su": _pct_from_w(su, denom),
        "vu": _pct_from_w(vu, denom),
        "no": _pct_from_w(no, denom),
        "nho": _pct_from_w(nho, denom),
        "tot_fav": _pct_from_w(tot_fav_w, denom),
        "tot_unfav": _pct_from_w(tot_unfav_w, denom),
        "net_fav": net_fav_pct,
        "denom_w": denom,
    }


# ----------------------------
# Build a "render plan" like Module 5, but with metadata needed to compute numbers
# ----------------------------
def _build_rows_plan(questions: list[dict], recodes: list[dict], ordered_vars: list[str]) -> pd.DataFrame:
    """
    Columns:
      RowID, VarOrder, Var, RowType, Text, Value, Bold, GroupCodes
    - For injected cQ group rows: Value=new_code, GroupCodes="2|3|4"
    - For normal choice rows: Value=choice code, GroupCodes=""
    """
    q_lookup = {q.get("qname"): q for q in questions if q.get("qname")}

    qimage_pos = _first_qimage_index(ordered_vars)
    inserted_qimage_placeholder = False

    # Build blocks (to prevent trailing blank at end)
    blocks = []
    qimage_vars = [v for v in ordered_vars if (str(v).upper().startswith("QIMAGE_") or str(v).upper().startswith("CQIMAGE_"))]
    if qimage_pos is not None:
        blocks.append("__QIMAGE_TABLE__")

    for v in ordered_vars:
        if str(v).upper().startswith("QIMAGE_"):
            continue
        if _find_qdef(questions, str(v).strip()):
            blocks.append(str(v).strip())

    def _is_last_block(block_name: str) -> bool:
        if not blocks:
            return True
        return blocks[-1] == block_name

    # Store image vars+labels for later
    image_labels = []
    for v in qimage_vars:
        qd = q_lookup.get(v, {})
        image_labels.append((qd.get("label") or "").strip() or v)
    st.session_state.m6_qimage_vars = qimage_vars
    st.session_state.m6_qimage_labels = image_labels

    rows = []
    rid = 1

    for vi, v in enumerate(ordered_vars):
        v = str(v).strip()
        var_order = (vi + 1) * 10

        # Insert Images block at first QIMAGE position
        if (qimage_pos is not None) and (vi == qimage_pos) and (not inserted_qimage_placeholder):
            rows.append({
                "RowID": f"RID{rid:06d}", "VarOrder": var_order, "Var": "__QIMAGE_TABLE__",
                "RowType": "qimage_header", "Text": "Images", "Value": "", "Bold": False, "GroupCodes": ""
            })
            rid += 1
            rows.append({
                "RowID": f"RID{rid:06d}", "VarOrder": var_order, "Var": "__QIMAGE_TABLE__",
                "RowType": "qimage_placeholder", "Text": "[QIMAGE_TABLE_PLACEHOLDER]", "Value": "", "Bold": False, "GroupCodes": ""
            })
            rid += 1

            if not _is_last_block("__QIMAGE_TABLE__"):
                rows.append({
                    "RowID": f"RID{rid:06d}", "VarOrder": var_order, "Var": "__QIMAGE_TABLE__",
                    "RowType": "blank", "Text": "", "Value": "", "Bold": False, "GroupCodes": ""
                })
                rid += 1

            inserted_qimage_placeholder = True

        # Skip individual image vars (QIMAGE_* and cQIMAGE_*)
        vu = v.upper()
        if vu.startswith("QIMAGE_") or vu.startswith("CQIMAGE_"):
            continue

        qdef = _find_qdef(questions, v)
        if not qdef:
            continue

        label = (qdef.get("label", "") or "").strip() or v
        prompt = (qdef.get("prompt", "") or "").strip()
        header_text = f"{v}: {label}: {prompt}".strip()

        rows.append({
            "RowID": f"RID{rid:06d}", "VarOrder": var_order, "Var": v,
            "RowType": "header", "Text": header_text, "Value": "", "Bold": False, "GroupCodes": ""
        })
        rid += 1

        # Inject cQ groups at top (bold)
        rec = _find_matching_recode(recodes, v)
        inject = bool(rec) and (str(rec.get("new_qname", "")).strip().lower() not in EXCLUDE_CQ_INJECT)

        if inject:
            for g in rec.get("groups", []) or []:
                t = str(g.get("new_text", "")).strip()
                new_code = str(g.get("new_code", "")).strip()
                from_codes = []
                for f in g.get("from", []) or []:
                    c = str(f.get("code", "")).strip()
                    if c:
                        from_codes.append(c)

                if not t or not new_code or not from_codes:
                    continue

                rows.append({
                    "RowID": f"RID{rid:06d}", "VarOrder": var_order, "Var": v,
                    "RowType": "choice_group", "Text": t, "Value": new_code,
                    "Bold": True, "GroupCodes": "|".join(from_codes)
                })
                rid += 1

        # Normal Q choices
        for ch in qdef.get("choices", []) or []:
            opt = str(ch.get("label", "")).strip()
            code = str(ch.get("code", "")).strip()
            if not opt or not code:
                continue
            if _is_terminate_code(code):
                continue
            rows.append({
                "RowID": f"RID{rid:06d}", "VarOrder": var_order, "Var": v,
                "RowType": "choice", "Text": opt, "Value": code,
                "Bold": False, "GroupCodes": ""
            })
            rid += 1

        if not _is_last_block(v):
            rows.append({
                "RowID": f"RID{rid:06d}", "VarOrder": var_order, "Var": v,
                "RowType": "blank", "Text": "", "Value": "", "Bold": False, "GroupCodes": ""
            })
            rid += 1

    return pd.DataFrame(rows, columns=["RowID", "VarOrder", "Var", "RowType", "Text", "Value", "Bold", "GroupCodes"])


def _build_rows_plan_from_table(questions: list[dict], recodes: list[dict], order_df: pd.DataFrame) -> pd.DataFrame:
    """Build a Module 6 render plan like Module 5 (SlotVar + optional InjectVar).

    Key behavior (what you asked for):
    - ONE question block per SlotVar (not separate blocks for the injected variable)
    - Injected groups appear ABOVE the SlotVar choices
    - Injected rows are emitted as RowType='choice_group' with GroupCodes so Module 6 can sum
      underlying codes when filling weighted percents.
    """
    qmap = _questions_map(questions)

    tmp = order_df.copy() if isinstance(order_df, pd.DataFrame) else pd.DataFrame(columns=["SlotVar", "InjectVar", "Order"])
    for col in ["SlotVar", "InjectVar", "Order"]:
        if col not in tmp.columns:
            tmp[col] = INJECT_NONE if col == "InjectVar" else (10_000 if col == "Order" else "")

    tmp["Order"] = pd.to_numeric(tmp["Order"], errors="coerce").fillna(10_000)
    tmp["SlotVar"] = tmp["SlotVar"].apply(_safe_str)
    tmp["InjectVar"] = tmp["InjectVar"].apply(_safe_str).replace({"": INJECT_NONE}).fillna(INJECT_NONE)
    tmp = tmp[tmp["SlotVar"] != ""].copy()
    tmp = tmp.sort_values(["Order", "SlotVar"], ascending=[True, True])

    ordered_slots = tmp["SlotVar"].tolist()
    qimage_pos = _first_qimage_index(ordered_slots)

    # labels for the Images table (used later in doc build)
    qimage_vars = [v for v in ordered_slots if str(v).upper().startswith("QIMAGE_")]
    qimage_labels = []
    for v in qimage_vars:
        qd = qmap.get(v, {})
        qimage_labels.append((_safe_str(qd.get("label")) or v))
    st.session_state.m6_qimage_vars = qimage_vars
    st.session_state.m6_qimage_labels = qimage_labels

    inject_map = {_safe_str(r["SlotVar"]): _safe_str(r["InjectVar"]) for _, r in tmp.iterrows() if _safe_str(r.get("SlotVar"))}

    # blocks to avoid trailing blank
    blocks = []
    if qimage_pos is not None:
        blocks.append("__QIMAGE_TABLE__")
    for slot in ordered_slots:
        if slot.upper().startswith("QIMAGE_"):
            continue
        blocks.append(slot)

    def _is_last_block(name: str) -> bool:
        return bool(blocks) and blocks[-1] == name

    rows = []
    rid = 1
    inserted_qimage = False

    for si, slot in enumerate(ordered_slots):
        slot = _safe_str(slot)
        var_order = (si + 1) * 10

        if (qimage_pos is not None) and (si == qimage_pos) and (not inserted_qimage):
            rows.append({
                "RowID": f"RID{rid:06d}", "VarOrder": var_order, "Var": "__QIMAGE_TABLE__",
                "RowType": "qimage_header", "Text": "Images", "Value": "", "Bold": False, "GroupCodes": ""
            }); rid += 1
            rows.append({
                "RowID": f"RID{rid:06d}", "VarOrder": var_order, "Var": "__QIMAGE_TABLE__",
                "RowType": "qimage_placeholder", "Text": "[QIMAGE_TABLE_PLACEHOLDER]", "Value": "", "Bold": False, "GroupCodes": ""
            }); rid += 1
            if not _is_last_block("__QIMAGE_TABLE__"):
                rows.append({
                    "RowID": f"RID{rid:06d}", "VarOrder": var_order, "Var": "__QIMAGE_TABLE__",
                    "RowType": "blank", "Text": "", "Value": "", "Bold": False, "GroupCodes": ""
                }); rid += 1
            inserted_qimage = True

        if slot.upper().startswith("QIMAGE_"):
            continue

        label, prompt = _get_var_label_prompt(slot, qmap, recodes)
        rows.append({
            "RowID": f"RID{rid:06d}", "VarOrder": var_order, "Var": slot,
            "RowType": "header", "Text": f"{slot}: {label}: {prompt}".strip(),
            "Value": "", "Bold": False, "GroupCodes": ""
        }); rid += 1

        inject_var = _safe_str(inject_map.get(slot, INJECT_NONE))
        injected_any = False

        # Inject groups (bold) as choice_group with GroupCodes so Module 6 can sum
        if inject_var and inject_var != INJECT_NONE:
            rec = _find_recode_by_newq(recodes, inject_var)
            if rec:
                for g in rec.get("groups", []) or []:
                    t = _safe_str(g.get("new_text"))
                    new_code = _safe_str(g.get("new_code"))
                    from_codes = []
                    for f in g.get("from", []) or []:
                        c = _safe_str(f.get("code"))
                        if c:
                            from_codes.append(c)
                    if not t or not new_code or not from_codes:
                        continue
                    rows.append({
                        "RowID": f"RID{rid:06d}", "VarOrder": var_order, "Var": slot,
                        "RowType": "choice_group", "Text": t, "Value": new_code,
                        "Bold": True, "GroupCodes": "|".join(from_codes)
                    }); rid += 1
                    injected_any = True

        # Spacer row after injection block (keeps same question)
        if injected_any:
            rows.append({
                "RowID": f"RID{rid:06d}", "VarOrder": var_order, "Var": slot,
                "RowType": "spacer", "Text": "", "Value": "", "Bold": False, "GroupCodes": ""
            }); rid += 1

        # Normal choice rows (same source logic as Module 5)
        choices = []
        if slot in qmap:
            choices = _choices_from_module1(qmap[slot])
        else:
            rec_slot = _find_recode_by_newq(recodes, slot)
            if rec_slot:
                choices = _choices_from_module2_recode(rec_slot, qmap)
            else:
                cat = st.session_state.get("var_catalog", {}) or {}
                if slot in cat and isinstance(cat.get(slot), dict):
                    ch = cat[slot].get("choices") or {}
                    if isinstance(ch, dict) and ch:
                        tmp2 = []
                        for code, lab in ch.items():
                            code_s = _safe_str(code)
                            lab_s = _safe_str(lab)
                            if code_s and lab_s and not _is_terminate_code(code_s):
                                tmp2.append((lab_s, code_s))
                        def _sort_key(pair):
                            lab, code = pair
                            ii = _safe_int(code)
                            return (0, ii) if ii is not None else (1, code.lower(), lab.lower())
                        choices = sorted(tmp2, key=_sort_key)

        if choices:
            for (opt, code) in choices:
                rows.append({
                    "RowID": f"RID{rid:06d}", "VarOrder": var_order, "Var": slot,
                    "RowType": "choice", "Text": opt, "Value": code,
                    "Bold": False, "GroupCodes": ""
                }); rid += 1
        else:
            rows.append({
                "RowID": f"RID{rid:06d}", "VarOrder": var_order, "Var": slot,
                "RowType": "choice", "Text": "(no labels found)", "Value": "",
                "Bold": False, "GroupCodes": ""
            }); rid += 1

        if not _is_last_block(slot):
            rows.append({
                "RowID": f"RID{rid:06d}", "VarOrder": var_order, "Var": slot,
                "RowType": "blank", "Text": "", "Value": "", "Bold": False, "GroupCodes": ""
            }); rid += 1

    return pd.DataFrame(rows, columns=["RowID", "VarOrder", "Var", "RowType", "Text", "Value", "Bold", "GroupCodes"])


# ----------------------------
# Main render
# ----------------------------
def render_module_6():
    st.header("Module 6: Filled topline (counts)")

    questions = st.session_state.get("questions", []) or []
    recodes = st.session_state.get("recodes", []) or []

    derived_25 = (
        st.session_state.get("derived_vars_25")
        or st.session_state.get("derived_vars")
        or st.session_state.get("derived_variables_25")
        or []
    )

    var_catalog = st.session_state.get("var_catalog", {}) or {}

    if not questions and not recodes and not derived_25 and not isinstance(st.session_state.get("m3_df_out"), pd.DataFrame):
        st.info("No variables found yet. Build variables in Modules 1/2/2.5 (or load a dataset in Module 3) first.")
        return

    # -------------------------
    # Step 1: Provide dataset (same idea as old Module 6)
    # -------------------------
    st.subheader("Step 1: Provide dataset")
    df_from_m4 = st.session_state.get("m4_df_out", None)
    wcol_from_m4 = st.session_state.get("m4_weight_col", None)

    df = None
    if isinstance(df_from_m4, pd.DataFrame) and len(df_from_m4) > 0:
        st.success(f"Using weighted dataset from Module 4 ({df_from_m4.shape[0]:,} rows, {df_from_m4.shape[1]:,} cols).")
        df = df_from_m4.copy()
    else:
        up = st.file_uploader("Upload weighted dataset (csv/xlsx)", type=["csv", "xlsx"], key="m6_upload_df")
        if up is not None:
            try:
                if up.name.lower().endswith(".csv"):
                    df = pd.read_csv(up)
                else:
                    df = pd.read_excel(up)
                st.success(f"Loaded dataset: {df.shape[0]:,} rows, {df.shape[1]:,} cols.")
            except Exception as e:
                st.error(f"Could not read file: {e}")
                return

    if df is None or not isinstance(df, pd.DataFrame) or len(df) == 0:
        st.info("Upload a dataset (or run Module 4) to continue.")
        return

    cols = df.columns.astype(str).tolist()
    weight_options = ["(none — unweighted)"] + cols

    default_w = wcol_from_m4 if (wcol_from_m4 in cols) else None
    if default_w is None:
        for guess in ["WEIGHT_CUM", "weight", "wgt", "wfinal", "wvar7", "wvar5", "wvar3", "wvar1"]:
            if guess in cols:
                default_w = guess
                break

    default_pick = "(none — unweighted)" if default_w is None else default_w
    weight_pick = st.selectbox(
        "Weight column",
        options=weight_options,
        index=weight_options.index(default_pick) if default_pick in weight_options else 0,
        key="m6_weight_pick",
    )
    weight_col = None if weight_pick.startswith("(none") else weight_pick

    st.divider()

    # ---------------- Header inputs (match Module 5) ----------------
    st.subheader("Header (optional)")
    c1, c2, c3 = st.columns([2, 1, 1])
    with c1:
        survey_title = st.text_input("Survey title", value=st.session_state.get("project", {}).get("name", ""), key="m6_survey_title")
    with c2:
        sample_n = st.text_input("Sample size (N)", value="", key="m6_sample_n")
    with c3:
        field_dates = st.text_input("Field dates", value="", key="m6_field_dates")

    st.divider()

    # ---------------- Order + inject UI (match Module 5) ----------------
    catalog_vars = _catalog_vars(questions, recodes, derived_25)
    if not catalog_vars:
        st.error("No variables available yet.")
        return

    # Create/keep a separate table df for Module 6 (so it doesn't collide with Module 5)
    if "m6_table_df" not in st.session_state or not isinstance(st.session_state.m6_table_df, pd.DataFrame):
        qnames_scripted = [_safe_str(q.get("qname")) for q in questions if _safe_str(q.get("qname"))]
        st.session_state.m6_table_df = _ensure_m5_table(qnames_scripted, recodes).copy()

    order_df = st.session_state.m6_table_df.copy()

    slot_options = list(catalog_vars)
    inject_options = [INJECT_NONE] + _recode_outputs(recodes)

    st.subheader("Topline order + optional cQ injection")
    st.caption(
        "Edits are batch-applied only when you click **Save** (no rerun per cell). "
        "Column 2 defaults to c{SlotVar} when that recode exists (except cQAGE/cQINCOME/cQEDUCATION). "
        "Set Column 2 to None to disable injection."
    )

    with st.form("m6_order_form", clear_on_submit=False):
        edited = st.data_editor(
            order_df,
            use_container_width=True,
            hide_index=True,
            num_rows="dynamic",
            column_config={
                "SlotVar": st.column_config.SelectboxColumn("Variable", options=slot_options, required=True),
                "InjectVar": st.column_config.SelectboxColumn("Inject cQ (optional)", options=inject_options, required=False),
                "Order": st.column_config.NumberColumn("Order", step=10),
            },
            key="m6_table_editor_v1",
        )

        cA, cB = st.columns(2)
        with cA:
            save_only = st.form_submit_button("💾 Save", use_container_width=True)
        with cB:
            save_and_preview = st.form_submit_button("💾 Save + Preview", use_container_width=True)

    if save_only or save_and_preview:
        edited = edited.copy()
        edited["SlotVar"] = edited["SlotVar"].apply(_safe_str)
        edited["InjectVar"] = edited["InjectVar"].apply(_safe_str).replace({"": INJECT_NONE}).fillna(INJECT_NONE)
        edited["Order"] = pd.to_numeric(edited["Order"], errors="coerce").fillna(10_000)
        edited = edited[edited["SlotVar"] != ""].copy()
        st.session_state.m6_table_df = edited
        st.success("Order saved.")

    st.divider()

    # ---------- helpers for weighted counts ----------
    def _as_weight_series(local_df: pd.DataFrame, wcol: str | None) -> pd.Series:
        if wcol and (wcol in local_df.columns):
            w = pd.to_numeric(local_df[wcol], errors="coerce").fillna(0.0)
            return w
        return pd.Series(np.ones(len(local_df), dtype=float), index=local_df.index)

    def _weighted_dist(local_df: pd.DataFrame, var: str, wcol: str | None):
        if var not in local_df.columns:
            return 0.0, {}
        w = _as_weight_series(local_df, wcol)
        s = local_df[var].astype(str)
        # treat blanks as NaN-like
        s = s.replace({"nan": np.nan, "": np.nan, "None": np.nan})
        mask = s.notna()
        # If we know the valid coded responses for this variable, treat anything
        # outside that code list as "not asked / not in universe" for denominator.
        choice_map = _choices_dict_for(var)
        if choice_map:
            valid = set(choice_map.keys())
            mask = mask & s.isin(valid)
        if mask.sum() == 0:
            return 0.0, {}
        denom = float(w[mask].sum())
        by_code = (w[mask].groupby(s[mask]).sum()).to_dict()
        # keys are strings
        return denom, {str(k): float(v) for k, v in by_code.items()}

    def _choices_dict_for(var: str) -> dict:
        d = (var_catalog.get(var, {}) if isinstance(var_catalog, dict) else {}) or {}
        ch = d.get("choices", {}) or {}
        if isinstance(ch, dict):
            return {str(k): str(v) for k, v in ch.items()}
        if isinstance(ch, list):
            out = {}
            for item in ch:
                if isinstance(item, dict):
                    out[str(item.get("code","")).strip()] = str(item.get("label","")).strip()
            return {k:v for k,v in out.items() if k}
        return {}

    def _var_label(var: str) -> str:
        d = (var_catalog.get(var, {}) if isinstance(var_catalog, dict) else {}) or {}
        lab = str(d.get("label","") or "").strip()
        return lab or var

    # ---------------- Preview (Module 5 vibe, but filled) ----------------
    st.subheader("Preview (filled)")
    st.caption("Shows weighted counts and percents (unweighted if no weight selected).")

    # build ordered vars list including inject vars, in the exact visible order
    cur = st.session_state.get("m6_table_df", pd.DataFrame()).copy()
    if not cur.empty:
        cur["Order"] = pd.to_numeric(cur["Order"], errors="coerce").fillna(10_000)
        cur = cur.sort_values(["Order", "SlotVar"], ascending=[True, True])

    ordered_vars = []
    for _, r in cur.iterrows():
        sv = _safe_str(r.get("SlotVar"))
        iv = _safe_str(r.get("InjectVar"))
        # Prefer showing the injected/recoded table ABOVE the base variable table
        # (e.g., QAGE2 should appear before QAGE when injected).
        if iv and iv != INJECT_NONE:
            ordered_vars.append(iv)
        if sv:
            ordered_vars.append(sv)

    # For SPSS syntax only, list base variable first, then the injected variable right after.
    spss_vars = []
    for _, r in cur.iterrows():
        sv = _safe_str(r.get("SlotVar"))
        iv = _safe_str(r.get("InjectVar"))
        if sv:
            spss_vars.append(sv)
        if iv and iv != INJECT_NONE:
            spss_vars.append(iv)

    # Remove duplicates but keep first occurrence
    seen = set()
    ordered_vars = [v for v in ordered_vars if not (v in seen or seen.add(v))]

    # De-dupe SPSS-only variable order (keep first occurrence)
    seen_spss = set()
    spss_vars = [v for v in spss_vars if not (v in seen_spss or seen_spss.add(v))]

    if not ordered_vars:
        st.info("Add at least one variable in the table above, then Save + Preview.")
        return

    # Pick a variable to preview (keeps the UI snappy)
    prev_var = st.selectbox("Preview variable", options=ordered_vars, key="m6_preview_var")
    denom, by_code = _weighted_dist(df, prev_var, weight_col)
    choice_map = _choices_dict_for(prev_var)

    # build display table in SPSS-like order: known codes, then any extra codes in data
    codes = list(choice_map.keys())
    extras = [c for c in by_code.keys() if c not in choice_map]
    all_codes = codes + sorted(extras)

    prev_rows = []
    for code in all_codes:
        wcount = by_code.get(code, 0.0)
        pct = (wcount / denom * 100.0) if denom > 0 else np.nan
        prev_rows.append({
            "Code": code,
            "Label": choice_map.get(code, code),
            "Count": float(wcount),
            "Percent": float(pct) if pct == pct else np.nan,
        })
    preview_df = pd.DataFrame(prev_rows)
    st.dataframe(preview_df, use_container_width=True, hide_index=True)

    st.divider()

    # ---------------- Exports ----------------
    st.subheader("Build + download")

    # SPSS syntax
    def _m6_build_spss_freq_syntax(var_list: list[str], wcol: str | None) -> str:
        """Return SPSS syntax in the exact style you asked for.

        Example:
            filter off.
            weight off.
            weight by wvar9.

            freq
            QREG
            QAGE
            .

            EXECUTE.
        """
        vars_clean = [v for v in (var_list or []) if str(v).strip()]
        if not vars_clean:
            return "* No variables selected.\n"

        lines: list[str] = []
        lines.append("filter off.")
        lines.append("weight off.")
        if wcol:
            lines.append(f"weight by {wcol}.")
        lines.append("")
        lines.append("freq")
        for v in vars_clean:
            lines.append(str(v).strip())
        lines.append("")
        lines.append("EXECUTE.")
        return "\n".join(lines) + "\n"

    # Excel SPSS-like output (block layout like SPSS viewer)
    def _m6_build_freq_excel(local_df: pd.DataFrame, var_list: list[str], wcol: str | None) -> BytesIO:
     

        wb = Workbook()
        ws = wb.active
        ws.title = "Frequencies"

        # Columns: A=Section (Valid), B=Label, C=Frequency, D=Percent, E=Valid Percent, F=Cumulative Percent
        col_widths = {1: 10, 2: 44, 3: 12, 4: 12, 5: 14, 6: 18}
        for c, w in col_widths.items():
            ws.column_dimensions[get_column_letter(c)].width = w

        def _is_missing(x) -> bool:
            if x is None:
                return True
            s = str(x).strip()
            return s == "" or s.lower() in {"nan", "none"}

        def _weighted_counts(local_df2: pd.DataFrame, var: str, wcol2: str | None):
            if var not in local_df2.columns:
                return 0.0, 0.0, {}
            w = _as_weight_series(local_df2, wcol2)
            s = local_df2[var]
            s_str = s.astype(str)
            miss = s_str.apply(_is_missing)
            total_all = float(w.sum())
            total_valid = float(w[~miss].sum())
            if total_valid <= 0:
                return total_all, total_valid, {}
            by_code = (w[~miss].groupby(s_str[~miss]).sum()).to_dict()
            return total_all, total_valid, {str(k): float(v) for k, v in by_code.items()}

        r = 1
        for v in [str(x).strip() for x in (var_list or []) if str(x).strip()]:
            if v not in local_df.columns:
                continue

            # Title (variable label)
            ws.cell(row=r, column=1, value=_var_label(v) or v)
            r += 1

            # Header row
            ws.cell(row=r, column=3, value="Frequency")
            ws.cell(row=r, column=4, value="Percent")
            ws.cell(row=r, column=5, value="Valid Percent")
            ws.cell(row=r, column=6, value="Cumulative Percent")
            r += 1

            total_all, total_valid, by_code_v = _weighted_counts(local_df, v, wcol)
            ch_map = _choices_dict_for(v)
            codes_v = list(ch_map.keys())
            extras_v = [c for c in by_code_v.keys() if c not in ch_map]
            all_codes_v = codes_v + sorted(extras_v)

            cum = 0.0
            first_valid_written = False
            for code in all_codes_v:
                freq = float(by_code_v.get(code, 0.0))
                if freq == 0:
                    # SPSS viewer generally suppresses 0s; keep behavior tidy
                    continue

                pct = (freq / total_all * 100.0) if total_all > 0 else np.nan
                v_pct = (freq / total_valid * 100.0) if total_valid > 0 else np.nan
                cum = cum + (v_pct if v_pct == v_pct else 0.0)

                ws.cell(row=r, column=1, value="Valid" if not first_valid_written else "")
                ws.cell(row=r, column=2, value=ch_map.get(code, code))
                ws.cell(row=r, column=3, value=round(freq, 0))
                ws.cell(row=r, column=4, value=round(pct, 1) if pct == pct else "")
                ws.cell(row=r, column=5, value=round(v_pct, 1) if v_pct == v_pct else "")
                ws.cell(row=r, column=6, value=round(cum, 1) if v_pct == v_pct else "")
                first_valid_written = True
                r += 1

            # Total row
            ws.cell(row=r, column=2, value="Total")
            ws.cell(row=r, column=3, value=round(total_valid, 0))
            ws.cell(row=r, column=4, value=round((total_valid / total_all * 100.0), 1) if total_all > 0 else "")
            ws.cell(row=r, column=5, value=100.0 if total_valid > 0 else "")
            r += 2  # blank row between sections

        bio = BytesIO()
        wb.save(bio)
        bio.seek(0)
        return bio

    # Build docx using existing planner, but ensure derived/recode vars have qdefs w/ list-style choices
    # (Planner expects question-style choice lists.)
    def _augment_questions_for_planner(qs: list[dict], cat: dict) -> list[dict]:
        existing = {str(q.get("qname","")).strip() for q in qs if str(q.get("qname","")).strip()}
        out = list(qs)
        if not isinstance(cat, dict):
            return out
        for qn, meta in cat.items():
            qn = str(qn).strip()
            if not qn or qn in existing:
                continue
            label = str((meta or {}).get("label","") or "").strip() or qn
            ch = (meta or {}).get("choices", {}) or {}
            choices_list = []
            if isinstance(ch, dict):
                for code, lab in ch.items():
                    code_s = str(code).strip()
                    if not code_s:
                        continue
                    choices_list.append({"code": code_s, "label": str(lab).strip() or code_s})
            elif isinstance(ch, list):
                for item in ch:
                    if isinstance(item, dict) and str(item.get("code","")).strip():
                        choices_list.append({"code": str(item.get("code")).strip(), "label": str(item.get("label","")).strip()})
            out.append({"qname": qn, "label": label, "choices": choices_list})
        return out

    # Buttons
    colA, colB, colC = st.columns(3)
    with colA:
        build_docx = st.button("🧮 Build topline (.docx)", use_container_width=True, key="m6_build_docx_like_m5")
    with colB:
        spss_txt = _m6_build_spss_freq_syntax(spss_vars, weight_col)
        st.download_button(
            "⬇️ SPSS freq syntax (.sps)",
            data=spss_txt.encode("utf-8"),
            file_name="module6_frequencies.sps",
            mime="text/plain",
            use_container_width=True,
            key="m6_dl_spss",
        )
    with colC:
        xbio = _m6_build_freq_excel(df, ordered_vars, weight_col)
        st.download_button(
            "⬇️ Frequencies (.xlsx)",
            data=xbio.getvalue(),
            file_name="module6_frequencies.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            use_container_width=True,
            key="m6_dl_xlsx",
        )

    if build_docx:
        # Build a Word doc that matches Module 5 formatting EXACTLY,
        # but with "Count Percent" filled into the % cells (rounded).
        questions_for_plan = _augment_questions_for_planner(questions, var_catalog)
        # IMPORTANT: Use the SlotVar/InjectVar table (like Module 5) so injected recodes
        # appear within the same question block (totals first), instead of becoming a
        # separate question (which was happening for QAGE2).
        plan = _build_rows_plan_from_table(questions_for_plan, recodes, st.session_state.get("m6_table_df", pd.DataFrame()))

        # Precompute weighted dists for all vars we might need
        vars_needed = sorted({str(v).strip() for v in plan["Var"].unique() if str(v).strip() and not str(v).strip().startswith("__")})

        # Resolve columns case-insensitively (prevents "couldn't find numbers" when dataset casing differs)
        _col_map = {str(c).strip().lower(): c for c in df.columns}

        def _resolve_col(name: str) -> str | None:
            key = str(name).strip().lower()
            return _col_map.get(key)

        dist_cache = {}  # var -> (denom_w, by_code_w)
        for v in vars_needed:
            col = _resolve_col(v)
            dist_cache[v] = _weighted_dist(df, col, weight_col) if col else (0.0, {})

        # Image stats
        image_stats_by_var = {}
        for v in st.session_state.get("m6_qimage_vars", []) or []:
            col = _resolve_col(v)
            image_stats_by_var[v] = _weighted_dist(df, col, weight_col) if col else (0.0, {})

        def _fmt_count_pct(count: float, denom: float) -> str:
            c = 0.0 if count is None else float(count)
            d = 0.0 if denom is None else float(denom)
            pct = (c / d * 100.0) if d > 0 else 0.0
            return _fmt_pct(pct)

        def _insert_qimage_table_filled(doc: Document, image_labels: list[str], image_vars: list[str]):
            headers = ["", "Total fav", "Total unfav", "Very fav", "Smwt fav",
                       "Smwt unfav", "Very unfav", "No opin", "NHO", "Net fav"]

            table = doc.add_table(rows=1, cols=len(headers))
            table.autofit = False
            table.alignment = WD_ALIGN_PARAGRAPH.CENTER

            col_widths = [1.55, 0.55, 0.56, 0.55, 0.55, 0.55, 0.53, 0.47, 0.51, 0.55]
            for i, w in enumerate(col_widths):
                table.columns[i].width = Inches(w)

            _remove_borders(table)

            hdr_cells = table.rows[0].cells
            for i, h in enumerate(headers):
                hdr_cells[i].text = h
                col = GREEN if h == "Total fav" else (ORANGE if h == "Total unfav" else None)
                _set_font_cell(hdr_cells[i], bold=True, color=col)
                for para in hdr_cells[i].paragraphs:
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _set_row_height(table.rows[0])

            for idx, (v, nm) in enumerate(zip(image_vars, image_labels)):
                denom_v, by_code_v = image_stats_by_var.get(v, (0.0, {}))
                fav = float(by_code_v.get("1", 0.0)) + float(by_code_v.get("2", 0.0))
                unf = float(by_code_v.get("3", 0.0)) + float(by_code_v.get("4", 0.0))
                very_fav = float(by_code_v.get("1", 0.0))
                smwt_fav = float(by_code_v.get("2", 0.0))
                smwt_unf = float(by_code_v.get("3", 0.0))
                very_unf = float(by_code_v.get("4", 0.0))
                no_op = float(by_code_v.get("5", 0.0))
                nho = float(by_code_v.get("6", 0.0))
                net = fav - unf

                vals = [None, fav, unf, very_fav, smwt_fav, smwt_unf, very_unf, no_op, nho, net]

                row_cells = table.add_row().cells
                row_cells[0].text = str(nm)
                _set_font_cell(row_cells[0])

                for j in range(1, len(headers)):
                    row_cells[j].text = _fmt_count_pct(vals[j], denom_v)
                    bold = True if j in [1, 2] else False
                    col = GREEN if j == 1 else (ORANGE if j == 2 else None)
                    _set_font_cell(row_cells[j], bold=bold, color=col)
                    for para in row_cells[j].paragraphs:
                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER

                if idx % 2 != 0:
                    _shade_row(row_cells, "F2F2F2")
                _set_row_height(table.rows[-1])

        # ---- Build the Word doc using Module 5's exact structure ----
        doc = Document()

        # Header (exactly Module 5)
        if any([survey_title.strip(), sample_n.strip(), field_dates.strip()]):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(8)
            p.paragraph_format.line_spacing = 1.0

            r = p.add_run(survey_title.strip() if survey_title.strip() else "Topline")
            r.bold = True
            r.font.name = "Arial"
            r.font.size = Pt(12)
            r._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")

            meta_parts = []
            if sample_n.strip():
                meta_parts.append(f"N = {sample_n.strip()}")
            if field_dates.strip():
                meta_parts.append(field_dates.strip())

            if meta_parts:
                p2 = doc.add_paragraph()
                p2.paragraph_format.space_before = Pt(0)
                p2.paragraph_format.space_after = Pt(12)
                p2.paragraph_format.line_spacing = 1.0
                r2 = p2.add_run(" • ".join(meta_parts))
                r2.font.name = "Arial"
                r2.font.size = Pt(11)
                r2._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")

        qimage_labels = st.session_state.get("m6_qimage_labels", []) or []
        qimage_vars = st.session_state.get("m6_qimage_vars", []) or []
        current_table = None
        qimage_inserted = False

        current_var = None
        injected_seen = False
        spacer_added = False

        for _, r in plan.iterrows():
            row_type = _safe_str(r.get("RowType"))
            text = _safe_str(r.get("Text"))
            bold = bool(r.get("Bold", False))
            var = _safe_str(r.get("Var"))

            if row_type == "qimage_header":
                _add_question_paragraph(doc, "Images", IMAGES_PROMPT_TEXT)
                current_table = None
                continue

            if row_type == "qimage_placeholder":
                if (not qimage_inserted) and qimage_labels and qimage_vars:
                    _insert_qimage_table_filled(doc, qimage_labels, qimage_vars)
                qimage_inserted = True
                continue

            if row_type == "header":
                current_table = None
                current_var = var
                injected_seen = False
                spacer_added = False

                t = _clean_text(text)
                parts = [p.strip() for p in t.split(":", 2)]
                if len(parts) == 3:
                    _, lab, prm = parts
                elif len(parts) == 2:
                    lab, prm = parts
                else:
                    lab, prm = t, ""
                _add_question_paragraph(doc, lab, prm)
                current_table = _make_2col_table(doc)
                continue

            if current_table is None:
                continue

            if row_type == "blank":
                current_table = None
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(6)
                p.paragraph_format.line_spacing = 0
                continue

            if row_type == "choice_group":
                injected_seen = True

            # Insert the same spacer row Module 5 uses after injected rows (once),
            # right before the first non-bold normal option.
            if (not bold) and injected_seen and (not spacer_added) and row_type in ("choice",):
                _m5_add_option_row(current_table, "", placeholder="", bold=False)
                spacer_added = True

            if row_type in ("choice", "choice_group"):
                denom_v, by_code_v = dist_cache.get(current_var, (0.0, {}))

                if row_type == "choice_group":
                    gc = _safe_str(r.get("GroupCodes"))
                    codes = [c.strip() for c in gc.split("|") if c.strip()]
                    count = float(sum(float(by_code_v.get(c, 0.0)) for c in codes))
                else:
                    code = _safe_str(r.get("Value"))
                    count = float(by_code_v.get(code, 0.0)) if code else 0.0

                placeholder = _fmt_count_pct(count, denom_v)
                _m5_add_option_row(current_table, text, placeholder=placeholder, bold=bold)

        if (not qimage_inserted) and qimage_labels and qimage_vars:
            _add_question_paragraph(doc, "Images", IMAGES_PROMPT_TEXT)
            _insert_qimage_table_filled(doc, qimage_labels, qimage_vars)

        bio = BytesIO()
        doc.save(bio)
        bio.seek(0)

        st.download_button(
            "⬇️ Download topline (.docx)",
            data=bio.getvalue(),
            file_name="module6_topline_filled.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
            key="m6_dl_docx",
        )

def _var_label(var: str) -> str:
    labels = st.session_state.get("var_labels", {})  # optional: {var: "Label"}
    return labels.get(var, var)

def _val_label(var: str, val) -> str:
    """Return the display label for a coded value.

    Priority:
      1) st.session_state.value_labels (if user provided)
      2) st.session_state.var_catalog[var]['choices'] (from Modules 1/2/2.5)
      3) fallback to stringified value
    """
    # --- 1) explicit mapping ---
    vlabels = st.session_state.get("value_labels", {}) or {}
    d = vlabels.get(var, {}) or {}

    # --- 2) var_catalog choices mapping ---
    vc = st.session_state.get("var_catalog", {}) or {}
    choices = {}
    try:
        choices = (vc.get(var, {}) or {}).get("choices", {}) or {}
    except Exception:
        choices = {}

    def _lookup(mapping, key):
        if key in mapping:
            return mapping[key]
        # try common numeric/string variants
        try:
            if isinstance(key, str):
                ks = key.strip()
                if ks in mapping:
                    return mapping[ks]
                f = float(ks)
            else:
                f = float(key)
            if f in mapping:
                return mapping[f]
            if int(f) in mapping:
                return mapping[int(f)]
            if str(int(f)) in mapping:
                return mapping[str(int(f))]
            if str(f) in mapping:
                return mapping[str(f)]
        except Exception:
            pass
        return None

    out = _lookup(d, val)
    if out is None:
        out = _lookup(choices, val)

    if out is None:
        return "" if pd.isna(val) else str(val)
    return str(out).strip()

def _ordered_categories(series: pd.Series) -> list:
    if pd.api.types.is_categorical_dtype(series):
        return list(series.cat.categories)
    vals = [v for v in series.dropna().unique().tolist()]
    try:
        return sorted(vals, key=lambda x: float(x))
    except Exception:
        return sorted(vals, key=lambda x: str(x).lower())


# -------------------------
# Weighted math helpers
# -------------------------
def _weighted_total(w: pd.Series) -> float:
    return float(np.nansum(w.values))

def _weighted_sum(mask: np.ndarray, w: pd.Series) -> float:
    return float(np.nansum(w.values[mask]))

def _pct(n: float, d: float) -> float:
    if d == 0 or np.isnan(d):
        return np.nan
    return n / d


def _build_unweighted_counts(df: pd.DataFrame, col_groups: list[tuple[str, list]], resolve_map: dict | None = None) -> dict:
    out = {}
    out[("Total", "Total")] = float(len(df))
    for demo_var, cats in col_groups:
        s = _get_series_safe(df, _m7_resolve_name(demo_var, resolve_map))
        for c in cats:
            out[(demo_var, c)] = float((s == c).sum())
    return out

def _build_counts(df: pd.DataFrame, col_groups: list[tuple[str, list]], weight_col: str | None, resolve_map: dict | None = None) -> dict:
    """Counts for the header count row.

    If weight_col is provided, returns weighted counts (sum of weights).
    Otherwise, returns unweighted counts (n).
    """
    if weight_col is None:
        return _build_unweighted_counts(df, col_groups, resolve_map=resolve_map)

    w = pd.to_numeric(_get_series_safe(df, weight_col), errors="coerce").fillna(0.0)
    out = {}
    out[("Total", "Total")] = float(w.sum())
    for demo_var, cats in col_groups:
        s = _get_series_safe(df, _m7_resolve_name(demo_var, resolve_map))
        for c in cats:
            out[(demo_var, c)] = float(w[s == c].sum())
    return out



def _build_pcts_for_var(
    df: pd.DataFrame,
    var: str,
    col_groups: list[tuple[str, list]],
    weight_col: str | None,
    resolve_map: dict | None = None,
) -> pd.DataFrame:
    """
    Returns DF:
      index = categories of `var`
      columns = MultiIndex [("Total","Total"), (demo_var, demo_cat)...]
      values = column % (weighted or unweighted)
    """
    s_var = _get_series_safe(df, _m7_resolve_name(var, resolve_map))
    rcats = _ordered_categories(s_var)

    cols = [("Total", "Total")]
    for dv, cats in col_groups:
        for c in cats:
            cols.append((dv, c))
    col_index = pd.MultiIndex.from_tuples(cols, names=["Group", "Category"])

    out = pd.DataFrame(index=rcats, columns=col_index, dtype=float)

    if weight_col is None:
        denom_total = len(df)
        for r in rcats:
            out.loc[r, ("Total", "Total")] = _pct((s_var == r).sum(), denom_total)

        for dv, cats in col_groups:
            s_demo = _get_series_safe(df, _m7_resolve_name(dv, resolve_map))
            for c in cats:
                sub = df[s_demo == c]
                denom = len(sub)
                for r in rcats:
                    out.loc[r, (dv, c)] = _pct((_get_series_safe(sub, _m7_resolve_name(var, resolve_map)) == r).sum(), denom)
    else:
        w = pd.to_numeric(_get_series_safe(df, weight_col), errors="coerce").fillna(0.0)
        denom_total = _weighted_total(w)
        var_vals = s_var.values

        for r in rcats:
            out.loc[r, ("Total", "Total")] = _pct(_weighted_sum((var_vals == r), w), denom_total)

        for dv, cats in col_groups:
            demo_vals = _get_series_safe(df, _m7_resolve_name(dv, resolve_map)).values
            for c in cats:
                mask_col = (demo_vals == c)
                denom = _weighted_sum(mask_col, w)
                for r in rcats:
                    mask = mask_col & (var_vals == r)
                    out.loc[r, (dv, c)] = _pct(_weighted_sum(mask, w), denom)

    return out


# -------------------------
# Reorder UI (Streamlit)
# -------------------------
def _reorder_list_ui(title: str, key: str, options: list[str], default: list[str] | None = None) -> list[str]:
    """
    Choose + reorder items using a single-select "Up/Down".
    """
    if default is None:
        default = []

    order_key = f"{key}_order"
    sel_key = f"{key}_sel"
    chosen_key = f"{key}_chosen"

    if order_key not in st.session_state:
        st.session_state[order_key] = default[:] if default else options[:]

    # keep valid + add missing
    st.session_state[order_key] = [x for x in st.session_state[order_key] if x in options]
    for x in options:
        if x not in st.session_state[order_key]:
            st.session_state[order_key].append(x)

    st.markdown(f"**{title}**")
    chosen = st.multiselect(
        "Pick items to include:",
        options=st.session_state[order_key],
        default=[x for x in st.session_state[order_key] if (x in default)] if default else st.session_state[order_key],
        key=chosen_key,
    )

    # item to move
    move_item = st.selectbox("Select one item to move:", options=chosen if chosen else ["(none)"], key=sel_key)

    c1, c2 = st.columns(2)
    with c1:
        if st.button("⬆️ Move up", key=f"{key}_up"):
            if move_item != "(none)":
                order = st.session_state[order_key]
                i = order.index(move_item)
                if i > 0:
                    order.pop(i)
                    order.insert(i - 1, move_item)
                    st.session_state[order_key] = order
    with c2:
        if st.button("⬇️ Move down", key=f"{key}_down"):
            if move_item != "(none)":
                order = st.session_state[order_key]
                i = order.index(move_item)
                if i < len(order) - 1:
                    order.pop(i)
                    order.insert(i + 1, move_item)
                    st.session_state[order_key] = order

    final = [x for x in st.session_state[order_key] if x in chosen]
    return final


# -------------------------
# Formatting (adapted from your code, Streamlit-safe)
# -------------------------



def _ordered_checkbox_box(options: list[str], key: str) -> list[str]:
    """
    Render a large checkbox "box" showing every variable, unchecked by default.
    The order a user checks variables becomes the selection order returned.
    """
    import hashlib

    order_key = f"{key}_order"
    prev_key = f"{key}_prev"

    if order_key not in st.session_state:
        st.session_state[order_key] = []
    if prev_key not in st.session_state:
        st.session_state[prev_key] = {}

    search = st.text_input("Search", value="", key=f"{key}_search", placeholder="Type to filter variables…")
    shown = [v for v in options if (search.lower() in str(v).lower())] if search else list(options)

    newly_checked: list[str] = []
    newly_unchecked: list[str] = []

    with st.container(border=True):
        st.caption(f"Showing {len(shown):,} of {len(options):,} variables")
        ncols = 2
        cols = st.columns(ncols, gap="small")
        for i, v in enumerate(shown):
            v_key = hashlib.md5(str(v).encode("utf-8")).hexdigest()[:10]
            cb_key = f"{key}_cb_{v_key}"

            # Initialize from current order list on first render
            if cb_key not in st.session_state:
                st.session_state[cb_key] = (v in st.session_state[order_key])

            with cols[i % ncols]:
                val = st.checkbox(str(v), key=cb_key)

            prev = st.session_state[prev_key].get(str(v), (v in st.session_state[order_key]))
            if val and not prev:
                newly_checked.append(v)
            elif (not val) and prev:
                newly_unchecked.append(v)

            st.session_state[prev_key][str(v)] = val

    order = [x for x in st.session_state[order_key] if x not in newly_unchecked]
    for v in newly_checked:
        if v not in order:
            order.append(v)

    st.session_state[order_key] = order

    if order:
        st.caption("Selected order:")
        st.code(" , ".join([str(x) for x in order]), language=None)

    return order


def standardize_formatting(ws):
    font = Font(name="Arial", size=12, color="000000")
    for row in ws.iter_rows():
        for cell in row:
            cell.font = font
            cell.fill = PatternFill(fill_type=None)

def restore_standard_borders(ws):
    thin = Side(border_style="thin", color="D9D9D9")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)
    for row in ws.iter_rows():
        for cell in row:
            cell.border = border

def shade_column_groups(ws):
    """Alternate shading by column GROUPS, starting at Column C.

    - Column A and B are label columns and should never be shaded.
    - Column C (Total) is treated as the first group and should be shaded.
    """
    header_row = 1

    col_groups = {}
    current_group = None

    # Only consider numeric columns starting at C
    for col in range(3, ws.max_column + 1):
        v = ws.cell(row=header_row, column=col).value
        if v not in [None, ""]:
            current_group = str(v).strip()
        if current_group:
            col_groups.setdefault(current_group, []).append(col)

    fill = PatternFill(start_color="F2F2F2", end_color="F2F2F2", fill_type="solid")
    for i, cols in enumerate(col_groups.values()):
        # First group (Total) is i=0 -> shaded
        if i % 2 == 0:
            for col in cols:
                for row in range(1, ws.max_row + 1):
                    ws.cell(row=row, column=col).fill = fill


def merge_and_center_total(ws):
    total_col = None
    for col in range(1, ws.max_column + 1):
        if ws.cell(row=1, column=col).value == "Total":
            total_col = col
            break
    if total_col:
        ws.merge_cells(start_row=1, start_column=total_col, end_row=2, end_column=total_col)
        ws.cell(row=1, column=total_col).alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)

def merge_header_sections(ws):
    col = 1
    while col <= ws.max_column:
        header_cell = ws.cell(row=1, column=col)
        if header_cell.value and str(header_cell.value).strip() != "":
            start_col = col
            end_col = col
            while end_col + 1 <= ws.max_column:
                next_header = ws.cell(row=1, column=end_col + 1).value
                below_next = ws.cell(row=2, column=end_col + 1).value
                if next_header in [None, ""] and below_next not in [None, ""]:
                    end_col += 1
                else:
                    break
            if end_col > start_col:
                ws.merge_cells(start_row=1, start_column=start_col, end_row=1, end_column=end_col)
                ws.cell(row=1, column=start_col).alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
            col = end_col + 1
        else:
            col += 1

def shade_row_groups(ws, start_row=3):
    """
    Alternating shading by blocks of identical Column A label.
    (Replaces fuzzywuzzy + popup with deterministic behavior.)
    """
    fill = PatternFill(start_color="FFB4B4", end_color="FFB4B4", fill_type="solid")
    current_label = None
    block_index = -1

    for row in range(start_row, ws.max_row + 1):
        v = ws.cell(row=row, column=1).value
        label = str(v).strip() if v not in [None, ""] else None

        # Only start a new block when a non-empty label changes
        if label is not None and label != current_label:
            current_label = label
            block_index += 1

        if block_index >= 0 and block_index % 2 == 0:
            for col in range(1, ws.max_column + 1):
                ws.cell(row=row, column=col).fill = fill

def apply_custom_layout(ws):
    # Merge A1:B2 (title)
    ws.merge_cells("A1:B2")
    # Merge A3:B3
    ws.merge_cells("A3:B3")

    ws.row_dimensions[1].height = 38

    ws["A1"].alignment = Alignment(vertical="center", wrap_text=True)

    # center row 2 headers from column C onward
    for col in range(3, ws.max_column + 1):
        ws.cell(row=2, column=col).alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)

    # vertically center Column A data
    for row in range(3, ws.max_row + 1):
        ws.cell(row=row, column=1).alignment = Alignment(vertical="center", wrap_text=True)

    ws.column_dimensions["A"].width = 30
    ws.column_dimensions["B"].width = 30

def set_numeric_column_widths(ws):
    for col in range(3, ws.max_column + 1):
        ws.column_dimensions[get_column_letter(col)].width = 13.4

def format_percent_cells(ws, data_start_row=4, data_start_col=3):
    # Format numeric cells as percent where appropriate (skip counts row 3)
    # Align the crosstab body: right (horizontal) + center (vertical).
    pct_align = Alignment(horizontal="right", vertical="center")
    for r in range(data_start_row, ws.max_row + 1):
        for c in range(data_start_col, ws.max_column + 1):
            cell = ws.cell(row=r, column=c)

            # Center alignment for numeric + placeholder cells
            if cell.value is not None and (isinstance(cell.value, (int, float)) or str(cell.value).strip() in {"—", "-"}):
                cell.alignment = pct_align

            if isinstance(cell.value, (int, float)) and cell.value is not None and not np.isnan(cell.value):
                # heuristic: percentages are between 0 and 1; counts are > 1
                if 0 <= float(cell.value) <= 1:
                    cell.number_format = "0.0%"
def apply_module7_formatting(ws, shaded: bool = True):
    standardize_formatting(ws)
    if shaded:
        shade_column_groups(ws)
    merge_and_center_total(ws)
    merge_header_sections(ws)
    if shaded:
        shade_row_groups(ws, start_row=3)
    restore_standard_borders(ws)
    apply_custom_layout(ws)
    set_numeric_column_widths(ws)
    format_percent_cells(ws, data_start_row=4, data_start_col=3)


# =========================================================
# MODULE 7: Crosstabs (Excel output)
#
# Purpose:
# - Build crosstab tables from a (weighted or unweighted) dataset
# - Export a formatted Excel workbook matching your topline style
# - ALSO generate SPSS syntax to reproduce the same crosstabs
#
# Notes:
# - Percentages are column % (distribution within each column category)
# - QIMAGE_* blocks get extra rows:
#     ***Name ID = sum of all rows except "Never heard of"
#     ***Net     = Total favorable - Total unfavorable
#   (These rows are NOT added for cQ* variables.)
# =========================================================

# -------------------------
# Module 7 label helpers
# -------------------------


def _m7_var_label(var: str) -> str:
    """Get a human-friendly variable label for Module 7.

    Priority:
      1) st.session_state.var_labels (if present)
      2) st.session_state.var_catalog[var]['label']
      3) fallback to var name
    """
    v = (var or "").strip()
    # direct labels dict (some modules may store this)
    vlabels = st.session_state.get("var_labels", {}) or {}
    lab = str(vlabels.get(v, "") or "").strip()
    if lab:
        return lab
    # var_catalog label
    vc = st.session_state.get("var_catalog", {}) or {}
    try:
        lab2 = str((vc.get(v, {}) or {}).get("label", "") or "").strip()
        if lab2:
            return lab2
    except Exception:
        pass
    return v




def render_module_7():
    st.header("Module 7: Crosstabs (weighted Excel → formatted export)")

    # ----------------------------
    # Dataset source (Module 4 output OR uploaded file)
    # ----------------------------
    df = None
    source_options = ["Upload a dataset file"]
    if "m4_df_out" in st.session_state and st.session_state.m4_df_out is not None and isinstance(st.session_state.m4_df_out, pd.DataFrame):
        source_options = ["Use Module 4 output (in-app)", "Upload a dataset file"]

    dataset_source = st.radio(
        "Dataset source",
        options=source_options,
        index=0,
        horizontal=True,
        key="m7_dataset_source",
    )

    if dataset_source == "Use Module 4 output (in-app)":
        df = st.session_state.m4_df_out.copy()
        st.caption("Using Module 4 output dataframe from this session.")
    else:
        up = st.file_uploader("Upload dataset (.xlsx or .csv)", type=["xlsx", "csv"], key="m7_upload")
        if not up:
            st.info("Upload a dataset file to begin.")
            return

        if up.name.lower().endswith(".csv"):
            df = pd.read_csv(up)
        else:
            xls = pd.ExcelFile(up)
            sheet = st.selectbox("Sheet", xls.sheet_names, key="m7_sheet")
            df = pd.read_excel(up, sheet_name=sheet)

    if df is None or df.empty:
        st.warning("Dataset is empty.")
        return

    st.caption(f"Rows: {len(df):,} | Columns: {len(df.columns):,}")

    # Available columns
    all_cols = df.columns.tolist()

    # Hide helper/open-end columns and avoid duplicate fuzzy-matched columns
    # - drop columns ending in '__text'
    # - if a dataset column is a strong (>=80%) fuzzy match for a canonical var_catalog name,
    #   we show the canonical name in the pickers and hide the raw column name
    base_cols = [c for c in all_cols if not str(c).endswith('__text')]

    var_catalog = st.session_state.get('var_catalog', {}) or {}
    canonical_vars = [v for v in list(var_catalog.keys()) if v and not str(v).endswith('__text')]

    # Cached resolve map so checkbox clicking doesn't redo fuzzy matching work.
    resolve_map, matched_raw = _m7_build_resolve_map_cached(tuple(base_cols), tuple(sorted(canonical_vars)), threshold=0.80)

    # Options shown in the checkbox boxes:
    #   - canonical vars that exist in the dataset (via resolve_map)
    #   - plus any remaining dataset columns not claimed by a canonical fuzzy match
    m7_options = list(resolve_map.keys()) + [c for c in base_cols if c not in matched_raw]
    # Weight selector (include explicit no-weight)
    # Prefer Module 4's weight col when using Module 4 output
    default_weight = None
    if "m4_weight_col" in st.session_state and st.session_state.m4_weight_col in all_cols:
        default_weight = st.session_state.m4_weight_col

    weight_choices = ["(No weight)"] + all_cols
    default_w_index = 0
    if default_weight in weight_choices:
        default_w_index = weight_choices.index(default_weight)

    weight_choice = st.selectbox(
        "Weight variable",
        options=weight_choices,
        index=default_w_index,
        key="m7_weight_choice",
        help="Choose a weight variable to compute weighted crosstabs, or '(No weight)' for unweighted.",
    )
    weight_col = None if weight_choice == "(No weight)" else weight_choice
    st.subheader("Select variables")

    left, right = st.columns(2, gap="large")

    with left:
        st.markdown("### Columns (BY variables)")
        st.write("Click variables to check them. The order you check them becomes the left → right column order (Totals will be added automatically).")
        col_demo_vars = _ordered_checkbox_box(
            options=m7_options,
            key="m7_cols_box",
        )

    with right:
        st.markdown("### Rows (TABLE variables)")
        st.write("Click variables to check them. The order you check them becomes the top → bottom row order.")
        row_vars_raw = _ordered_checkbox_box(
            options=m7_options,
            key="m7_rows_box",
        )

    # Keep selections as-is.
    # If a variable is selected in both Columns and Rows, we keep it in BOTH;
    # during export we will simply skip self-pairs (row == column) when building tables.
    qvars = list(row_vars_raw)
    # Optional: demographic tables appended AFTER question tables.
    # (Earlier versions referenced row_demo_vars but never defined it.)
    row_demo_vars = []
    overlap = set(row_vars_raw) & set(col_demo_vars)
    if overlap:
        st.info("Some variables are selected in both Columns and Rows. They will remain selected in both; exports will include them as selected.")
    # NOTE: Don't compute category lists (and other heavy work) on every checkbox click.
    # We'll build col_groups only when Preview/Generate is pressed.
    col_groups = None

    st.subheader("Export settings")
    title_text = st.text_input(
        "Title (A1 block)",
        value="Georgia’s 14th Congressional District 2026 Election Research",
        key="m7_title",
    )
    subtitle_text = st.text_area(
        "Subtitle (included in A1 block under title)",
        value="December 2nd – December 4th, 2025\nCrosstabs\nMOE = +/- 4.9%",
        height=90,
        key="m7_subtitle",
    )

    col1, col2 = st.columns([1, 1])
    with col1:
        if st.button("Preview first question", use_container_width=True, key="m7_preview"):
            if not qvars:
                st.warning("Pick at least one question.")
            elif not col_demo_vars:
                st.warning("Pick at least one column demo.")
            else:
                col_groups = [(dv, _ordered_categories(_get_series_safe(df, _m7_resolve_name(dv, resolve_map)))) for dv in col_demo_vars]
                p = _build_pcts_for_var(df, qvars[0], col_groups, weight_col, resolve_map=resolve_map)
                st.write(f"Preview: {_var_label(qvars[0])}")
                st.dataframe(p)

    with col2:
        if st.button("Generate Crosstabs Excel", use_container_width=True, key="m7_export"):
            if not qvars:
                st.error("Select at least one question variable.")
                return
            if not col_demo_vars:
                st.error("Select at least one column demo variable.")
                return

            # Heavy work begins here (only on Generate)
            col_groups = [(dv, _ordered_categories(_get_series_safe(df, _m7_resolve_name(dv, resolve_map)))) for dv in col_demo_vars]

            wb = Workbook()
            ws = wb.active
            ws.title = "Crosstabs"

            # ---- Header block like your example ----
            ws["A1"] = f"{title_text}\n{subtitle_text}"
            ws["A1"].alignment = Alignment(wrap_text=True, vertical="top")

            # ---- Build 2-row header ----
            # Row 1 & 2 are headers; A/B reserved for left labels.
            # C is Total column.
            ws.cell(row=1, column=3, value="Total")
            ws.cell(row=2, column=3, value="Total")

            col_cursor = 4
            # Write demo group header (row 1) and category labels (row 2)
            for demo_var, cats in col_groups:
                ws.cell(row=1, column=col_cursor, value=_m7_var_label(demo_var))
                for i, c in enumerate(cats):
                    ws.cell(row=2, column=col_cursor + i, value=_val_label(demo_var, c))
                # merge group header across its categories
                if len(cats) > 1:
                    ws.merge_cells(start_row=1, start_column=col_cursor, end_row=1, end_column=col_cursor + len(cats) - 1)
                    ws.cell(row=1, column=col_cursor).alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
                col_cursor += len(cats)

            # ---- Count row (Row 3) ----
            counts = _build_counts(df, col_groups, weight_col, resolve_map=resolve_map)
            ws.cell(row=3, column=1, value=("Weighted Count" if weight_col is not None else "Unweighted Count"))
            ws.cell(row=3, column=2, value="")
            ws.cell(row=3, column=3, value=round(counts[("Total", "Total")], 0))

            cc = 4
            for demo_var, cats in col_groups:
                for c in cats:
                    ws.cell(row=3, column=cc, value=round(counts[(demo_var, c)], 0))
                    cc += 1

            # ---- Write blocks starting Row 4 ----
            cur_row = 4

            def write_block(varname: str, cur: int) -> int:
                # Build % table for this row variable across all selected column groups
                pct = _build_pcts_for_var(df, varname, col_groups, weight_col, resolve_map=resolve_map)

                # No separate title row. We'll merge the row-group label into Column A across the category rows.
                group_start = cur

                for cat in pct.index.tolist():
                    # Column A will be merged later; we can leave it blank here to avoid repeats
                    ws.cell(row=cur, column=1, value="")
                    lab = _val_label(varname, cat)
                    ws.cell(row=cur, column=2, value=lab if lab != "" else str(cat))

                    # Total col
                    vtot = pct.loc[cat, ("Total", "Total")]
                    ws.cell(row=cur, column=3, value=float(vtot) if pd.notna(vtot) else np.nan)

                    # demo cols (skip self-pairs: if a demo var equals the row var, don't build that panel)
                    cc2 = 4
                    for dv, cats2 in col_groups:
                        for c2 in cats2:
                            v = pct.loc[cat, (dv, c2)]
                            ws.cell(row=cur, column=cc2, value=float(v) if pd.notna(v) else np.nan)
                            cc2 += 1

                    cur += 1

                # Total row at end of block (intentional)
                # Total row: write COUNTS (denominators), not percents
                ws.cell(row=cur, column=1, value="")
                ws.cell(row=cur, column=2, value="Total")
                ws.cell(row=cur, column=3, value=float(round(counts.get(("Total", "Total"), np.nan), 0)) if pd.notna(counts.get(("Total", "Total"), np.nan)) else np.nan)
                cc3 = 4
                for dv, cats2 in col_groups:
                    for cat2 in cats2:
                        val_ct = counts.get((dv, cat2), np.nan)
                        ws.cell(row=cur, column=cc3, value=float(round(val_ct, 0)) if pd.notna(val_ct) else np.nan)
                        cc3 += 1
                total_row = cur
                cur += 1

                # Merge + center the row-group label in Column A (over category rows + Total row)
                group_end = total_row
                if group_end >= group_start:
                    ws.merge_cells(start_row=group_start, start_column=1, end_row=group_end, end_column=1)
                    c = ws.cell(row=group_start, column=1)
                    c.value = _m7_var_label(varname)
                    c.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)

                # (No spacer row; this removes the extra blank row you were seeing per group.)
                return cur

            # questions first
            for q in qvars:
                cur_row = write_block(q, cur_row)

            # demos last (row blocks appended at end)
            for d in row_demo_vars:
                cur_row = write_block(d, cur_row)

            # Remove final spacer if last row is blank
            # (we can just clear it; avoids an extra blank "table row" feel)
            last = ws.max_row
            if (ws.cell(row=last, column=1).value in [None, ""] and
                ws.cell(row=last, column=2).value in [None, ""]):
                # leave it, but it will be empty; your formatter shades by blocks so it's fine
                pass

            # Save a base copy (data + merges only), then export 2 variants:
            # 1) Full formatting with shading
            # 2) Same formatting but NO shading (clean version)

            base = BytesIO()
            wb.save(base)
            base.seek(0)

            # --- Shaded version ---
            wb_shaded = openpyxl.load_workbook(base)
            ws_shaded = wb_shaded.active
            apply_module7_formatting(ws_shaded, shaded=True)
            bio_shaded = BytesIO()
            wb_shaded.save(bio_shaded)
            bio_shaded.seek(0)

            
            # ----------------------------
            # SPSS syntax export (Module 7)
            # ----------------------------
            st.divider()

            st.download_button(
                "Download Module 7 Crosstabs (.xlsx)",
                data=bio_shaded,
                file_name="module7_crosstabs.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                use_container_width=True,
                key="m7_download_shaded",
            )

            # --- No-shading version ---
            base.seek(0)
            wb_clean = openpyxl.load_workbook(base)
            ws_clean = wb_clean.active
            apply_module7_formatting(ws_clean, shaded=False)
            bio_clean = BytesIO()
            wb_clean.save(bio_clean)
            bio_clean.seek(0)

            st.download_button(
                "Download Module 7 Crosstabs (no shading) (.xlsx)",
                data=bio_clean,
                file_name="module7_crosstabs_no_shading.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                use_container_width=True,
                key="m7_download_clean",
            )

















# ============================
# Render selected module ----------------------------------------------------never delete. keep here
# ============================
if active_module == "project":
    render_module_0()
elif active_module == "scripting":
    render_module_1()
elif active_module == "recodes":
    render_module_2() 
elif active_module == "derived_vars":
    render_module_25() 
elif active_module == "import_match":
    render_module_3()  
elif active_module == "weighting":
    render_module_4() 
elif active_module == "topline_shell":
    render_module_5() 
elif active_module == "weighted_topline":
    render_module_6()
elif active_module == "crosstabs":
    render_module_7()
